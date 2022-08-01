# -*- coding = utf-8 -*-
# @Software : PyCharm
from tkinter import messagebox
import matplotlib.pyplot as plt
import math
import numpy as np
import pandas as pd
import seaborn as sns
import scipy.stats as st
import scipy.integrate
import timeit
import pickle
import multiprocessing as mp
import smt.surrogate_models
import shutil
import glob
import os
from scipy.stats import qmc
import subprocess
from fitter import Fitter
from pathlib import Path
from scipy.optimize import fsolve
from smt.surrogate_models import RBF, IDW, RMTB, RMTC, LS, QP, KRG, KPLS, KPLSK
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from tkinter import *
import tkinter.ttk as ttk
import tkinter.filedialog
import warnings
import threading
import time
import re
import chardet

# 系统CPU核数
Num_CPU = int(mp.cpu_count() / 2)

txt = ""
btnAbaqus = ""

# fluent
FluentPath = ""
ScriptPath = ""
ScriptRootPath = ""
sema: int = 4
cpuCore = os.cpu_count()
def popupmenu(event):
    mainmenu.post(event.x_root, event.y_root)
    print("tight")

def ChFluentPath():
    global FluentPath
    FluentPath = tkinter.filedialog.askdirectory(initialdir='C:/Program Files/ANSYS Inc/')
    print(FluentPath)
    if FluentPath != '':
        lb1.config(text='Fluent路径：'+FluentPath)
    else:
        lb1.config(text='Fluent路径：未选择')

def ChScriptPath():
    global ScriptPath
    global ScriptRootPath
    ScriptPath = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目',filetypes=[("JOU", "jou")])
    if ScriptPath != '':
        wordlist = re.split('/', ScriptPath)
        ScriptRootPath = ""
        for i in range(0, len(wordlist) - 1):
            if i!=len(wordlist) - 2:
                ScriptRootPath += wordlist[i] + '/'
            else:
                print("!2")
                ScriptRootPath += wordlist[i]
        lb2.config(text='脚本文件：' + ScriptPath)
        lb3.config(text='脚本文件根目录路径：' + ScriptRootPath)
    else:
        lb2.config(text='脚本文件：未选择')

def ChScriptRootPath():
    global ScriptRootPath
    ScriptRootPath = tkinter.filedialog.askdirectory(initialdir='C:/Users/Leo/Desktop/WLB项目')
    if ScriptRootPath != '':
        lb3.config(text='脚本文件根目录路径：' + ScriptRootPath)
    else:
        lb3.config(text='脚本文件根目录路径：未选择')

def FluentCmdStart():
    txt.config(state=NORMAL)
    command = "exec.bat"
    ret = subprocess.Popen(command, stdin=subprocess.PIPE, stderr=subprocess.PIPE, shell=False, stdout=subprocess.PIPE)
    try:
        for i in iter(ret.stdout.readline, 'b'):
            # line_list = i.split('\r\n')
            if i:
                ret = chardet.detect(i)
                # 爬虫
                txt.insert(END, i.decode(ret['encoding']))
                txt.see(END)
                print(i.decode(ret['encoding']))
            else:
                break
        for i in iter(ret.stderr.readline, 'b'):
            # line_list = i.split('\r\n')
            if i:
                ret = chardet.detect(i)
                txt.insert(END, i.decode(ret['encoding']))
                txt.see(END)
                print(i.decode(ret['encoding']))
            else:
                break
    finally:
        txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 结束\n")
        txt.see(END)
        print(time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 结束")
        txt.config(state=DISABLED)
        btn2.config(state=NORMAL)

def Start():
    # txt.delete(1.0, END)
    global txt
    txt = txtFluent
    btnStatic2.config(state=DISABLED)
    txt.config(state=NORMAL)
    txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 开始分析\n")
    file_handle = open('exec.bat', 'w')
    file_handle.write('@ set PATH=' + FluentPath + '/;%PATH%\n')
    file_handle.write('@ cd ' + ScriptRootPath + '\n')
    if DelCheckVar.get() == 1:
        file_handle.write('@ fluent -wait 3ddp -t'+str(sema)+' -i ' + ScriptPath + '\n')
    else:
        file_handle.write('@ fluent -hidden -wait 3ddp -t'+str(sema)+' -i ' + ScriptPath + '\n')
    file_handle.write('@ del *.trn\n')
    file_handle.close()
    try:
        t1 = threading.Thread(target=FluentCmdStart)
        t1.setDaemon(True)
        t1.start()
    except:
        txt.insert(END, '开始失败\n')
        txt.config(state=DISABLED)
        btn2.config(state=NORMAL)

def pop_up_box():
    global sema
    # global btn4
    winNew = Toplevel(root)
    winNew.geometry('300x100')
    winNew.title('设置核数')
    def com():
        try:
            num = int(e1.get())  # 获取e1的值，转为浮点数，如果不能转捕获异常
            if(num>cpuCore):
                messagebox.showwarning('警告', '超出范围')
            else:
                sema = int(e1.get())
                lb3.config(text='计算核数'+str(sema)+'<计算核数'+str(cpuCore))
                winNew.destroy()
        except:
            messagebox.showwarning('警告', '请输入数字')
    var = tkinter.StringVar()  # 这即是输入框中的内容
    var.set(str(sema))
    e1 = Entry(winNew,textvariable=var)
    e1.pack()
    Button(winNew, text='保存', command=com).pack()
    l1 = Label(winNew, text='只能数字。最大值为'+str(cpuCore))
    l1.pack()

def edit_Semaphore():
    winNew = Toplevel(root)
    winNew.geometry('500x200')
    winNew.title('设置核数')

    var = IntVar()
    scl = Scale(winNew, orient=HORIZONTAL, length=400, from_=1, to=200, tickinterval=20, resolution=1,
                variable=var)
    scl.set(sema)
    scl.bind('<ButtonRelease-1>', lambda event:edit(var.get(),lb))
    scl.pack()
    lb = Label(winNew, text='当前设置数值：%d' % sema)
    lb.pack()
    btClose = Button(winNew, text='关闭', command=winNew.destroy)
    btClose.place(relx=0.7, rely=0.5)
def edit(num,lb):
    global sema
    sema = num
    lb.config(text='当前设置数值：%d' % sema)

# fluent end

# abaqus
#路径
AbaqusPath = ""
AbaqusScriptPath = ""


def ChAbaqusPathT():
    global AbaqusPath
    AbaqusPath = tkinter.filedialog.askdirectory(initialdir='C:/SIMULIA/Commands')
    print(AbaqusPath)
    if AbaqusPath != '':
        lbThermal1.config(text='Abaqus路径：'+AbaqusPath)
    else:
        lbThermal1.config(text='Abaqus路径：未选择')

def ChScriptPathT():
    global AbaqusScriptPath
    global ScriptRootPathT
    AbaqusScriptPath = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目',filetypes=[("PY", "py")])
    if AbaqusScriptPath != '':
        print(AbaqusScriptPath)
        wordlist = re.split('/', AbaqusScriptPath)
        ScriptRootPathT = ""
        for i in range(0, len(wordlist) - 1):
            if i!=len(wordlist) - 2:
                ScriptRootPathT += wordlist[i] + '/'
            else:
                ScriptRootPathT += wordlist[i]
        lbThermal2.config(text='脚本文件：' + AbaqusScriptPath)
        lbThermal3.config(text='脚本文件根目录路径：' + ScriptRootPathT)
    else:
        lbThermal2.config(text='脚本文件：未选择')

def ChAbaqusPathS():
    global AbaqusPath
    AbaqusPath = tkinter.filedialog.askdirectory(initialdir='C:/SIMULIA/Commands')
    if AbaqusPath != '':
        lbStatic1.config(text='Abaqus路径：'+AbaqusPath)
    else:
        lbStatic1.config(text='Abaqus路径：未选择')

def ChScriptPathS():
    global ScriptPathS
    global ScriptRootPathS
    ScriptPathS = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目',filetypes=[("PY", "py")])
    if ScriptPathS != '':
        wordlist = re.split('/', ScriptPathS)
        ScriptRootPathS = ""
        for i in range(0, len(wordlist) - 1):
            if i!=len(wordlist) - 2:
                ScriptRootPathS += wordlist[i] + '/'
            else:
                
                ScriptRootPathS += wordlist[i]
        lbStatic2.config(text='脚本文件：' + ScriptPathS)
        lbStatic3.config(text='脚本文件根目录路径：' + ScriptRootPathS)
    else:
        lbStatic2.config(text='脚本文件：未选择')

def ChAbaqusFreqPathT():
    global AbaqusPath
    AbaqusPath = tkinter.filedialog.askdirectory(initialdir='C:/SIMULIA/Commands')
    if AbaqusPath != '':
        lbFreqThermal1.config(text='Abaqus路径：'+AbaqusPath)
    else:
        lbFreqThermal1.config(text='Abaqus路径：未选择')

def ChScriptFreqPathT():
    global AbaqusFreqScriptPath
    global ScriptFreqRootPathT
    AbaqusFreqScriptPath = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目',filetypes=[("PY", "py")])
    if AbaqusFreqScriptPath != '':
        wordlist = re.split('/', AbaqusFreqScriptPath)
        ScriptFreqRootPathT = ""
        for i in range(0, len(wordlist) - 1):
            if i!=len(wordlist) - 2:
                ScriptFreqRootPathT += wordlist[i] + '/'
            else:
                ScriptFreqRootPathT += wordlist[i]
        lbFreqThermal2.config(text='脚本文件：' + AbaqusFreqScriptPath)
        lbFreqThermal3.config(text='脚本文件根目录路径：' + ScriptFreqRootPathT)
    else:
        lbFreqThermal2.config(text='脚本文件：未选择')

def ChAbaqusFreqPathS():
    global AbaqusPath
    AbaqusPath = tkinter.filedialog.askdirectory(initialdir='C:/SIMULIA/Commands')
    if AbaqusPath != '':
        lbFreqStatic1.config(text='Abaqus路径：'+AbaqusPath)
    else:
        lbFreqStatic1.config(text='Abaqus路径：未选择')

def ChScriptFreqPathS():
    global ScriptFreqPathS
    global ScriptFreqRootPathS
    ScriptFreqPathS = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目',filetypes=[("PY", "py")])
    if ScriptFreqPathS != '':
        wordlist = re.split('/', ScriptFreqPathS)
        ScriptFreqRootPathS = ""
        for i in range(0, len(wordlist) - 1):
            if i!=len(wordlist) - 2:
                ScriptFreqRootPathS += wordlist[i] + '/'
            else:
                
                ScriptFreqRootPathS += wordlist[i]
        lbFreqStatic2.config(text='脚本文件：' + ScriptFreqPathS)
        lbFreqStatic3.config(text='脚本文件根目录路径：' + ScriptFreqRootPathS)
    else:
        lbFreqStatic2.config(text='脚本文件：未选择')

def ChAbaqusMorePathS():
    global AbaqusPath
    AbaqusPath = tkinter.filedialog.askdirectory(initialdir='C:/SIMULIA/Commands')
    if AbaqusPath != '':
        lbMoreStatic1.config(text='Abaqus路径：'+AbaqusPath)
    else:
        lbMoreStatic1.config(text='Abaqus路径：未选择')

def ChScriptMorePathS():
    global ScriptMorePathS
    global ScriptMoreRootPathS
    ScriptMorePathS = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目',filetypes=[("PY", "py")])
    if ScriptMorePathS != '':
        wordlist = re.split('/', ScriptMorePathS)
        ScriptMoreRootPathS = ""
        for i in range(0, len(wordlist) - 1):
            if i!=len(wordlist) - 2:
                ScriptMoreRootPathS += wordlist[i] + '/'
            else:
                
                ScriptMoreRootPathS += wordlist[i]
        lbMoreStatic2.config(text='脚本文件：' + ScriptMorePathS)
        lbMoreStatic3.config(text='脚本文件根目录路径：' + ScriptMoreRootPathS)
    else:
        lbMoreStatic2.config(text='脚本文件：未选择')

def ChAbaqusFreqMorePathS():
    global AbaqusPath
    AbaqusPath = tkinter.filedialog.askdirectory(initialdir='C:/SIMULIA/Commands')
    if AbaqusPath != '':
        lbFreqMoreStatic1.config(text='Abaqus路径：'+AbaqusPath)
    else:
        lbFreqMoreStatic1.config(text='Abaqus路径：未选择')

def ChScriptFreqMorePathS():
    global ScriptFreqMorePathS
    global ScriptFreqMoreRootPathS
    ScriptFreqMorePathS = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目',filetypes=[("PY", "py")])
    if ScriptFreqMorePathS != '':
        wordlist = re.split('/', ScriptFreqMorePathS)
        ScriptFreqMoreRootPathS = ""
        for i in range(0, len(wordlist) - 1):
            if i!=len(wordlist) - 2:
                ScriptFreqMoreRootPathS += wordlist[i] + '/'
            else:
                
                ScriptFreqMoreRootPathS += wordlist[i]
        lbFreqMoreStatic2.config(text='脚本文件：' + ScriptFreqMorePathS)
        lbFreqMoreStatic3.config(text='脚本文件根目录路径：' + ScriptFreqMoreRootPathS)
    else:
        lbFreqMoreStatic2.config(text='脚本文件：未选择')


def CmdStart(str):
    txt.config(state=NORMAL)
    command = str
    ret = subprocess.Popen(command, stdin=subprocess.PIPE, stderr=subprocess.PIPE, shell=False, stdout=subprocess.PIPE)
    try:
        for i in iter(ret.stdout.readline, 'b'):
            # line_list = i.split('\r\n')
            if i:
                ret = chardet.detect(i)
                txt.insert(END, i.decode(ret['encoding']))
                txt.see(END)
                print(i.decode(ret['encoding']))
            else:
                break
    finally:
        txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 结束\n")
        txt.see(END)
        print(time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 结束")
        txt.config(state=DISABLED)
        btnAbaqus.config(state=NORMAL)

def StartAbaqusThermal():
    global txt
    txt = txtAbaqus1
    global btnAbaqus
    btnAbaqus = btnT2
    # txt.delete(1.0, END)
    btnAbaqus.config(state=DISABLED)
    txt.config(state=NORMAL)
    txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 开始分析\n")
    file_handle = open('MakeStaticFreqCae.bat', 'w')
    file_handle.write('@ echo off\n')
    file_handle.write('@ set PATH=' + AbaqusPath + '/;%PATH%\n')
    file_handle.write('@ cd '+ ScriptRootPathT+' \n')
    file_handle.write('rem abaqus.bat cae noGUI=test.py\n')
    file_handle.write('rem abaqus.bat cae script=' + AbaqusScriptPath + '\n')
    file_handle.write('abaqus.bat cae noGUI=' + AbaqusScriptPath + '\n\n')
    file_handle.write('pause')
    file_handle.close()
    try:
        t1 = threading.Thread(target=CmdStart, args=("MakeStaticFreqCae.bat",))
        t1.setDaemon(True)
        t1.start()
    except:
        txt.insert(END, '开始失败\n')
        txt.config(state=DISABLED)
        btnAbaqus.config(state=NORMAL)

def StartAbaqusStatic():
    # txt.delete(1.0, END)
    global txt
    txt = txtAbaqus2
    global btnAbaqus
    btnAbaqus = btnStatic3
    btnAbaqus.config(state=DISABLED)
    txt.config(state=NORMAL)
    txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 开始分析\n")
    file_handle = open('MakeStaticCae.bat', 'w')
    file_handle.write('@ echo off\n')
    file_handle.write('@ set PATH=' + AbaqusPath + '/;%PATH%\n')
    file_handle.write('@ cd '+ ScriptRootPathS +' \n')
    file_handle.write('abaqus.bat cae noGUI=' + ScriptPathS + '\n\n')
    file_handle.write('pause')
    file_handle.close()
    try:
        t1 = threading.Thread(target=CmdStart,args=("MakeStaticCae.bat",))
        t1.setDaemon(True)
        t1.start()
    except:
        txt.insert(END, '开始失败\n')
        txt.config(state=DISABLED)
        btnAbaqus.config(state=NORMAL)

def StartFreqThermal():
    global txt
    txt = txtAbaqusFreq1
    global btnAbaqus
    btnAbaqus = btnFT2
    # txt.delete(1.0, END)
    btnAbaqus.config(state=DISABLED)
    txt.config(state=NORMAL)
    txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 开始分析\n")
    file_handle = open('MakeFreqThermalCae.bat', 'w')
    file_handle.write('@ echo off\n')
    file_handle.write('@ set PATH=' + AbaqusPath + '/;%PATH%\n')
    file_handle.write('@ cd '+ScriptFreqRootPathT+' \n')
    file_handle.write('rem abaqus.bat cae noGUI=test.py\n')
    file_handle.write('rem abaqus.bat cae script=' + AbaqusFreqScriptPath + '\n')
    file_handle.write('abaqus.bat cae noGUI=' + AbaqusFreqScriptPath + '\n\n')
    file_handle.write('pause')
    file_handle.close()
    try:
        t1 = threading.Thread(target=CmdStart, args=("MakeFreqThermalCae.bat",))
        t1.setDaemon(True)
        t1.start()
    except:
        txt.insert(END, '开始失败\n')
        txt.config(state=DISABLED)
        btnAbaqus.config(state=NORMAL)

def StartFreqStatic():
    # txt.delete(1.0, END)
    global txt
    txt = txtAbaqusFreq2
    global btnAbaqus
    btnAbaqus = btnFreqStatic3
    btnAbaqus.config(state=DISABLED)
    txt.config(state=NORMAL)
    txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 开始分析\n")
    file_handle = open('MakeFreqStaticCae.bat', 'w')
    file_handle.write('@ echo off\n')
    file_handle.write('@ set PATH=' + AbaqusPath + '/;%PATH%\n')
    file_handle.write('@ cd '+ScriptFreqRootPathS +' \n')
    file_handle.write('abaqus.bat cae noGUI=' + ScriptFreqPathS + '\n\n')
    file_handle.write('pause')
    file_handle.close()
    try:
        t1 = threading.Thread(target=CmdStart,args=("MakeFreqStaticCae.bat",))
        t1.setDaemon(True)
        t1.start()
    except:
        txt.insert(END, '开始失败\n')
        txt.config(state=DISABLED)
        btnAbaqus.config(state=NORMAL)

def StartMoreStatic():
        # txt.delete(1.0, END)
    global txt
    txt = txtAbaqusMore
    global btnAbaqus
    btnAbaqus = btnMoreStatic3
    btnAbaqus.config(state=DISABLED)
    txt.config(state=NORMAL)
    txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 开始分析\n")
    file_handle = open('MakeMoreCae.bat', 'w')
    file_handle.write('@ echo off\n')
    file_handle.write('@ set PATH=' + AbaqusPath + '/;%PATH%\n')
    file_handle.write('@ cd '+ScriptMoreRootPathS +' \n')
    file_handle.write('abaqus.bat cae noGUI=' + ScriptMorePathS + '\n\n')
    file_handle.write('pause')
    file_handle.close()
    try:
        t1 = threading.Thread(target=CmdStart,args=("MakeMoreCae.bat",))
        t1.setDaemon(True)
        t1.start()
    except:
        txt.insert(END, '开始失败\n')
        txt.config(state=DISABLED)
        btnAbaqus.config(state=NORMAL)
    
def StartMoreFreq():
    # txt.delete(1.0, END)
    global txt
    txt = txtAbaqusMore2
    global btnAbaqus
    btnAbaqus = btnFreqMoreStatic3
    btnAbaqus.config(state=DISABLED)
    txt.config(state=NORMAL)
    txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 开始分析\n")
    file_handle = open('MakeFreqMoreCae.bat', 'w')
    file_handle.write('@ echo off\n')
    file_handle.write('@ set PATH=' + AbaqusPath + '/;%PATH%\n')
    file_handle.write('@ cd '+ ScriptFreqMoreRootPathS +' \n')
    file_handle.write('abaqus.bat cae noGUI=' + ScriptFreqMorePathS + '\n\n')
    file_handle.write('pause')
    file_handle.close()
    try:
        t1 = threading.Thread(target=CmdStart,args=("MakeFreqMoreCae.bat",))
        t1.setDaemon(True)
        t1.start()
    except:
        txt.insert(END, '开始失败\n')
        txt.config(state=DISABLED)
        btnAbaqus.config(state=NORMAL)


# abaqus end


#Alldef

#分析对象
Object1="Turbine"#涡轮
Object2="涡轮"#涡轮

def ClearFileDir():
    txt.insert(END, '\t 清理不必要的文件 \n')
    txt.see(END)
    print('\t 清理不必要的文件 \n')
    filedir = []
    # 以SMT-开头的文件夹：构建IDW, RBF, RMTB, RMTC模型，临时保存目录data_dir
    # 以.pkl结尾的文件：除了(IDW, RBF, RMTB, RMTC)之外的模型，临时保存在pkl文件
    # *.png和KrigingOut*.txt是输出
    filedir = glob.glob("SMT-*") + \
              glob.glob("*.pkl") + \
              glob.glob("*.png") + \
              glob.glob("KrigingOut*.txt")
    # print(filedir)
    os.chdir(RootPath) # 保存文件的根路径
    for path in filedir:
        if os.path.isdir(path):  # 目录用shutil.rmtree删除
            shutil.rmtree(path)
        elif os.path.isfile(path):  # 文件用os.remove删除
            os.remove(path)
    # EndFor


# EndDef

def FigConfig():
    config = {
      "font.family": 'serif',
      "font.size": 9,
      "legend.fontsize":7,#图例字体尺寸
      "axes.titlesize":7,#坐标轴字体尺寸
      "font.weight": 'normal',
      "font.serif": ['SimSun'],#中文宋体
      "savefig.dpi":500,#保存图片的分辨率
      "figure.figsize":(6.0, 3.0),#保存图片的尺寸
      "axes.unicode_minus": 'False'#正常显示坐标轴正负号（minus），unicode显示负号异常
      }
    plt.rcParams.update(config)


# EndDef

# 根据Abaqus输入参数的随机特征，进行大样本抽样，样本数目为Num1。
# 目前针对密度、弹性模量和转速进行抽样
def ReliabSampling(Num1, FileGauss, _Text):
    FigList = []
    FigCaption = []

    # 用于可靠性分析的抽样
    # 读取文本文件，存入矩阵。分隔符为' '，注释符为'#'
    # 第1列为平均值，第2列为变异系数
    data = np.loadtxt(FileGauss, comments='#', encoding='utf-8-sig')
    length = data.shape[0]
    data = data[0:length - 1, :]

    # 样本平均值
    Mean1 = data[:, 0]
    # 样本方差
    Vars1 = (data[:, 0] * data[:, 1]) ** 2  # (平均值*变异系数)^2
    # print(data[:,0]*data[:,1])

    # 构建单位对角线矩阵
    Cov1 = np.eye(Mean1.shape[0])
    # 提取对角线元素的坐标
    row, col = np.diag_indices_from(Cov1)
    # 修改对角线元素为方差
    Cov1[row, col] = Vars1

    # 采用多维正态分布抽样方法，构建样本
    data = np.random.multivariate_normal(Mean1, Cov1, Num1)

    # 画图
    # 检测密度、弹性模量、转速的抽样结果，保存为png图片
    data1 = data[:, 0]
    data2 = data[:, 1]
    data3 = data[:, 2]
    StrVar = ['密度', '弹性模量', '转速']
    StrUnit = [' [kg/m^3]', ' [GPa]', ' [r/min]']
    StrFilename = ['rho', 'E', 'rpm']
    for index in range(len(data[0, :])):
        DistX = data[:, index]

        fig = plt.figure()  # 定义一个图像窗口
        # 直方图+概率密度+拟合曲线
        dist = sns.distplot(DistX, hist=True, kde=False, fit=st.norm,
                            fit_kws={'color': 'red', 'label': '正态分布拟合', 'linestyle': '-'})
        plt.xlabel(StrVar[index] + StrUnit[index])
        plt.ylabel('')
        Str1 = StrVar[index] + '抽样'
        Str2 = '(拟合平均值%.2f' % (st.norm.fit(DistX)[0])
        Str3 = '标准差%.2f' % (st.norm.fit(DistX)[1]) + ')'
        StrTitle = _Text + Str1 + Str2 + ',' + Str3
        plt.title(StrTitle)
        FigName = "Fig-Sample-Dist-" + StrFilename[index] + ".png"
        ##显示图例
        plt.legend()
        # plt.show()
        plt.savefig(FigName, bbox_inches='tight')
        plt.close()
        FigList.append(FigName)
        FigCaption.append(StrTitle)
    # EndFor
    return data, FigList, FigCaption


##end def

# 从文件中读取GH4169材料的4个疲劳系数
# (用于局部应力应变法求疲劳寿命)
def FuncReadFatigueP(FileMater):
    global sigma_f, b, epsilon_f, c
    data = np.loadtxt(FileMater, comments='#', encoding='utf-8-sig')
    # 疲劳强度系数[MPa]745.4
    sigma_f = data[0]
    # 疲劳强度指数-0.092
    b = data[1]
    # 疲劳塑性系数0.161
    epsilon_f = data[2]
    # 疲劳塑性指数-0.419
    c = data[3]


# def

# 定义局部应力应变法求疲劳寿命的函数
# SWT疲劳寿命函数Func1(x)=0
def Func1(x):
    return sigma_f ** 2 / E * x ** (2 * b) + epsilon_f * sigma_f * x ** (b + c) - epsilon_a * sigma_max


##end def

# python并行计算用的循环体：循环内命令彼此独立，不相干。
# 读取data中的数据，作为疲劳寿命计算的参数，
# 调用fsolve求解疲劳寿命函数，求得疲劳计算中的工作次数
def do_something(data):
    global E, epsilon_a, sigma_max, sigma_f, b, epsilon_f, c
    E = data[0]  # 弹模模量
    epsilon_a = data[1]  # 应变幅
    sigma_max = data[2]  # 最大应力
    sigma_f = data[3]  # 疲劳强度系数[MPa]
    b = data[4]  # 疲劳强度指数
    epsilon_f = data[5]  # 疲劳塑性系数
    c = data[6]  # 疲劳塑性指数
    # print(E,epsilon_a,sigma_max,sigma_f,b,epsilon_f,c)
    # Func1函数中自变量x=2*NFatigue
    # NFatigue为疲劳计算中的工作次数
    NFatigue = fsolve(Func1, 1, xtol=1e-8) / 2
    return NFatigue[0]


# end def

def GetDataFreq(StrFile):
    data = np.loadtxt(StrFile, dtype=float, skiprows=1, encoding='utf-8-sig')
    X0 = data[:, len(data[0]) - 3:len(data[0])]  # 密度、弹模、转速
    # 前6阶节径振型频率存储的位置
    newArr = en.get().split(",")
    locs = np.array([int(x) for x in newArr])
    locs = locs - 1  # python 数组第1个下标是0，
    # print(locs)

    # 从频率文件中提取前6阶节径振型固有频率
    for index in range(len(locs)):
        ii = locs[index]
        if index == 0:
            Freq0 = np.c_[data[:, ii]]
        else:
            Freq0 = np.c_[Freq0, data[:, ii]]  # 拼接矩阵
    # EndFor
    txt.insert(END, '\t 该文件共有%d组数据 \n' % (len(X0[:, 0])))
    txt.see(END)
    print('\t 该文件共有%d组数据 \n' % (len(X0[:, 0])))

    return X0, Freq0


# EndDef

def Limits(X):
    for index in range(X.shape[1]):
        xmin = min(X[:, index])
        xmax = max(X[:, index])
        temp = np.array([[xmin, xmax]])
        if index == 0:
            xlimits = np.r_[temp]
        else:
            xlimits = np.r_[xlimits, temp]
        # EndIf
    # EndFor
    return xlimits


# EndDef

def mkdir(path):
    # 去除首位空格
    path = path.strip()
    # 去除尾部 \ 符号
    path = path.rstrip("\\")
    # 判断路径是否存在
    isExists = os.path.exists(path)
    # 判断结果
    if not isExists:
        os.makedirs(path)
        txt.insert(END, '\t\t\t %s创建成功\n' % (path))
        txt.see(END)
        print('\t\t\t %s创建成功\n' % (path))
        return True
    else:
        # 如果目录存在则不创建，并提示目录已存在
        shutil.rmtree(path)
        txt.insert(END, '\t\t\t %s虽然存在，但已经删除，并重新创建\n' % (path))
        txt.see(END)
        print('\t\t\t %s虽然存在，但已经删除，并重新创建\n' % (path))
        os.makedirs(path)
        return False
    # EndIf


# EndDef

# 搜索最优的SMT方法
def OptSMT(X, Y, X1, Y1, Strs, ratio):
    FigList = []
    FigCaption = []
    TxtList = []
    TxtCaption = []
    TxtHead = []

    xlimits = Limits(X)  # 确定输入量X的极值，n行2列的数据
    StrVar = Strs[0]  # SMax,EMax,Mode1
    if StrVar[0] == 'S':
        StrVarOut = '最大Miese应力'
        StrType = 'Static'
        StrUnit = 'MPa'
    elif StrVar[0] == 'E':
        StrVarOut = '最大Miese应变'
        StrType = 'Static'
        StrUnit = ''
    elif StrVar[0] == 'M':
        StrVarOut = '第' + StrVar[-1] + '节径振型固有频率'
        StrType = 'Frequency'
        StrUnit = 'Hz'
    # EndIf
    # 一些设置参数
    IDW_p = 5.0
    RMTB_order = 3
    RMTB_num = 15
    RBF_d0 = 5
    RBF_poly = 1
    # 用来构建代理模型的几个方法名称
    SMTModels = ['RBF', 'IDW', 'RMTB', 'RMTC', 'LS', 'QP', 'KRG', 'KPLS', 'KPLSK']
    # SMTModels=['RMTC']
    # 存储相对均方根误差，求其最小值。
    RMSEMin = []
    # 存储SMT的plk文件名
    StrFileDir = []
    for index in range(len(SMTModels)):
        t0 = timeit.default_timer()
        StrSMT = SMTModels[index]
        txt.insert(END, '\t\t 测试SMT模型：%s \n' % (StrSMT))
        txt.see(END)
        print('\t\t 测试SMT模型：%s \n' % (StrSMT))

        # IDW, RBF, RMTB, RMTC模型用data_dir=FileDir可以临时保存，
        # 其它SMT模型需要保存到pkl文件，
        # 后续提取最佳SMT模型时直接调用data_dir，加载pkl文件，不用再训练
        if (StrSMT[0] != 'I' and StrSMT[0] != 'R'):
            FileDir = 'SMT-' + StrType + '-' + StrVar + '-' + StrSMT + '.pkl'
        else:
            FileDir = 'SMT-' + StrType + '-' + StrVar + '-' + StrSMT
        # EndIf

        # 调用SMT中的方法，getattr函数能够根据字符串调用函数
        str1 = getattr(smt.surrogate_models, StrSMT)
        if StrSMT == 'RBF':  # RBF, data_dir用于保存模型
            mkdir(FileDir)
            StrFileDir.append(FileDir)
            sm = str1(print_global=False, data_dir=FileDir, d0=RBF_d0, poly_degree=RBF_poly)
            txt.insert(END, '\t\t\t %s保存到%s文件夹 \n' % (StrSMT, FileDir))
            txt.see(END)
            print('\t\t\t %s保存到%s文件夹 \n' % (StrSMT, FileDir))
        elif StrSMT == 'RMTB':  # RMTB, data_dir用于保存模型
            mkdir(FileDir)
            StrFileDir.append(FileDir)
            sm = str1(print_global=False, data_dir=FileDir, xlimits=xlimits, order=RMTB_order, num_ctrl_pts=RMTB_num)
            txt.insert(END, '\t\t\t %s保存到%s文件夹 \n' % (StrSMT, FileDir))
            txt.see(END)
            print('\t\t\t %s保存到%s文件夹 \n' % (StrSMT, FileDir))
        elif StrSMT == 'RMTC':  # RMTC, data_dir用于保存模型
            mkdir(FileDir)
            StrFileDir.append(FileDir)
            sm = str1(print_global=False, data_dir=FileDir, xlimits=xlimits)
            txt.insert(END, '\t\t\t %s保存到%s文件夹 \n' % (StrSMT, FileDir))
            txt.see(END)
            print('\t\t\t %s保存到%s文件夹 \n' % (StrSMT, FileDir))
        elif StrSMT == 'IDW':  # IDW, data_dir用于保存模型
            mkdir(FileDir)
            StrFileDir.append(FileDir)
            sm = str1(print_global=False, data_dir=FileDir, p=IDW_p)
            txt.insert(END, '\t\t\t %s保存到%s文件夹 \n' % (StrSMT, FileDir))
            txt.see(END)
            print('\t\t\t %s保存到%s文件夹 \n' % (StrSMT, FileDir))
        elif StrSMT[0] == 'K':  # Kriging 族
            sm = str1(print_global=False, poly='quadratic', corr='squar_exp')
            StrFileDir.append(FileDir)
        else:  # LS, QP
            sm = str1(print_global=False)
            StrFileDir.append(FileDir)
        # EndIf

        # 设置smt输入数据
        sm.set_training_values(X, Y)
        # 训练smt
        sm.train()

        # 保存除(IDW, RBF, RMTB, RMTC)之外的SMT模型到pkl文件，后续可以加载
        if (StrSMT[0] != 'I' and StrSMT[0] != 'R'):
            with open(FileDir, "wb") as f:
                pickle.dump(sm, f)
            txt.insert(END, '\t\t\t %s保存到%s文件夹 \n' % (StrSMT, FileDir))
            print('\t\t\t %s保存到%s文件 \n' % (StrSMT, FileDir))
            # EndIf
        txt.insert(END, '\t\t\t 校核SMT模型 \n')
        print('\t\t\t 校核SMT模型 \n')
        # 根据检测数据的输入X1，预测
        Y1p = sm.predict_values(X1)
        ymin = min(Y1) * ratio
        ymax = max(Y1) * ratio

        # 判断Kriging预测值与Abaqus计算值之间的差别：相对均方根误差:均方根误差/真实值的平均值
        RMSE = np.sqrt(np.sum((1 - Y1p / Y1) ** 2) / len(Y1))
        RMSEMin.append(RMSE)

        StrRMSE=Object2+StrVarOut+'SMT模型校核的相对RMSE='+'%6.4f' %(RMSE*1e3) +'‰('+StrSMT+'方法)'
        StrTitle = StrRMSE
        txt.insert(END, '\t\t\t 对校核结果画图 \n')
        txt.see(END)
        print('\t\t\t 对校核结果画图 \n')
        # ratio=[1,1e3]
        StrYlabels = Strs[1]  # '最大Mises应力 [MPa]',  '最大Mises应变 [‰]', '第1节径振型固有频率[Hz]'
        FigName = 'Fig-' + StrType + '-RMSE-' + StrVar + '-' + StrSMT + '.png'
        # ['Fig-Static-RMSE-SMax.png','Fig-Static-RMSE-EMax.png']
        fig = plt.figure()  # 定义一个图像窗口
        plt.plot(ratio * Y1p, 'o', markersize=4, color='red', markerfacecolor='none')
        plt.plot(ratio * Y1, '.', markersize=2, color='blue')
        plt.xlabel('抽样点编号')
        plt.ylabel(StrYlabels)
        plt.legend(["预测值", "真实值"], loc='best')
        plt.title(StrTitle)
        plt.ylim(ymin - (ymax - ymin) * 0.05, ymax + (ymax - ymin) * 0.05)
        plt.savefig(FigName, bbox_inches='tight')
        plt.close()
        FigList.append(FigName)
        FigCaption.append(StrTitle)

        # 保存检测结果到txt文件
        txt.insert(END, '\t\t\t 保存校核结果到txt文件 \n')
        print('\t\t\t 保存校核结果到txt文件 \n')
        XX=np.arange(1,len(Y1)+1,1)
        data=np.c_[XX,Y1p,Y1]
        TxtName = 'KrigingOut-' + StrType + '-' + StrVar + '-' + StrSMT + '.txt'
        TabCaption = StrTitle
        TabHead=['样本编号']+ ['%s方法预测值[%s]' %(StrSMT,StrUnit),'Abaqus值[%s]' %(StrUnit)]
        f = open(TxtName, 'w', encoding='utf-8-sig')  # 打开文件， 用'w'写文件
        f.write('#%s\n' % (TabCaption))
        f.write('#    %s    ' % (TabHead[0]))
        f.write('  %s       %s\n' %(TabHead[1],TabHead[2]))

        np.savetxt(f, X=data, delimiter=' ', encoding='utf-8-sig')
        f.close()
        TxtList.append(TxtName)
        TxtCaption.append(TabCaption)
        TxtHead.append(TabHead)
        tn = timeit.default_timer()
        txt.insert(END, '\t\t\t 完成！耗时:'+str(tn - t0)+'s.\n')
        txt.see(END)
        print('\t\t\t 完成！耗时:', str(tn - t0), 's.\n')
    # EndFor

    # 查找误差最小的SMT方法
    index = np.argmin(RMSEMin)
    StrSMT = SMTModels[index]
    str1 = getattr(smt.surrogate_models, StrSMT)
    str2='%s%s的最佳SMT模型为:%s方法' %(Object2,StrVarOut,StrSMT)
    OptSMT_str=[str2]
    print('\t %s \n' %(str2)) 
    txt.insert(END, '\t %s \n' %(str2))
    FileDir = StrFileDir[index]
    t0 = timeit.default_timer()

    # 加载之前保存的SMT模型
    if (StrSMT[0] != 'I' and StrSMT[0] != 'R'):  # 加载除(IDW, RBF, RMTB, RMTC)之外的SMT模型
        with open(FileDir, "rb") as f:
            sm = pickle.load(f)
    else:
        if StrSMT == 'RBF':  # 加载(IDW, RBF, RMTB, RMTC)中的SMT模型
            sm = str1(print_global=False, data_dir=FileDir, d0=RBF_d0, poly_degree=RBF_poly)
        elif StrSMT == 'RMTB':  # RMTB, data_dir用于保存pkl文件
            sm = str1(print_global=False, data_dir=FileDir, xlimits=xlimits, order=RMTB_order, num_ctrl_pts=RMTB_num)
        elif StrSMT == 'RMTC':  # RMTC, data_dir用于保存pkl文件
            sm = str1(print_global=False, data_dir=FileDir, xlimits=xlimits)
        elif StrSMT == 'IDW':  # 加载(IDW, RBF, RMTB, RMTC)中的SMT模型
            sm = str1(print_global=False, data_dir=FileDir, p=IDW_p)
        sm.set_training_values(X, Y)  # 输入参数与之前相同
        sm.train()  # 实际并未训练，直接加载
    tn = timeit.default_timer()
    txt.insert(END, '\t 完成！耗时:'+str(tn - t0)+'s.\n')
    txt.see(END)
    print('\t 完成！耗时:', str(tn - t0), 's.\n')
    ####用于再次校核
    ##Y1p=sm.predict_values(X1)
    ##RMSE=np.sqrt(np.sum((1-Y1p/Y1)**2)/len(Y1))
    ##print(RMSE)
    return sm,OptSMT_str,FigList,FigCaption,TxtList,TxtCaption,TxtHead


# EndDef

def StressStrain(Exec, Num1, FileGauss, FileAbaqus1, FileAbaqus2):
    ######################
    # 构建SMT模型
    ######################

    # 读取Abaqus计算结果1，用于构建kriging模型
    txt.insert(END, '\t 读取Abaqus计算得到的最大Mises应力和Mises应变 \n')
    txt.insert(END, '\t\t 读取用于构建SMT的数据 \n')
    print('\t 读取Abaqus计算得到的最大Mises应力和Mises应变 \n')
    print('\t\t 读取用于构建SMT的数据 \n')
    data = np.loadtxt(FileAbaqus1, dtype=float, skiprows=1, encoding='utf-8-sig')
    X0 = data[:, 2:5]  # 密度、弹模、转速
    Y0 = data[:, 0:2]  # 最大Mises应力和应变
    txt.insert(END, '\t\t 共有%d个样本点 \n' % (len(X0[:, 0])))
    print('\t\t 共有%d个样本点 \n' % (len(X0[:, 0])))

    print('\t\t 读取用于校核SMT的样本数据 \n')
    txt.insert(END, '\t\t 读取用于校核SMT的样本数据 \n')
    data =np.loadtxt(FileAbaqus2, dtype=float, skiprows=1,encoding='utf-8-sig')
    X1   =data[:,2:5]#密度、弹模、转速
    Y1   =data[:,0:2]#最大Mises应力和应变
    print('\t\t 共有%d个样本点 \n' %(len(X1[:,0])))
    txt.insert(END, '\t\t 共有%d个样本点 \n' % (len(X1[:, 0])))

    ##处理校核SMT模型的样本，用于生成docx报告
    TxtName='KrigingOut-Samples.txt'
    TabCaption=Object2+'校核SMT模型的样本'
    TabHead=['样本编号','密度 [kg/m^3]','弹性模量 [GPa]','转速 [r/min] ']+\
        ['最大Miese应力 [MPa]','最大Miese应变']
    f = open(TxtName, 'w', encoding='utf-8-sig') # 打开文件， 用'w'写文件
    f.write('#%s\n'   %(TabCaption))
    f.write('#    %s    ' %(TabHead[0]))
    f.write('  %s       %s       %s' %(TabHead[1],TabHead[2],TabHead[3]))
    f.write('  %s       %s\n' %(TabHead[4],TabHead[5]))
    XX=np.arange(1,len(Y1[:,0])+1,1)
    data=np.c_[XX,X1,Y1]
    np.savetxt(f, X=data, delimiter=' ', encoding='utf-8-sig')
    f.close()
    TxtList0=[TxtName]
    TxtCaption0=[TabCaption]
    TxtHead0=[TabHead]

    txt.insert(END, '\t\t 读取用于校核SMT的样本数据 \n')
    print('\t\t 读取用于校核SMT的样本数据 \n')
    data = np.loadtxt(FileAbaqus2, dtype=float, skiprows=1, encoding='utf-8-sig')
    X1 = data[:, 2:5]  # 密度、弹模、转速
    Y1 = data[:, 0:2]  # 最大Mises应力和应变
    txt.insert(END, '\t\t 共有%d个样本点 \n' % (len(X1[:, 0])))
    print('\t\t 共有%d个样本点 \n' % (len(X1[:, 0])))

    # 构建最大Mises应力的SMT模型，从几个SMT中搜索最优模型
    txt.insert(END, '\n\t 搜索最大Mises应力的最佳SMT模型\n')
    print('\n\t 搜索最大Mises应力的最佳SMT模型\n')
    Strs = ['SMax', '最大Mises应力 [MPa]']
    ratio = 1
    SMax = np.c_[Y0[:, 0]]
    SMax1 = np.c_[Y1[:, 0]]
    sm1,OptSMT_str1,FigList1,FigCaption1,TxtList1,TxtCaption1,TxtHead1=\
        OptSMT(X0, SMax, X1, SMax1, Strs, ratio)

    # 构建最大Mises应变的SMT模型，从几个SMT中搜索最优模型
    txt.insert(END, '\n\t 搜索最大Mises应变的最佳SMT模型\n')
    txt.see(END)
    print('\n\t 搜索最大Mises应变的最佳SMT模型\n')
    Strs = ['EMax', '最大Mises应变 [‰]']
    ratio = 1e3
    EMax = np.c_[Y0[:, 1]]
    EMax1 = np.c_[Y1[:, 1]]
    sm2, OptSMT_str2,FigList2, FigCaption2, TxtList2, TxtCaption2, TxtHead2 = \
        OptSMT(X0, EMax, X1, EMax1, Strs, ratio)

    ##疲劳可靠性计算抽样
    txt.insert(END, '\t 抽样数目%3.2e组 \n' % Num1)
    print('\t 抽样数目%3.2e组 \n' % Num1)
    _Text = Object2+Exec
    XR, FigList3, FigCaption3 = ReliabSampling(Num1, FileGauss, _Text)

    # Kriging预测抽样样本的最大Mises应力
    txt.insert(END, '\t Kriging预测抽样样本的最大Mises应力 \n')
    print('\t Kriging预测抽样样本的最大Mises应力 \n')
    # YR=np.zeros([XR.shape[0],2])
    temp1 = sm1.predict_values(XR)

    # Kriging预测抽样样本的最大Mises应变
    txt.insert(END, '\t Kriging预测抽样样本的最大Mises应变 \n')
    txt.see(END)
    print('\t Kriging预测抽样样本的最大Mises应变 \n')
    temp2 = sm2.predict_values(XR)
    YR = np.c_[temp1, temp2]
    OptOut_str=OptSMT_str1+OptSMT_str2

    FigList = FigList1 + FigList2 + FigList3
    FigCaption = FigCaption1 + FigCaption2 + FigCaption3
    TxtList   =TxtList0   +TxtList1   +TxtList2
    TxtCaption=TxtCaption0+TxtCaption1+TxtCaption2
    TxtHead   =TxtHead0   +TxtHead1   +TxtHead2
    return XR,YR,OptOut_str,FigList,FigCaption,TxtList,TxtCaption,TxtHead


# EndDef

def StressReliability(YR, FileGauss):
    FigList = []
    FigCaption = []
    # 读取GH4169的抗拉强度分布特性
    data = np.loadtxt(FileGauss, skiprows=8, comments='#', encoding='utf-8-sig')
    mu1 = data[0]  # data[0]是平均值
    sigma1 = data[0] * data[1]  # data[1]是标准差

    # 确定抗拉强度分布
    x1_min = mu1 - 5.0 * sigma1
    x1_max = mu1 + 5.0 * sigma1
    x1 = np.linspace(x1_min, x1_max, 1000)
    y1 = st.norm.pdf(x1, mu1, sigma1)

    txt.insert(END, '\t 绘制最大Mises应力概率分布图 \n')
    print('\t 绘制最大Mises应力概率分布图 \n')
    fig = plt.figure()
    # 最大Mises应力概率分布
    dist = sns.distplot(YR, hist=True, kde=False, fit=st.norm, hist_kws={'facecolor': 'gray', 'alpha': 0.3},
                        fit_kws={'lw': 0.5, 'color': 'purple', 'label': '正态分布拟合', 'linestyle': '-'})

    plt.plot(x1, y1, color='red', linewidth=0.5)
    plt.fill_between(x1, y1, facecolor='gray', alpha=0.3)

    StrTitle = Object2+ "最大Mises应力分布"
    plt.xlabel('应力 [MPa]')
    plt.ylabel('概率密度')
    plt.title(StrTitle)
    plt.legend(['最大Mises应力', '抗拉强度'], loc='best')
    FigName = 'Fig-Static-Dist.png'
    plt.savefig(FigName, bbox_inches='tight')
    plt.close()
    FigList.append(FigName)
    FigCaption.append(StrTitle)

    # 用norm分布拟合最大Mises应力数据，返回norm分布参数
    (mu2, sigma2) = st.norm.fit(YR)

    beta = abs(mu1 - mu2) / np.sqrt(sigma1 ** 2 + sigma2 ** 2)
    R_S = st.norm.cdf(beta)  # 静强度可靠度

    txt.insert(END, '\n保存强度可靠度到txt文本\n')
    txt.see(END)
    print('\n保存强度可靠度到txt文本\n')
    TxtName = 'KrigingOut-Static-Reliability.txt'
    f = open(TxtName, 'w', encoding='utf-8-sig')  # 打开文件， 用'w'写文件
    f.write('#静强度可靠度\n')
    f.write('%f\n' % (R_S))
    f.close()

    return R_S, FigList, FigCaption


# EndDef

def FrequencyReliability(IFreq, X, Y, X1, Y1, delta, Num1, FileGauss):
    # 第1列为平均值，第2列为变异系数
    data = np.loadtxt(FileGauss, comments='#', encoding='utf-8-sig')
    rpm_avg = data[2, 0]
    # 转速变异系数
    rpm_cv = data[2, 1]
    # 转速标准差
    rpm_std = rpm_avg * rpm_cv

    txt.insert(END, '处理第%d节径振型固有频率 \n' % (IFreq))
    print('处理第%d节径振型固有频率 \n' % (IFreq))

    Strs = ['Mode%d' % (IFreq), '第%d节径振型固有频率[Hz]' % (IFreq)]
    ratio = 1
    sm,OptSMT_str,FigList1,FigCaption1,TxtList,TxtCaption,TxtHead=OptSMT(X,Y,X1,Y1,Strs,ratio)
    ##振动可靠性计算抽样
    txt.insert(END, '\t 振动可靠性计算抽样 \n')
    print('\t 振动可靠性计算抽样 \n')
    _Text = Object2+'第%d节径振动可靠性计算' % (IFreq)
    XR, FigList2, FigCaption2 = ReliabSampling(Num1, FileGauss, _Text)

    FigList = FigList1 + FigList2
    FigCaption = FigCaption1 + FigCaption2

    # Kriging预测抽样样本的固有频率
    txt.insert(END, '\t Kriging预测抽样样本的固有频率 \n')
    print('\t Kriging预测抽样样本的固有频率 \n')
    YR = sm.predict_values(XR)  # 样本固有频率 ω_n
    txt.insert(END, '\t 三重点共振可靠度 \n')
    txt.see(END)
    print('\t 三重点共振可靠度 \n')
    Freq_1 = np.c_[XR[:, 2]] / 60.0  # 1倍转动频率 ω
    Error = YR / Freq_1 - IFreq  # n节径共振时的三重点共振条件:ω_n/ω=k=n。IFreq=n。若error=0，表示共振
    Num2 = np.sum(Error > 0.01)  # 统计非零元素数目
    R_tri = Num2 / Num1  # 非零元素比例：未共振的可靠度。
    # print(R_tri)

    # 用norm分布拟合数据，返回norm分布参数
    (mu, sigma) = st.norm.fit(YR)

    # 共振避开率可靠度，采用应力干涉法
    # delta=0.05#避开率Δ
    # 高于(1+Δ)倍k次谐波的正态分布
    mu2 = (1 + delta) * IFreq * rpm_avg / 60.0
    sigma2 = (1 + delta) * IFreq * rpm_std / 60.0
    # 低于(1-Δ)倍k次谐波的正态分布
    mu1 = (1 - delta) * IFreq * rpm_avg / 60.0
    sigma1 = (1 - delta) * IFreq * rpm_std / 60.0

    beta1 = abs(mu - mu1) / np.sqrt(sigma ** 2 + sigma1 ** 2)
    beta2 = abs(mu - mu2) / np.sqrt(sigma ** 2 + sigma2 ** 2)

    txt.insert(END, '\t 第%d节径振型固有频率概率分布图 \n' % (IFreq))
    txt.see(END)
    print('\t 第%d节径振型固有频率概率分布图 \n' % (IFreq))
    fig = plt.figure()
    # 根据实际数据画疲劳寿命概率分布
    dist = sns.distplot(YR, hist=True, kde=False, fit=st.norm, hist_kws={'facecolor': 'gray', 'alpha': 0.3},
                        fit_kws={'lw': 0.5, 'color': 'purple', 'label': '正态分布拟合', 'linestyle': '-'})

    # 根据norm拟合各节径振型固有频率概率分布
    x0 = np.linspace(min(YR), max(YR), 1000)
    loc = mu  # 平均值
    scale = sigma  # 标准差
    y0 = st.norm.pdf(x0, loc, scale)

    # 根据转速的正态分布，确定高于(1+Δ)倍k次谐波的分布
    x1_min = mu1 - 5.0 * sigma1
    x1_max = mu1 + 5.0 * sigma1
    x1 = np.linspace(x1_min, x1_max, 1000)
    y1 = st.norm.pdf(x1, mu1, sigma1)

    x2_min = mu2 - 5.0 * sigma2
    x2_max = mu2 + 5.0 * sigma2
    x2 = np.linspace(x2_min, x2_max, 1000)
    y2 = st.norm.pdf(x2, mu2, sigma2)
    plt.plot(x1, y1, color='red', linewidth=0.5)
    plt.plot(x2, y2, color='blue', linewidth=0.5)
    plt.fill_between(x1, y1, facecolor='gray', alpha=0.3)
    plt.fill_between(x2, y2, facecolor='gray', alpha=0.3)
    StrTitle = '%s第%d节径振型固有频率概率分布' % (Object2,IFreq)
    plt.xlabel('频率 [Hz]')
    plt.ylabel('概率密度')
    plt.title(StrTitle)
    plt.legend(['%d节径振型固有频率' % (IFreq), '%d次谐波频率×%3.2f' % (IFreq, 1 - delta),
                '%d次谐波频率×%3.2f' % (IFreq, 1 + delta)], loc='best')
    FigName = 'Fig-Frequency-Dist-Mode' + '%d' % (IFreq) + '.png'
    plt.savefig(FigName, bbox_inches='tight')
    plt.close()
    FigList.append(FigName)
    FigCaption.append(StrTitle)

    # 根据标准正态分布的累积密度，计算避开k次谐波频率可靠度
    R_delta1 = st.norm.cdf(beta1)  # 低于(1-Δ)倍k次谐波的可靠度
    R_delta2 = st.norm.cdf(beta2)  # 高于(1+Δ)倍k次谐波的可靠度

    # 节径振动避开率可靠度
    R_delta = R_delta1 * R_delta2
    return R_tri, R_delta, OptSMT_str, FigList, FigCaption, TxtList, TxtCaption, TxtHead


# EndDef

def FatigueReliability(XR, YR, FileMater):
    ##读取材料疲劳参数
    txt.insert(END, '\t 读取材料疲劳参数 \n')
    print('\t 读取材料疲劳参数 \n')
    FuncReadFatigueP(FileMater)

    Num1 = XR.shape[0]  # 样本数量

    ##构建疲劳寿命计算的参数矩阵
    txt.insert(END, '\t 构建计算疲劳寿命的参数矩阵 \n')
    print('\t 构建计算疲劳寿命的参数矩阵 \n')
    E0 = XR[:, 1] * 1e3  # 弹性模量[MPa]
    sigma0 = YR[:, 0]  # 最大Mises应力[MPa]
    epsilon0 = YR[:, 1] / 2  # Mises应变幅
    sigma_f0 = np.ones(Num1) * sigma_f  # 疲劳强度系数[MPa]
    b0 = np.ones(Num1) * b  # 疲劳强度指数
    epsilon_f0 = np.ones(Num1) * epsilon_f  # 疲劳塑性系数
    c0 = np.ones(Num1) * c  # 疲劳塑性指数

    ##疲劳寿命计算(多核并行)
    txt.insert(END, '\t 疲劳寿命计算(%d核并行)，样本数目=%3.2e (耗时稍久)\n' % (Num_CPU - 1, Num1))
    txt.see(END)
    print('\t 疲劳寿命计算(%d核并行)，样本数目=%3.2e (耗时稍久)\n' % (Num_CPU - 1, Num1))
    items = zip(E0, epsilon0, sigma0, sigma_f0, b0, epsilon_f0, c0)  # 封装数据
    p = mp.Pool(Num_CPU - 1)
    t0 = timeit.default_timer()
    print('t0')
    NFatig = p.map(do_something, items)
    print('NFatig')
    p.close()
    p.join()
    tn = timeit.default_timer()
    txt.insert(END, '\t\t 完成！耗时:'+str(tn - t0)+'s.\n')
    txt.see(END)
    print('\t\t 完成！耗时:', str(tn - t0), 's.\n')
    np.savetxt('Fatigue.txt',NFatig,fmt="%d",delimiter=" ")
    FigList, FigCaption, TxtList, TxtCaption, TxtHead = OptFit(NFatig)

    return FigList, FigCaption, TxtList, TxtCaption, TxtHead


# EndDef

# 对疲劳寿命分布进行拟合
def OptFit(data):
    FigList = []
    FigCaption = []
    TxtList = []
    TxtCaption = []
    TxtHead = []
    print('\n采用Fitter对疲劳寿命分布进行拟合（耗时略多）\n')
    txt.insert(END, '\n采用Fitter对疲劳寿命分布进行拟合（耗时略多）\n')
    txt.see(END)
    #偏态分布，x最大值太大了，需要截取:根据累积概率密度曲线，取接近最大累积密度的某点(99.9%或99.8%)为x坐标轴最大值
    fig1 = plt.figure()
    x1,y1=sns.kdeplot(data,cumulative=True).get_lines()[0].get_data()
    #plt.show()
    plt.close()
    
    y1max=max(y1)
    index=np.where(y1>y1max*0.995)
    xmax=x1[index[0][0]]
    xmin=int(min(data))

    # 绘制直方图及核密度估计曲线
    fig = plt.figure()
    # 画直方图，核密度
    sns.distplot(data,norm_hist=True,kde=False,bins=500,color="y")
    sns.kdeplot(data,clip=[0,xmax],cumulative=False)
    StrTitle = Object2+'疲劳寿命分布直方图及核密度估计曲线'
    FigName = 'Fig-Fatigue-HistKde.png'
    plt.xlabel('疲劳循环次数')
    plt.ylabel('概率密度')
    plt.title(StrTitle)
    plt.xlim(0, xmax)
    plt.savefig(FigName, bbox_inches='tight')
    plt.close()
    FigList.append(FigName)
    FigCaption.append(StrTitle)

    # 使用Fitter拟合概率分布核密度曲线
    # https://docs.scipy.org/doc/scipy/reference/stats.html
    # https://www.zhihu.com/question/284452109
    distributions = ['genhalflogistic', 'logistic', 'genlogistic', 'gamma', 'invweibull', 'erlang', 'exponweib',
                     'fatiguelife', 'weibull_min', 'weibull_max']
    f = Fitter(data, xmin=0, xmax=xmax, distributions=distributions, timeout=30, bins=300, density=True)  # 创建Fitter类
    f.fit()  # 调用fit函数拟合分布
    df = f.summary(Nbest=5, plot=False)  # 前5个最佳拟合结果，存储到df中(dataframe格式)，不直接绘图
    # f.hist()
    # f.plot_pdf(Nbest=5, lw=2, method='sumsquare_error')

    # 根据拟合参数，绘制前5个最佳核密度拟合曲线
    # 用f.summary(plot=True)也能绘制，这里想试试getattr函数
    x1 = np.linspace(0, xmax, num=int(xmax))
    # https://blog.csdn.net/u014281392/article/details/75331570
    distributions_opt = df.index  # 提取前5个最佳拟合结果的分布名称

    fig = plt.figure()
    # 只画直方图
    sns.distplot(data,norm_hist=True,kde=False,bins=500,color="y")
    # 提取拟合分布的建模数据,并画图
    for index in range(len(distributions_opt)):
        str0 = distributions_opt[index]
        param = f.fitted_param[str0]
        str1 = getattr(st, str0)
        y1 = str1.pdf(x1, *param)
        plt.plot(x1, y1, linewidth=1.0)
    # EndFor
    StrTitle = Object2+'疲劳寿命分布直方图及核密度估计曲线'
    FigName = 'Fig-Fatigue-HistKdeFit.png'
    plt.xlabel('疲劳循环次数')
    plt.ylabel('概率密度')
    plt.title(StrTitle)
    plt.xlim(0, xmax)
    plt.legend(distributions_opt, loc='best')
    plt.savefig(FigName, bbox_inches='tight')
    plt.close()
    FigList.append(FigName)
    FigCaption.append(StrTitle)

    ##累积概率密度曲线画图
    ##先画核密度概率分布曲线
    fig = plt.figure()
    # 概率核密度曲线从0开始的第1段加密，将数据点保存为x1,y1
    x1,y1=sns.kdeplot(data,bw_adjust=0.1,clip=[0,xmin],gridsize=xmin+1,cumulative=False).get_lines()[0].get_data()
    #概率核密度曲线的第2段放粗一些，将数据点保存为x2,y2
    x2,y2=sns.kdeplot(data,bw_adjust=0.1,clip=[xmin+1,xmax],gridsize=300,cumulative=False).get_lines()[1].get_data()
    plt.close()

    # 将两端数据拼接在一起
    x1 = np.r_[x1, x2]  # 拼接数组，增加行数用np.r_[a,b]，增加列数用np.c_[a,b]
    y1 = np.r_[y1, y2]

    ##再对核密度曲线求积分得到累积概率密度
    area = []
    # https://www.cnpython.com/qa/1477247
    # https://vimsky.com/examples/usage/python-scipy.integrate.trapz.html
    # scipy.integrate.trapz(y, x=None, dx=1.0, axis=-1)
    for i in range(len(x1)):
        x = x1[0:i + 1]  # 取x1的第0,1,2...,i列，注意没有取第i+1列
        y = y1[0:i + 1]
        area_tmp = scipy.integrate.trapz(y, x)
        area.append(area_tmp)
    # EndFor
    ##根据累积概率密度曲线，确定可靠性曲线
    y2 = 1.0 - np.c_[area]

    txt.insert(END, '\n画疲劳寿命累积概率密度曲线\n')
    txt.see(END)
    print('\n画疲劳寿命累积概率密度曲线\n')
    fig = plt.figure()
    plt.plot(x1, area, color='red', linewidth=1.0)
    StrTitle = Object2+'疲劳寿命累积概率密度曲线'
    FigName = 'Fig-Fatigue-Cumulative.png'
    plt.xlabel('疲劳循环次数')
    plt.ylabel('概率')
    plt.title(StrTitle)
    plt.xlim(0, xmax)
    plt.savefig(FigName, bbox_inches='tight')
    plt.close()
    FigList.append(FigName)
    FigCaption.append(StrTitle)

    txt.insert(END, '\n画疲劳寿命可靠度曲线\n')
    txt.see(END)
    print('\n画疲劳寿命可靠度曲线\n')
    fig = plt.figure()
    plt.plot(x1, y2, color='red', linewidth=1.0)
    StrTitle = Object2+'疲劳寿命可靠度曲线'
    FigName = 'Fig-Fatigue-Reliability.png'
    plt.xlabel('疲劳循环次数')
    plt.ylabel('可靠度')
    plt.title(StrTitle)
    plt.xlim(0, xmax)
    plt.savefig(FigName, bbox_inches='tight')
    plt.close()
    FigList.append(FigName)
    FigCaption.append(StrTitle)

    txt.insert(END, '\n保存疲劳寿命概率密度曲线到txt文本\n')
    txt.see(END)
    print('\n保存疲劳寿命概率密度曲线到txt文本\n')
    data = np.c_[x1, y1]
    TxtName = 'KrigingOut-Fatigue-KDE.txt'
    TabCaption = Object2+'疲劳寿命概率密度'
    TabHead = ['循环次数', '概率密度']
    comments = '#%s \n#%s     %s' % (TabCaption, TabHead[0], TabHead[1])
    formats='%d      %12.10f'
    SaveTxt(data, TxtName, comments, formats)
    TxtList.append(TxtName)
    TxtCaption.append(TabCaption)
    TxtHead.append(TabHead)

    txt.insert(END, '\n保存疲劳寿命累积概率密度曲线到txt文本\n')
    txt.see(END)
    print('\n保存疲劳寿命累积概率密度曲线到txt文本\n')
    data = np.c_[x1, area]
    TxtName = 'KrigingOut-Fatigue-Cumulative.txt'
    TabCaption = Object2+'疲劳寿命累积概率密度'
    TabHead = ['循环次数', '累积概率密度']
    comments = '#%s \n#%s     %s' % (TabCaption, TabHead[0], TabHead[1])
    formats='%d      %12.10f'
    SaveTxt(data, TxtName, comments, formats)
    TxtList.append(TxtName)
    TxtCaption.append(TabCaption)
    TxtHead.append(TabHead)

    txt.insert(END, '\n保存疲劳寿命可靠度曲线到txt文本\n')
    txt.see(END)
    print('\n保存疲劳寿命可靠度曲线到txt文本\n')
    data = np.c_[x1, y2]
    TxtName = 'KrigingOut-Fatigue-Reliability.txt'
    TabCaption = Object2+'疲劳寿命可靠度'
    TabHead = ['循环次数', '可靠度']
    comments = '#%s \n#%s     %s' % (TabCaption, TabHead[0], TabHead[1])
    formats='%d      %12.10f'
    SaveTxt(data, TxtName, comments, formats)
    TxtList.append(TxtName)
    TxtCaption.append(TabCaption)
    TxtHead.append(TabHead)

    return FigList, FigCaption, TxtList, TxtCaption, TxtHead


# EndDef

def SaveTxt(data, TxtName, comments, formats):
    f = open(TxtName, 'w', encoding='utf-8-sig')  # 打开文件， 用'w'写文件
    f.write('%s\n' % (comments))
    np.savetxt(f, X=data, fmt=formats, delimiter=' ', encoding='utf-8-sig')
    f.close()


# EndDEF


def InsertFig(Doc,Bookmark_Name,Bookmark_Text, Fig, Caption, index):
    # 插入图片
    paragraph = Doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(Fig, width=Inches(5.4))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)  # 段后0磅间距

    # 插入题注与标签
    # str0='图'
    # InsertCaption(doc, caption, index,str0)
    # Bookmark_Name = 'Fig' + str(index)
    # Bookmark_Text = '图'
    # InsertCaption(Doc, Bookmark_Name, Bookmark_Text, index, Caption)
    InsertCaption2(Doc,Bookmark_Name,Bookmark_Text,index,Caption)


# EndDef

# 插入表格
# https://www.freesion.com/article/7261971185/
# https://blog.csdn.net/liaosen/article/details/121242200
def InsertTab(Doc,Bookmark_Name,Bookmark_Text,TxtName, TxtCaption, TxtHead, index):
    # 插入题注与标签
    # InsertCaption(Doc, TxtCaption, index,str0)
    # Bookmark_Name = 'Tab' + str(index)
    # Bookmark_Text = '表'
    # InsertCaption(Doc, Bookmark_Name, Bookmark_Text, index, TxtCaption)
    InsertCaption1(Doc,Bookmark_Name,Bookmark_Text,index,TxtCaption)#题注不含章标题编号

    data = np.loadtxt(TxtName, comments='#', encoding='utf-8-sig')
    rows=data.shape[0]#原始数据的行数
    cols=data.shape[1]#原始数据的列数
    NumAll=rows*cols#原始数据总数
    
    if   cols==2:##如果只有2列，扩展成6列
        rmax=rows#原始数据最大行数
        rows=math.ceil(rows/3)#调整行数
        cols=cols*3#调整列数
        TxtHead=TxtHead+TxtHead+TxtHead
        data0=data#重组data数组
        data=np.zeros([rows,cols])
        #print(rmax,rows,cols)
        data[:,0:2]=data0[0:rows,:]
        data[:,2:4]=data0[rows:rows*2,:]
        data[0:rmax-rows*2,4:6]=data0[rows*2:rmax,:]
    elif cols==3:##如果只有3列，扩展成6列
        rmax=rows#原始数据最大行数
        rows=math.ceil(rows/2)#调整行数
        cols=cols*2#调整列数
        TxtHead=TxtHead+TxtHead
        data0=data#重组data数组
        data=np.zeros([rows,cols])
        #print(rmax,rows,cols)
        data[:,0:3]=data0[0:rows,:]
        data[0:rmax-rows,3:6]=data0[rows:rmax,:]
    elif cols==4:##如果只有4列，扩展成8列
        rmax=rows#原始数据最大行数
        rows=math.ceil(rows/2)#调整行数
        cols=cols*2#调整列数
        TxtHead=TxtHead+TxtHead
        data0=data#重组data数组
        data=np.zeros([rows,cols])
        #print(rmax,rows,cols)
        data[:,0:4]=data0[0:rows,:]
        data[0:rmax-rows,4:8]=data0[rows:rmax,:]

    # 创建表格
    table = Doc.add_table(rows=rows + 1, cols=cols, style='Table Grid')
    # 将所有的单元格抽取出来  #https://theprogrammingexpert.com/write-table-fast-python-docx/
    table_cells = table._cells

    for i_col in range(len(TxtHead)):
        cells=table.cell(0,i_col)
        cells.text="%s" %(TxtHead[i_col])
        cells.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cells.paragraphs[0].runs[0].font.size=Pt(7)
    # EndFor
    for i_row in range(len(data[:, 0])):
        for i_col in range(len(data[0, :])):
            ii=(i_col+1)+i_row*data.shape[1]#计数器，处理到第n个数据
            i_cell=i_col + (i_row+1) * data.shape[1]#表格位置，注意第1行的编号包含进去了
            cells=table_cells[i_cell]
            cells.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ii<=NumAll: #最后一行，可能有空列，需要根据数据总数判断
                cells.text="%8.6e" %(data[i_row,i_col])
            else:
                cells.text="-"
            cells.paragraphs[0].runs[0].font.size=Pt(7)
            cells.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        # EndFor
    # EndFor
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表后插入空白行
    paragraph = Doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(0)  # 段后0磅间距


# EndDef


# 书签锚点起始
def Bookmark_Start(Tag, index, Bookmark_Name):
    bmrk = OxmlElement('w:bookmarkStart')
    bmrk.set(qn('w:id'), str(index))
    bmrk.set(qn('w:name'), Bookmark_Name)
    # return bmrk
    Tag.append(bmrk)


# EndDef

# 书签锚点结束
def Bookmark_End(Tag, index):
    bmrk = OxmlElement('w:bookmarkEnd')
    bmrk.set(qn('w:id'), str(index))
    # return bmrk
    Tag.append(bmrk)


# EndDef

# 创建题注域代码: "图{ SEQ 图 \* ARABIC }"   或者
# 创建题注域代码: "表{ SEQ 表 \* ARABIC }"
def CaptionField1(Tag,str0,index):
    # 创建第1部分："图"  或者 "表"
    text = OxmlElement('w:r')
    text.text = str0
    text.set(qn('xml:space'), 'preserve')
    Tag.append(text)

    # 创建第2部分："{"
    r2 = OxmlElement('w:r')
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r2.append(fldChar)
    Tag.append(r2)

    # 创建第3部分：" SEQ 图 \* ARABIC "  或者 " SEQ 表 \* ARABIC "
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ ' + str0 + ' \* ARABIC '
    instrText.set(qn('xml:space'), 'preserve')
    r2.append(instrText)
    Tag.append(r2)

    # 创建第4部分：
    r2 = OxmlElement('w:r')
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    r2.append(fldChar)
    Tag.append(r2)

    # 创建第5部分：
    r2 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = str(index)
    r2.append(t)
    Tag.append(r2)

    # 创建第6部分："}"
    r2 = OxmlElement('w:r')
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r2.append(fldChar)
    Tag.append(r2)
    ####


# EndDef


#创建题注域代码: "图{ STYLEREF 1 }-{ SEQ 图 \* ARABIC \s 1 }"   或者 
#创建题注域代码: "表{ STYLEREF 1 }-{ SEQ 表 \* ARABIC \s 1 }"
#  包含章标题编号
def CaptionField2(Tag,str0,index):
  
  #创建第1部分："图"  或者 "表"
  text = OxmlElement('w:r')
  text.text = str0
  text.set(qn('xml:space'), 'preserve')
  Tag.append(text)
  
  #创建第2部分：章标题编号
  #创建第2.1部分："{"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r2.append(fldChar)
  Tag.append(r2)
  
  #创建第2.2部分：" STYLEREF 1 "
  r2 = OxmlElement('w:r')
  instrText = OxmlElement('w:instrText')
  instrText.text = ' STYLEREF 1 \s'
  instrText.set(qn('xml:space'), 'preserve')
  r2.append(instrText)
  Tag.append(r2)

  #创建第2.3部分：
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'separate')
  r2.append(fldChar)
  Tag.append(r2)

  #创建第2.4部分：
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = str(index)
  r2.append(t)
  Tag.append(r2)
  
  #创建第2.5部分："}"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r2.append(fldChar)
  Tag.append(r2)
  ####
  
  #创建第3部分："-"
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = '-'
  r2.append(t)
  Tag.append(r2)
  
  #创建第4部分：题注编号
  #创建第4.1部分："{"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r2.append(fldChar)
  Tag.append(r2)
  
  #创建第4.2部分：" SEQ 图 \* ARABIC \s 1 "  或者 " SEQ 表 \* ARABIC \s 1 "
  r2 = OxmlElement('w:r')
  instrText = OxmlElement('w:instrText')
  instrText.text = ' SEQ '+str0+' \* ARABIC \s 1 '
  instrText.set(qn('xml:space'), 'preserve')
  r2.append(instrText)
  Tag.append(r2)

  #创建第4.3部分：
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'separate')
  r2.append(fldChar)
  Tag.append(r2)

  #创建第4.4部分：
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = str(index)
  r2.append(t)
  Tag.append(r2)
  
  #创建第4.5部分："}"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r2.append(fldChar)
  Tag.append(r2)
  ####
#EndDef


#在CaptionField2的题注中包含章标题编号
#下面插入章标题
def InsertHead(Doc):
  #创建段落
  paragraph = Doc.add_paragraph()
  paragraph.add_run().font.color.rgb = RGBColor(0,0,0)
  Tag=paragraph._p
  #段落居中
  #加个空格，防止在题注前回车，出现2行书签
  r1 = OxmlElement('w:pPr')
  fldChar = OxmlElement('w:pStyle')
  fldChar.set(qn('w:val'), 'Heading1')
  r1.append(fldChar)#创建1级标题
  
  r2 = OxmlElement('w:numPr')
  fldChar = OxmlElement('w:ilvl')
  fldChar.set(qn('w:val'), '0')
  r2.append(fldChar)
  fldChar = OxmlElement('w:numId')
  fldChar.set(qn('w:val'), '7')
  r2.append(fldChar)
  r1.append(r2)#创建自动编号
  
  #1级标题的格式
  fldChar = OxmlElement('w:ind')
  fldChar.set(qn('w:left'), '425')
  fldChar.set(qn('w:leftChars'), '0')
  fldChar.set(qn('w:hanging'), '425')
  fldChar.set(qn('w:firstLineChars'), '0')
  r1.append(fldChar)
  ##居中
  #fldChar = OxmlElement('w:jc')
  #fldChar.set(qn('w:val'), 'center')
  #r1.append(fldChar)
  
  #自动编号的颜色为黑色
  r2 = OxmlElement('w:rPr')
  fldChar = OxmlElement('w:color')
  fldChar.set(qn('w:val'), '000000')
  r2.append(fldChar)
  r1.append(r2)
  
  Tag.append(r1)
  
  
  r1 = OxmlElement('w:r')
  #1级标题颜色为黑色
  r2 = OxmlElement('w:rPr')
  fldChar = OxmlElement('w:color')
  fldChar.set(qn('w:val'), '000000')
  r2.append(fldChar)
  r1.append(r2)
  #1级标题的题目
  t = OxmlElement('w:t')
  t.text = "章"
  r1.append(t)
  Tag.append(r1)
  
  #另起一段
  Doc.add_paragraph()
  
#EndDef

def InsertCaption1(Doc,Bookmark_Name,Bookmark_Text,index,Caption):
  #创建段落
  paragraph = Doc.add_paragraph()
  #段落居中
  paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
  #加个空格，防止在题注前回车，出现2行书签
  text = OxmlElement('w:r')
  text.text = ' '
  text.set(qn('xml:space'), 'preserve')
  paragraph._p.append(text)
  
  #书签的位置锚点
  Tag=paragraph._p
  #书签开始
  Bookmark_Start(Tag,index, Bookmark_Name)
  #创建题注域代码
  CaptionField1(Tag,Bookmark_Text,index)#不包含章标题编号
  #书简结束
  Bookmark_End(Tag,index)
  
  #题注文字部分
  paragraph.add_run(" " + Caption)
  paragraph.paragraph_format.space_after=Pt(0)  #段后0磅间距
#EndDef

# 插入题注与标签
#插入域，再插入书签，最后引用书签
#https://blog.csdn.net/weixin_39742727/article/details/109914739
#https://blog.csdn.net/igoizzz/article/details/117905625
def InsertCaption2(Doc,Bookmark_Name,Bookmark_Text,index,Caption):
    # 创建段落
    paragraph = Doc.add_paragraph()
    # 段落居中
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 加个空格，防止在题注前回车，出现2行书签
    text = OxmlElement('w:r')
    text.text = ' '
    text.set(qn('xml:space'), 'preserve')
    paragraph._p.append(text)

    # 书签的位置锚点
    Tag = paragraph._p
    # 书签开始
    Bookmark_Start(Tag, index, Bookmark_Name)
    # 创建题注域代码
    CaptionField2(Tag, Bookmark_Text, index)
    # 书签结束
    Bookmark_End(Tag, index)

    # 题注文字部分
    paragraph.add_run(" " + Caption)
    paragraph.paragraph_format.space_after = Pt(0)  # 段后0磅间距


# EndDef

# 创建交叉引用的域代码："{ REF Fig1 \h }"
def ReferenceField(Tag, Bookmark_Name, Bookmark_Text, index):
    # 创建第1部分："{"
    r2 = OxmlElement('w:r')
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r2.append(fldChar)
    Tag.append(r2)

    # 创建第2部分：" REF Fig1 \h "，其中Bookmark_Name是书签名称
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' REF ' + Bookmark_Name + ' \h '
    r2.append(instrText)
    Tag.append(r2)

    # 创建第3部分：
    r2 = OxmlElement('w:r')
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    r2.append(fldChar)
    Tag.append(r2)

    # 创建第4部分：" REF Fig1 \h "，其中Bookmark_Text是 引用的书签内容
    r2 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = Bookmark_Text + str(index)
    r2.append(t)
    Tag.append(r2)

    # 创建第5部分："}"
    r2 = OxmlElement('w:r')
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r2.append(fldChar)
    Tag.append(r2)


# EndDef

def CrossReference(Doc, _Text, Bookmark_Name, Bookmark_Text, index):
    # 创建段落
    paragraph = Doc.add_paragraph()

    run = paragraph.add_run(_Text + "如")
    Tag = run._r
    ReferenceField(Tag, Bookmark_Name, Bookmark_Text, index)
    paragraph.add_run("所示。")
    paragraph.paragraph_format.space_after = Pt(0)  # 段后0磅间距


# EndDef

#在页脚中添加页码
#https://stackoverflow.com/questions/50776715/setting-pgnumtype-property-in-python-docx-is-without-effect
def AddFooterNumber(run):
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'Page'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "Seq"
    fldChar2.append(t)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

#在文档中设置页脚
#https://baijiahao.baidu.com/s?id=1665454009794833226&wfr=spider&for=pc
def InsertPageNumber(Doc):
  footer = Doc.sections[0].footer # 获取第一个节的页脚
  footer.is_linked_to_previous = True  #编号续前一节
  paragraph = footer.paragraphs[0] # 获取页脚的第一个段落
  paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#页脚居中对齐
  run_footer=paragraph.add_run() # 添加页脚内容
  AddFooterNumber(run_footer)
  font = run_footer.font
  font.name = 'Times New Roman'#新罗马字体
  font.size = Pt(10)#10号字体
  font.bold = False#加粗
#EndDef

# 在docx中新建：图的交叉引用，表的交叉引用，图及其题注，表及其题注，
def AddFigTab(Exec,OptOut_str,FigList,FigCaption,TxtList,TxtCaption,TxtHead,DocxName):
    # path = os.getcwd()
    # print(path)
    Doc = Document()
    Doc.styles['Normal'].font.name = u'宋体'
    Doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    Doc.styles['Normal'].font.size = Pt(10)  # 小四
    Doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    #设置文档页码
    InsertPageNumber(Doc)

    #插入章标题，自动编号，后续题注中使用
    #打开docx文件时，需要"更新域"2次
    InsertHead(Doc)

    #文字说明
    for i in range(0, len(OptOut_str)):
        paragraph = Doc.add_paragraph()
        run = paragraph.add_run(OptOut_str[i])
        paragraph.paragraph_format.space_after=Pt(0)  #段后0磅间距

    # 建立图的交叉引用
    for i in range(0, len(FigList)):
        index = i + 1
        Bookmark_Name=Object2+Exec+'Fig'+str(index)
        Bookmark_Text = '图'
        CrossReference(Doc, FigCaption[i], Bookmark_Name, Bookmark_Text, index)

    #新建图及其题注
    for i in range(0, len(FigList)):
        index=i+1
        Bookmark_Name=Object2+Exec+'Fig'+str(index)
        Bookmark_Text='图'
        InsertFig(Doc,Bookmark_Name,Bookmark_Text,FigList[i],FigCaption[i],index)
        
    paths=DocxName+"-图.docx"
    Doc.save(paths)
  
    ##输出表格
    Doc = Document()
    Doc.styles['Normal'].font.name = u'宋体'
    Doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    Doc.styles['Normal'].font.size = Pt(10)# 小四
    Doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)

    #设置文档页码
    InsertPageNumber(Doc)

    # 建立表的交叉引用
    for i in range(0, len(TxtList)):
        index = i + 1
        Bookmark_Name=Object2+Exec+'Tab'+str(index)
        Bookmark_Text = '表'
        CrossReference(Doc, TxtCaption[i], Bookmark_Name, Bookmark_Text, index)

    paragraph = Doc.add_paragraph()
    #段落居中
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # run = paragraph.add_run("附表")
    # run.bold=True
    # run.font.name = u'宋体'
    # run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # run.font.color.rgb = RGBColor(0, 0, 0)

    # 新建表及其题注
    for i in range(0, len(TxtList)):
        # TxtCaption是2维列表
        index = i + 1
        Bookmark_Name=Object2+Exec+'Tab'+str(index)
        Bookmark_Text='表'
        InsertTab(Doc,Bookmark_Name,Bookmark_Text,TxtList[i],TxtCaption[i],TxtHead[i],index)

    paths = DocxName + "-表.docx"
    Doc.save(paths)
    return paths
# EndDef
# 以上为自定义的辅助函数AllDef

def popupmenu(event):
    mainmenu.post(event.x_root, event.y_root)

#关于
def about():
    aboutNew = Toplevel(root)
    aboutNew.geometry('300x200')
    aboutNew.title('关于')
    l0 = Label(aboutNew, text='', anchor="center")
    l0.pack(anchor="center")
    l1 = Label(aboutNew,bitmap="info",anchor="center")
    l1.pack(anchor="center")
    l2 = Label(aboutNew, text='涡轮泵可靠性计算',anchor="center",font="Helvetic 20 bold")
    l2.pack(anchor="center")
    l0 = Label(aboutNew, text='', anchor="center")
    l0.pack(anchor="center")
    l3 = Label(aboutNew, text='版本号：V1.01    July 3,2022创建',anchor="center")
    l3.pack(anchor="center")
    l4 = Label(aboutNew, text='技术支持：合肥工业大学WLB',anchor="center")
    l4.pack(anchor="center")

def guidebook():
    guideNew = Toplevel(root)
    guideNew.geometry('400x400')
    guideNew.title('指南')
    l0 = Label(guideNew, text='', anchor="center")
    l0.pack()
    m1 = Message(guideNew,text ='1."样本文件1"用于建立小样本模型，"样本文件2"用于验证代理模型，"概率分布文件"用于代理模型进行可靠性计算。',
                 width="360")
    m1.pack(padx=20,anchor="nw")
    m2 = Message(guideNew, text = '2.拉丁超立方抽样的抽样点数需为素数的平方，本软件内置算法将智能寻找距离输入点最接近的坐标。', width="360")
    m2.pack(padx=20, anchor="nw")
# 文件路径
# 输入量
Num1 = int(1e5)  # 可靠性计算大样本抽样数量
point = 300 # 拉丁超立方抽样默认抽样点数量
# root=""
RootPath = ""   # 保存目录
File1RootPath = ""  # 样本文件1的路径，如果没有保存目录，则设为保存目录

FileLHS = ""
LHSRoot = ""

FileStaGauss = ""  # 保存随机变量信息的文本，用于可靠性计算大样本抽样
FileStatic1 = ""  # Abaqus输出静强度样本文件1，用于构建SMT模型
FileStatic2 = ""  # Abaqus输出静强度样本文件2，用于校核SMT模型

FileFreqGauss = ""
FileFreq1 = ""  # Abaqus输出固有频率文件1，用于构建SMT模型
FileFreq2 = ""  # Abaqus输出固有频率文件2，用于校核SMT模型

FileFatGauss = ""
FileStaticFat1 = ""
FileStaticFat2 = ""
FileMater="" # 存储GH4169材料的4个疲劳系数

#拉丁超立方抽样
def pop_up_boxLHS():

    # global btn4
    winNew = Toplevel(root)
    winNew.geometry('300x100')
    winNew.title('设置抽样点数')

    def com():
        global point
        try:
            num = int(e1.get())  # 获取e1的值，转为浮点数，如果不能转捕获异常
            if (num > 1000):
                messagebox.showwarning('警告', '超出范围')
            else:
                point = int(e1.get())
                lbLHS2.config(text='抽样点数：' + str(point))
                winNew.destroy()
        except:
            messagebox.showwarning('警告', '请输入数字')

    var = tkinter.StringVar()  # 这即是输入框中的内容
    var.set(str(point))
    e1 = Entry(winNew, textvariable=var)
    e1.pack()
    Button(winNew, text='保存', command=com).pack()
    l1 = Label(winNew, text='只能数字。最大值为1000')
    l1.pack()


def chFileHLS():
    global FileLHS
    FileLHS = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目')
    lbLHS4.config(text='')
    if FileLHS != '':
        lbLHS1.config(text='抽样样本文件：' + FileLHS)
    else:
        lbLHS1.config(text='抽样样本文件：未选择')

def chLhsRoot():
    global RootPath
    RootPath = tkinter.filedialog.askdirectory()
    lbLHS4.config(text='')
    if RootPath != '':
        lbLHS3.config(text='保存路径：' + RootPath)
    else:
        lbLHS3.config(text='保存路径：未选择')

def StartLHS():
    if FileLHS != "" and RootPath != "": #if判断文件/文件路径是否为空，下同
        warnings.filterwarnings("ignore")
        btnLHS4.config(state=DISABLED)
        global txt
        txt = txt0
        txt.config(state=NORMAL)
        try:
            t1 = threading.Thread(target=LHS)
            t1.setDaemon(True)
            t1.start()
        except:
            txt.insert(END, '开始失败\n')
            txt.see(END)
            txt.config(state=DISABLED)
            btnLHS4.config(state=NORMAL)
    else:
        lbLHS4.config(text='请选择文件或文件路径')

def LHS():
    #程序标记
    Exec1="Sampling"
    Exec2="抽样"
    ##输入量
    DocxName = Object2+Exec2
    FigList = []
    FigCaption = []
    TxtList = []
    TxtCaption = []
    TxtHead = []

    # 清除不必要的文件
    ClearFileDir()

    # 自定义画图设置
    FigConfig()

    print('读取原始数据 \n')
    txt.insert(END, '读取原始数据 \n')
    txt.see(END)
    # 读取文本文件，存入矩阵。分隔符为','，注释符为'#'
    data = np.loadtxt(FileLHS, comments='#', encoding='utf-8-sig')

    print('生成正态分布的平均值和标准差 \n')
    txt.insert(END, '生成正态分布的平均值和标准差 \n')
    txt.see(END)
    # 密度平均值(kg/m3)
    rho_avg = data[0, 0]
    # 密度变异系数
    rho_cv = data[0, 1]
    # 弹模平均值(GPa)
    E_avg = data[1, 0]
    # 弹模变异系数
    E_cv = data[1, 1]
    # 转速平均值(rpm)
    rpm_avg = data[2, 0]
    # 转速变异系数
    rpm_cv = data[2, 1]

    # 密度标准差
    rho_std = rho_avg * rho_cv
    # 密度方差
    rho_var = rho_std * rho_std
    # 弹模标准差
    E_std = E_avg * E_cv
    # 弹模方差
    E_var = E_std * E_std
    # 转速标准差
    rpm_std = rpm_avg * rpm_cv
    # 转速方差
    rpm_var = rpm_std * rpm_std

    # 根据正态分布的3σ准则确定样本抽样的上下限
    A = np.array([rho_avg, E_avg, rpm_avg])
    B = np.array([rho_std, E_std, rpm_std])
    l_bounds = A - 4.0 * B  # 样本缩放的下限
    u_bounds = A + 4.0 * B  # 样本缩放的上限
    # print(l_bounds,'\n')
    # print(u_bounds,'\n')
    # num0必须是素数：2, 3, 5, 7, 11, 13, 17, 19, 23, 29, 31
    num0 = [2, 3, 5, 7, 11, 13, 17, 19, 23, 29, 31]
    Num0 = np.c_[num0] ** 2

    # 在素数平方中找到距离NumSample最近的素数平方
    temp = abs(Num0 - point)
    print(point)
    b = np.where(temp == min(temp))
    index = b[0][0]  # 最接近的位置坐标
    Num01 = Num0[index, 0]  # 抽样组的数目(采用强度2时，数目必须素数的平方)

    dim1 = A.shape[0]  # 抽样维度的数目
    # print(dim1)
    StrNum = '%03d' % (Num01)

    print('生成%d个优化拉丁超立方样本，强度为2 \n' % (Num01))
    txt.insert(END, '生成%d个优化拉丁超立方样本，强度为2 \n' % (Num01))
    txt.see(END)
    # 2维优化拉丁超立方，强度为2
    sampler = qmc.LatinHypercube(d=dim1, strength=2, optimization="random-cd")
    sample0 = sampler.random(n=Num01)
    error1 = qmc.discrepancy(sample0)
    sample = qmc.scale(sample0, l_bounds, u_bounds)  # 样本缩放到上下限。
    # print(l_bounds)
    # print(u_bounds)
    # print(np.min(sample[:,0]),np.min(sample[:,1]),np.min(sample[:,2]))
    # print(np.max(sample[:,0]),np.max(sample[:,1]),np.max(sample[:,2]))

    print('绘制样本第1、2维数据图 \n')
    txt.insert(END, '绘制样本第1、2维数据图 \n')
    txt.see(END)
    str0 = '%.3e' % error1
    StrTitle = Object2+'优化拉丁超立方抽样第1、2维数据(强度2,' + '差异度' + str0 + ')'
    fig = plt.figure()  # 定义一个图像窗口
    plt.plot(sample[:, 0], sample[:, 1], '.', markersize=4, color='blue')
    # plt.legend(loc="upper left")
    # plt.title(str1,fontsize=7)
    plt.title(StrTitle)
    plt.xlabel('密度 [kg/m^3]')
    plt.ylabel('弹性模量 [GPa]')
    FigName = 'Fig-LatinHypercube-Num' + StrNum + '-Dim1vDim2.png'
    plt.savefig(FigName, bbox_inches='tight')
    # plt.savefig('fig1.pdf', bbox_inches='tight')
    plt.close()
    FigList.append(FigName)
    FigCaption.append(StrTitle)

    print('绘制样本第2、3维数据图 \n')
    txt.insert(END, '绘制样本第2、3维数据图 \n')
    txt.see(END)
    fig = plt.figure()  # 定义一个图像窗口
    StrTitle = Object2+ '优化拉丁超立方抽样第2、3维数据(强度2,' + '差异度' + str0 + ')'
    plt.plot(sample[:, 1], sample[:, 2], '.', markersize=4, color='blue')
    # plt.legend(loc="upper left")
    plt.title(StrTitle)
    plt.xlabel('弹性模量 [GPa]')
    plt.ylabel('转速 [r/min]')
    FigName = 'Fig-LatinHypercube-Num' + StrNum + '-Dim2vDim3.png'
    plt.savefig(FigName, bbox_inches='tight')
    plt.close()
    FigList.append(FigName)
    FigCaption.append(StrTitle)

    TxtName = 'AbaqusInputLatinHypercube-Num' + StrNum + '.txt'
    TabCaption= Object2+'抽样数据(n='+'%d'%(Num1)+')'
    TabHead=['样本编号','密度[kg/m^3]','弹性模量[GPa]','转速[r/min]']
    print('保存样本到文件: "%s" \n' % (TxtName))
    txt.insert(END, '保存样本到文件: "%s" \n' % (TxtName))
    txt.see(END)
    XX=np.arange(1,len(sample[:,0])+1,1)
    data=np.c_[XX,sample]
    f = open(TxtName, 'w', encoding='utf-8-sig')  # 打开文件， 用'w'写文件
    f.write('#%s\n' % (TabCaption))
    f.write('#      %s         ' % (TabHead[0]))
    f.write('       %s         ' % (TabHead[1]))
    f.write('       %s         ' %(TabHead[2]))
    f.write('       %s         \n' %(TabHead[3]))
    np.savetxt(f, X=data, delimiter=' ')
    f.close()
    TxtList.append(TxtName)
    TxtCaption.append(TabCaption)
    TxtHead.append(TabHead)

    OptOut_str=[]
    print('图片和表格输出到文件: "%s.docx" \n' % (DocxName))
    txt.insert(END, '图片和表格输出到文件: "%s.docx" \n' % (TxtName))
    txt.see(END)
    AddFigTab(Exec1,OptOut_str,FigList, FigCaption, TxtList, TxtCaption, TxtHead, DocxName)
    print('完成 \n')
    txt.insert(END, '完成 \n\n\n')
    txt.see(END)
    txt.insert(END, '完成 \n\n\n')
    txt.see(END)
    lbLHS4.config(text='完成，图片和表格输出到文件: "%s.docx" \n' % (DocxName))
    txt.config(state=DISABLED)
    btnLHS4.config(state=NORMAL)

#静强度可靠性
# 静强度可靠性界面选择抽样样本文件
def chFileStaGauss():
    global FileStaGauss
    FileStaGauss = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目')
    lbSta5.config(text='')
    if FileStaGauss != '':
        lbSta3.config(text='抽样点样本文件：' + FileStaGauss)
    else:
        lbSta3.config(text='抽样点样本文件：未选择')

# 选择静强度样本文件1
def chStaticFile1():
    global FileStatic1
    FileStatic1 = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目')
    lbSta5.config(text='')
    if FileStatic1 != '':
        lbSta1.config(text='样本文件1：' + FileStatic1)
    else:
        lbSta1.config(text='样本文件1：未选择')

# 选择静强度样本文件2
def chStaticFile2():
    global FileStatic2
    FileStatic2 = tkinter.filedialog.askopenfilename()
    lbSta5.config(text='')
    if FileStatic2 != '':
        lbSta2.config(text='样本文件2：' + FileStatic2)
    else:
        lbSta2.config(text='样本文件2：未选择')

def chStaRoot():
    global RootPath
    RootPath = tkinter.filedialog.askdirectory()
    lbSta5.config(text='')
    if RootPath != '':
        lbSta4.config(text='保存路径：' + RootPath)
    else:
        lbSta4.config(text='保存路径：未选择')

# 开始计算静强度可靠性
def StartStatic():
    if FileStaGauss != "" and FileStatic1 != "" and FileStatic2 != "" and RootPath != "":
        btn21.config(state=DISABLED)
        global txt
        txt = txt1
        txt.config(state=NORMAL)
        # 忽略warnings,不然在cmd窗口会输出一堆warning
        warnings.filterwarnings("ignore")
        try:
            t1 = threading.Thread(target=static)
            t1.setDaemon(True)
            t1.start()
        except:
            txt.insert(END, '开始失败\n')
            txt.config(state=DISABLED)
            btn21.config(state=NORMAL)
    else:
        lbSta5.config(text='请选择文件或文件路径')

def static():
    #程序标记
    Exec1="Static"
    Exec2="静强度可靠性"

    # 清除不必要的文件
    ClearFileDir()

    # 自定义画图设置
    FigConfig()

    DocxName = Object2+Exec2

    txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + " 开始分析\n")
    txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + '\n 预测大样本的最大Mises应力应变 \n')
    txt.see(END)
    XR,YR,OptOut_str1,FigList1,FigCaption1,TxtList,TxtCaption,TxtHead= \
        StressStrain(Exec2,Num1, FileStaGauss, FileStatic1, FileStatic2)
    S=YR[:,0]#提取应力

    txt.insert(END, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime()) + '\n 计算大样本的静强度可靠度 \n')
    R, FigList2, FigCaption2 = StressReliability(S, FileStaGauss)

    txt.insert(END, '图片和表格输出到文件: "%s.docx" \n' % (DocxName))
    OptOut_str=OptOut_str1
    FigList = FigList1 + FigList2
    FigCaption = FigCaption1 + FigCaption2
    path = AddFigTab(Exec1,OptOut_str,FigList, FigCaption, TxtList, TxtCaption, TxtHead, DocxName)
    txt.insert(END, '完成！ \n\n\n')
    txt.see(END)
    lbSta5.config(text='完成，图片和表格输出到文件: "%s.docx" \n' % (DocxName))
    txt.config(state=DISABLED)
    btn21.config(state=NORMAL)

# 固有频率
# 选择抽样样本文件(固有频率)
def chFileFreqGauss():
    global FileFreqGauss
    FileFreqGauss = tkinter.filedialog.askopenfilename()
    lbFreq5.config(text='')
    if FileFreqGauss != '':
        lbFreq3.config(text='抽样点样本文件：' + FileFreqGauss)
    else:
        lbFreq3.config(text='脚本文件：未选择')

# 选择固有频率样本文件1
def chFreqFile1():
    global FileFreq1
    FileFreq1 = tkinter.filedialog.askopenfilename()
    lbFreq5.config(text='')
    if FileFreq1 != '':
        lbFreq1.config(text='样本文件1：' + FileFreq1)
    else:
        lbFreq1.config(text='样本文件1：未选择') 

# 选择固有频率样本文件2
def chFreqFile2():
    global FileFreq2
    FileFreq2 = tkinter.filedialog.askopenfilename()
    lbFreq5.config(text='')
    if FileFreq2 != '':
        lbFreq2.config(text='样本文件2：' + FileFreq2)
    else:
        lbFreq2.config(text='样本文件2：未选择')

def chFreqRoot():
    global RootPath
    RootPath = tkinter.filedialog.askdirectory()
    lbFreq5.config(text='')
    if RootPath != '':
        lbFreq4.config(text='保存路径：' + RootPath)
    else:
        lbFreq4.config(text='保存路径：未选择')

# 开始计算振动可靠性
def StartFrequency():
    if FileFreq1 != "" and FileFreq2 != "" and FileFreqGauss != "" and RootPath != "":
        btn22.config(state=DISABLED)
        global txt
        txt = txt2
        txt.config(state=NORMAL)
        # 忽略warnings,不然在cmd窗口会输出一堆warning
        warnings.filterwarnings("ignore")
        try:
            t1 = threading.Thread(target=frequency)
            t1.setDaemon(True)
            t1.start()
        except:
            txt.insert(END, '开始失败\n')
            txt.config(state=DISABLED)
            btn22.config(state=NORMAL)
    else:
        lbFreq5.config(text='请选择文件或文件路径\n')

def frequency():
    #程序标记
    Exec1="Freq"
    Exec2="振动可靠性"

    DocxName=Object2+Exec2
    # 清除不必要的文件
    ClearFileDir()
    # 自定义画图设置
    FigConfig()

    ######################
    # 构建Kriging代理模型
    ######################
    # 读取Abaqus计算的频率结果，构建kriging模型
    txt.insert(END, '\n\n读取Abaqus计算的频率结果 \n')
    txt.see(END)
    print('\n\n读取Abaqus计算的频率结果 \n')
    # 读取第1个abaqus计算的频率结果，样本数少，用于构建kriging模型
    txt.insert(END, '\t 读取数据构建kriging模型 \n')
    txt.see(END)
    print('\t 读取数据构建kriging模型 \n')
    X1, Freq1 = GetDataFreq(FileFreq1)
    # 读取第2个abaqus计算的频率结果，样本数多，用于校核kriging模型
    txt.insert(END, '\t 读取数据校核kriging模型 \n')
    txt.see(END)
    print('\t 读取数据校核kriging模型 \n')
    X2, Freq2 = GetDataFreq(FileFreq2)

    # 谐波频率避开率Δ
    delta = 0.05  # 避开率Δ

    # 提取各阶节径振型固有频率，分别构建kriging模型，并校核
    II = []
    R1 = []
    R2 = []
    OptOut_str=[]
    FigList = []
    FigCaption = []
    TxtList = []
    TxtCaption = []
    TxtHead = []
    for index in range(len(Freq1[0])):
        Y1 = np.c_[Freq1[:, index]]  # 提取某阶节径振型固有频率，并转为列向量
        Y2 = np.c_[Freq2[:, index]]  # 校核用的
        II.append(index + 1)
        R_tri,R_delta,OptSMT_str1,FigList1,FigCaption1,TxtList1,TxtCaption1,TxtHead1=\
            FrequencyReliability(index+1,X1,Y1,X2,Y2,delta,Num1,FileFreqGauss)
        R1.append(R_tri)
        R2.append(R_delta)
        OptOut_str.extend(OptSMT_str1)
        FigList.extend(FigList1)
        FigCaption.extend(FigCaption1)
        TxtList.extend(TxtList1)
        TxtCaption.extend(TxtCaption1)
        TxtHead.extend(TxtHead1)
        # FigCaption=FigCaption+FigCaption1
    # EndFor
    R = np.c_[R1] * np.c_[R2]

    # 保存可靠性结果
    txt.insert(END, '\n保存各节径振动可靠度到txt文本\n')
    txt.see(END)
    print('\n保存各节径振动可靠度到txt文本\n')
    data=np.c_[II,II,R,R1,R2]
    StrFile = 'KrigingOut-Frequency-Reliability.txt'
    TxtCaption1=Object2+'前%d节径振动可靠度' %(index+1)
    TxtHead1=['编号','节径','总可靠度','三重点共振可靠度', '谐波频率避开率%3.2f%%可靠度' %(delta*100)]
    TxtList.append(StrFile)
    TxtCaption.append(TxtCaption1)
    TxtHead.append(TxtHead1)
    f = open(StrFile, 'w', encoding='utf-8-sig')  # 打开文件， 用'w'写文件
    f.write('#%s\n'   %(TxtCaption1))
    f.write('#    %s     %s    %s    %s    %s\n' %(TxtHead1[0],TxtHead1[1],TxtHead1[2],TxtHead1[3],TxtHead1[4]))
    #np.savetxt(f, X=data, fmt='%d     %d       %f            %f            %f', delimiter=' ', encoding='utf-8-sig')
    np.savetxt(f, X=data, delimiter=' ', encoding='utf-8-sig')
    f.close()
    path = AddFigTab(Exec1,OptOut_str,FigList, FigCaption, TxtList, TxtCaption, TxtHead, DocxName)
    txt.insert(END, '图片和表格输出到文件: "%s.docx" \n' % (DocxName))
    txt.insert(END, '完成！ \n\n\n' )
    txt.see(END)
    lbFreq5.config(text='完成，图片和表格输出到文件: "%s.docx" \n' % (DocxName))
    btn22.config(state=NORMAL)
    txt.config(state=DISABLED)

# 疲劳强度
# 选择抽样样本文件(疲劳强度)
def chFileFatGauss():
    global FileFatGauss
    FileFatGauss = tkinter.filedialog.askopenfilename()
    lbFat6.config(text='')
    if FileFatGauss != '':
        lbFat3.config(text='抽样点样本文件：' + FileFatGauss)
    else:
        lbFat3.config(text='抽样点样本文件：未选择')

# 选择疲劳强度样本文件1
def chFatFile1():
    global FileStaticFat1
    FileStaticFat1 = tkinter.filedialog.askopenfilename()
    lbFat6.config(text='')
    if FileStaticFat1 != '':
        lbFat1.config(text='样本文件1：' + FileStaticFat1)
    else:
        lbFat1.config(text='样本文件1：未选择')

# 选择疲劳强度样本文件2
def chFatFile2():
    global FileStaticFat2
    FileStaticFat2 = tkinter.filedialog.askopenfilename()
    lbFat6.config(text='')
    if FileStaticFat2 != '':
        lbFat2.config(text='样本文件2：' + FileStaticFat2)
    else:
        lbFat2.config(text='样本文件2：未选择')

# 存储GH4169材料的4个疲劳系数
def chFileMater():
    global FileMater
    FileMater = tkinter.filedialog.askopenfilename(initialdir='C:/Users/Leo/Desktop/WLB项目')
    lbFat6.config(text='')
    if FileMater != '':
        lbFat4.config(text='材料系数文件：' + FileMater)
    else:
        lbFat4.config(text='材料系数文件：未选择')

def chFatRoot():
    global RootPath
    RootPath = tkinter.filedialog.askdirectory()
    lbFat6.config(text='')
    if RootPath != '':
        lbFat5.config(text='保存路径：' + RootPath)
    else:
        lbFat5.config(text='保存路径：未选择')

#计算疲劳强度可靠性
def StartFatigue():
    if FileStaticFat1 != "" and FileStaticFat2 != "" and FileFatGauss != "" and FileMater != "" and RootPath != "":
        btn23.config(state=DISABLED)
        global txt
        txt = txt3
        txt.config(state=NORMAL)
        # 忽略warnings,不然在cmd窗口会输出一堆warning
        warnings.filterwarnings("ignore")

        try:
            t1 = threading.Thread(target=fatigue)
            t1.setDaemon(True)
            t1.start()
        except:
            txt.insert(END, '开始失败\n')
            txt.config(state=DISABLED)
            btn23.config(state=NORMAL)
    else:
        lbFat6.config(text='请选择文件或文件路径')

def fatigue():
    #程序标记
    Exec1="Fatigue"
    Exec2="疲劳可靠性"
    DocxName = Object2+Exec2
    ClearFileDir()
    FigConfig()
    txt.insert(END, '\n 预测大样本的最大Mises应力应变 \n')
    txt.see(END)
    print('\n 预测大样本的最大Mises应力应变 \n')
    XR,YR,OptOut_str1,FigList1,FigCaption1,TxtList1,TxtCaption1,TxtHead1 = StressStrain(Exec2,Num1, FileFatGauss, FileStaticFat1, FileStaticFat2)

    txt.insert(END, '\n 预测大样本的疲劳寿命 \n')
    txt.see(END)
    print('\n 预测大样本的疲劳寿命 \n')
    FigList2, FigCaption2, TxtList2, TxtCaption2, TxtHead2 = FatigueReliability(XR, YR, FileMater)
    txt.insert(END, '图片和表格输出到文件: "%s.docx" \n\n\n' % (DocxName))
    txt.see(END)
    txt.insert(END, '完成！\n\n\n\n\n')
    txt.see(END)
    print('图片和表格输出到文件: "%s.docx" \n' % (DocxName))
    OptOut_str=OptOut_str1
    FigList = FigList1 + FigList2
    FigCaption = FigCaption1 + FigCaption2
    TxtList = TxtList1 + TxtList2
    TxtCaption = TxtCaption1 + TxtCaption2
    TxtHead = TxtHead1 + TxtHead2
    path = AddFigTab(Exec1,OptOut_str,FigList, FigCaption, TxtList, TxtCaption, TxtHead, DocxName)
    lbFat6.config(text='完成，图片和表格输出到文件: "%s.docx" \n' % (DocxName))
    btn23.config(state=NORMAL)
    txt.config(state=DISABLED)


if __name__ == '__main__':
    # global root
    mp.freeze_support()
    root = Tk()
    root.title('涡轮一体化设计平台')
    root.geometry('800x600')
    root.state('zoomed')
    tab_main = ttk.Notebook()  # 创建分页栏
    tab_main.place(relx=0, rely=0, relwidth=1, relheight=1)
    FluentFrame = Frame(tab_main)  # 创建fluent框架
    tab_main.add(FluentFrame, text='流动分析')  # 将fluent页插入分页栏中
    AbaqusFrame = Frame(tab_main)  # 创建第一页框架
    tab_main.add(AbaqusFrame, text='结构热分析')  # 将第一页插入分页栏中
    AbaqusFrame2 = Frame(tab_main)  # 创建第一页框架
    tab_main.add(AbaqusFrame2, text='静强度分析')  # 将第一页插入分页栏中
    AbaqusFrame3 = Frame(tab_main)  # 创建第一页框架
    tab_main.add(AbaqusFrame3, text='整圈热分析')  # 将第一页插入分页栏中
    AbaqusFrame4 = Frame(tab_main)  # 创建第一页框架
    tab_main.add(AbaqusFrame4, text='振动分析')  # 将第一页插入分页栏中
    Frame0 = Frame(tab_main)  # 创建第一页框架
    tab_main.add(Frame0, text='拉丁超立方抽样')  # 将第一页插入分页栏中
    AbaqusFrame5 = Frame(tab_main)  # 创建第一页框架
    tab_main.add(AbaqusFrame5, text='静强度批量分析')  # 将第一页插入分页栏中
    AbaqusFrame6 = Frame(tab_main)  # 创建第一页框架
    tab_main.add(AbaqusFrame6, text='振动批量分析')  # 将第一页插入分页栏中
    Frame1 = Frame(tab_main)  # 创建第二页框架
    tab_main.add(Frame1, text='静强度可靠性')  # 将第二页插入分页栏中
    Frame2 = Frame(tab_main)  # 创建第三页框架
    tab_main.add(Frame2, text='振动可靠性')  # 将第三页插入分页栏中
    Frame3 = Frame(tab_main)  # 创建第四页框架
    tab_main.add(Frame3, text='疲劳可靠性')  # 将第四页插入分页栏中

    # 导航栏
    mainmenu = Menu(root)
    menuFile = Menu(mainmenu)  # 菜单分组 menuFile
    helpLabel = Menu(mainmenu)
    mainmenu.add_cascade(label="文件", menu=menuFile)
    menuFile.add_command(label="退出", command=root.destroy)
    mainmenu.add_cascade(label="帮助", menu=helpLabel)
    helpLabel.add_command(label="指南",command=guidebook)
    helpLabel.add_command(label="关于",command=about)

    # Fluent窗口
    lb1 = Label(FluentFrame, text='Fluent路径：未选择')
    lb1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lb2 = Label(FluentFrame, text='脚本文件路径：未选择')
    lb2.place(relx=0.2, rely=0.1, relwidth=0.8, relheight=0.1)
    lb3 = Label(FluentFrame, text='计算核数'+str(sema)+'<计算核数'+str(cpuCore))
    lb3.place(relx=0.2, rely=0.18, relwidth=0.8, relheight=0.1)

    DelCheckVar = IntVar()
    ch = Checkbutton(FluentFrame, text='分析时显示界面', variable=DelCheckVar, onvalue=1, offvalue=0)
    ch.place(relx=0.1, rely= 0.35)

    btn1 = Button(FluentFrame, text='选择Fluent路径', command=ChFluentPath)
    btn1.place(relx=0.1, rely=0.04, relwidth=0.2, relheight=0.05)
    btn2 = Button(FluentFrame, text='选择脚本文件', command=ChScriptPath)
    btn2.place(relx=0.1, rely=0.12, relwidth=0.2, relheight=0.05)
    btn4 = Button(FluentFrame, text='设置核数', command=pop_up_box)
    btn4.place(relx=0.1, rely=0.20, relwidth=0.2, relheight=0.05)
    btn2 = Button(FluentFrame, text='开始分析', command=Start)
    btn2.place(relx=0.1, rely=0.4, relwidth=0.8, relheight=0.05)

    txtFluent = Text(FluentFrame)
    txtFluent.config(state=DISABLED)
    txtFluent.place(rely=0.5, relheight=0.5, relwidth=1.0)
    txtFluent.tag_config("tag_1", foreground="red")

    #sress abaqus makeThermal
    lbThermal1 = Label(AbaqusFrame, text='Abaqus路径：未选择')
    lbThermal1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lbThermal2 = Label(AbaqusFrame, text='脚本文件路径：未选择')
    lbThermal2.place(relx=0.2, rely=0.1, relwidth=0.8, relheight=0.1)
    lbThermal3 = Label(AbaqusFrame, text='')
    lbThermal3.place(relx=0.2, rely=0.18, relwidth=0.8, relheight=0.1)
    lbThermal4 = Label(AbaqusFrame, text='网格文件：TurbineThermalStatic.ansa.zst.inp',anchor='w',justify='left')
    lbThermal4.place(relx=0.1, rely=0.26, width=320, relheight=0.05)
    lbThermal5 = Label(AbaqusFrame, text='温度数据文件：FSI-Temp.txt',anchor='w',justify='left')
    lbThermal5.place(relx=0.1, rely=0.31, width=320, relheight=0.05)
    #lbThermal6 = Label(AbaqusFrame, text='压力数据文件：FSI-Pres.txt',anchor='w',justify='left')
    #lbThermal6.place(relx=0.1, rely=0.36, width=320, relheight=0.05)

    btnThermal1 = Button(AbaqusFrame, text='选择Abaqus路径', command=ChAbaqusPathT)
    btnThermal1.place(relx=0.1, rely=0.04, relwidth=0.2, relheight=0.05)
    btnThermal2 = Button(AbaqusFrame, text='选择脚本文件', command=ChScriptPathT)
    btnThermal2.place(relx=0.1, rely=0.12, relwidth=0.2, relheight=0.05)
    btnT2 = Button(AbaqusFrame, text='开始分析', command=StartAbaqusThermal)
    btnT2.place(relx=0.1, rely=0.4, relwidth=0.8, relheight=0.05)

    txtAbaqus1 = Text(AbaqusFrame)
    txtAbaqus1.config(state=DISABLED)
    txtAbaqus1.place(rely=0.5, relheight=0.5, relwidth=1.0)
    txtAbaqus1.tag_config("tag_1", foreground="red")
    
    # stress abaqus makeStatic
    lbStatic1 = Label(AbaqusFrame2, text='Abaqus路径：未选择')
    lbStatic1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lbStatic2 = Label(AbaqusFrame2, text='脚本文件路径：未选择')
    lbStatic2.place(relx=0.2, rely=0.1, relwidth=0.8, relheight=0.1)
    lbStatic3 = Label(AbaqusFrame2, text='')
    lbStatic3.place(relx=0.2, rely=0.18, relwidth=0.8, relheight=0.1)
    lbStatic4 = Label(AbaqusFrame2, text='网格文件：TurbineThermalStatic.ansa.zst.inp',anchor='w',justify='left')
    lbStatic4.place(relx=0.1, rely=0.26, width=320, relheight=0.05)
    lbStatic5 = Label(AbaqusFrame2, text='odb文件：TurbineThermal.odb',anchor='w',justify='left')
    lbStatic5.place(relx=0.1, rely=0.31, width=320, relheight=0.05)
    lbStatic6 = Label(AbaqusFrame2, text='压力数据文件：FSI-Pres.txt',anchor='w',justify='left')
    lbStatic6.place(relx=0.1, rely=0.36, width=320, relheight=0.05)
    
    btnStatic1 = Button(AbaqusFrame2, text='选择Abaqus路径', command=ChAbaqusPathS)
    btnStatic1.place(relx=0.1, rely=0.04, relwidth=0.2, relheight=0.05)
    btnStatic2 = Button(AbaqusFrame2, text='选择脚本文件', command=ChScriptPathS)
    btnStatic2.place(relx=0.1, rely=0.12, relwidth=0.2, relheight=0.05)
    # btn23 = Button(AbaqusFrame2, text='选择脚本文件根目录', command=ChScriptRootPath)
    # btn23.place(relx=0.1, rely=0.20, relwidth=0.2, relheight=0.05)
    btnStatic3 = Button(AbaqusFrame2, text='开始分析', command=StartAbaqusStatic)
    btnStatic3.place(relx=0.1, rely=0.4, relwidth=0.8, relheight=0.05)
    
    txtAbaqus2 = Text(AbaqusFrame2)
    txtAbaqus2.config(state=DISABLED)
    txtAbaqus2.place(rely=0.5, relheight=0.5, relwidth=1.0)
    txtAbaqus2.tag_config("tag_1", foreground="red")
    
    #freq abaqus makeThermal
    lbFreqThermal1 = Label(AbaqusFrame3, text='Abaqus路径：未选择')
    lbFreqThermal1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lbFreqThermal2 = Label(AbaqusFrame3, text='脚本文件路径：未选择')
    lbFreqThermal2.place(relx=0.2, rely=0.1, relwidth=0.8, relheight=0.1)
    lbFreqThermal3 = Label(AbaqusFrame3, text='')
    lbFreqThermal3.place(relx=0.2, rely=0.18, relwidth=0.8, relheight=0.1)
    lbFreqThermal4 = Label(AbaqusFrame3, text='单级网格文件：Turbine-Freq-SingleBlade.inp',anchor='w',justify='left')
    lbFreqThermal4.place(relx=0.1, rely=0.26, width=320, relheight=0.05)
    lbFreqThermal5 = Label(AbaqusFrame3, text='整体网格文件：Turbine-Freq-Whole.inp',anchor='w',justify='left')
    lbFreqThermal5.place(relx=0.1, rely=0.31, width=320, relheight=0.05)
    lbFreqThermal6 = Label(AbaqusFrame3, text='温度数据文件：FSI-Temp.txt',anchor='w',justify='left')
    lbFreqThermal6.place(relx=0.1, rely=0.36, width=320, relheight=0.05)

    btnFreqThermal1 = Button(AbaqusFrame3, text='选择Abaqus路径', command=ChAbaqusFreqPathT)
    btnFreqThermal1.place(relx=0.1, rely=0.04, relwidth=0.2, relheight=0.05)
    btnFreqThermal2 = Button(AbaqusFrame3, text='选择脚本文件', command=ChScriptFreqPathT)
    btnFreqThermal2.place(relx=0.1, rely=0.12, relwidth=0.2, relheight=0.05)
    btnFT2 = Button(AbaqusFrame3, text='开始分析', command=StartFreqThermal)
    btnFT2.place(relx=0.1, rely=0.4, relwidth=0.8, relheight=0.05)

    txtAbaqusFreq1 = Text(AbaqusFrame3)
    txtAbaqusFreq1.config(state=DISABLED)
    txtAbaqusFreq1.place(rely=0.5, relheight=0.5, relwidth=1.0)
    txtAbaqusFreq1.tag_config("tag_1", foreground="red")
    
    #freq abaqus makeStatic
    lbFreqStatic1 = Label(AbaqusFrame4, text='Abaqus路径：未选择')
    lbFreqStatic1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lbFreqStatic2 = Label(AbaqusFrame4, text='脚本文件路径：未选择')
    lbFreqStatic2.place(relx=0.2, rely=0.1, relwidth=0.8, relheight=0.1)
    lbFreqStatic3 = Label(AbaqusFrame4, text='')
    lbFreqStatic3.place(relx=0.2, rely=0.18, relwidth=0.8, relheight=0.1)
    lbFreqStatic4 = Label(AbaqusFrame4, text='网格文件：Turbine-Freq-Whole.inp',anchor='w',justify='left')
    lbFreqStatic4.place(relx=0.1, rely=0.26, width=320, relheight=0.05)
    lbFreqStatic5 = Label(AbaqusFrame4, text='odb文件：TurbineThermal.odb',anchor='w',justify='left')
    lbFreqStatic5.place(relx=0.1, rely=0.31, width=320, relheight=0.05)
    lbFreqStatic6 = Label(AbaqusFrame4, text='压力数据文件：FSI-Pres.txt',anchor='w',justify='left')
    lbFreqStatic6.place(relx=0.1, rely=0.36, width=320, relheight=0.05)
    
    btnFreqStatic1 = Button(AbaqusFrame4, text='选择Abaqus路径', command=ChAbaqusFreqPathS)
    btnFreqStatic1.place(relx=0.1, rely=0.04, relwidth=0.2, relheight=0.05)
    btnFreqStatic2 = Button(AbaqusFrame4, text='选择脚本文件', command=ChScriptFreqPathS)
    btnFreqStatic2.place(relx=0.1, rely=0.12, relwidth=0.2, relheight=0.05)
    # btn23 = Button(AbaqusFrame2, text='选择脚本文件根目录', command=ChScriptRootPath)
    # btn23.place(relx=0.1, rely=0.20, relwidth=0.2, relheight=0.05)
    btnFreqStatic3 = Button(AbaqusFrame4, text='开始分析', command=StartFreqStatic)
    btnFreqStatic3.place(relx=0.1, rely=0.4, relwidth=0.8, relheight=0.05)
    
    txtAbaqusFreq2 = Text(AbaqusFrame4)
    txtAbaqusFreq2.config(state=DISABLED)
    txtAbaqusFreq2.place(rely=0.5, relheight=0.5, relwidth=1.0)
    txtAbaqusFreq2.tag_config("tag_1", foreground="red")

    # 拉丁超立方抽样窗口
    lbLHS1 = Label(Frame0, text='概率分布文件：未选择')
    lbLHS1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lbLHS2 = Label(Frame0, text='抽样点数：'+str(point))
    lbLHS2.place(relx=0.2, rely=0.10, relwidth=0.8, relheight=0.1)
    lbLHS3 = Label(Frame0, text='文件保存路径：未选择')
    lbLHS3.place(relx=0.2, rely=0.21, relwidth=0.8, relheight=0.1)
    lbLHS5 = Label(Frame0, text='脚本会自动寻找距离给定点数最近的拉丁超立方坐标。例：给定300，找到最近坐标为289，将抽样289个样本点')
    lbLHS5.place(relx=0.03, rely=0.18, relwidth=0.6, relheight=0.05)

    btnLHS1 = Button(Frame0, text='选择概率分布文件（用于样本抽样）', command=chFileHLS)
    btnLHS1.place(relx=0.1, rely=0.04, relwidth=0.2, relheight=0.05)
    btnLHS2 = Button(Frame0, text='设置抽样点数', command=pop_up_boxLHS)
    btnLHS2.place(relx=0.1, rely=0.12, relwidth=0.2, relheight=0.05)
    btnLHS3 = Button(Frame0, text='选择文件保存路径', command=chLhsRoot)
    btnLHS3.place(relx=0.1, rely=0.24, relwidth=0.2, relheight=0.05)
    btnLHS4 = Button(Frame0, text='开始分析', command=StartLHS)
    btnLHS4.place(relx=0.1, rely=0.32, relwidth=0.8, relheight=0.05)
    lbLHS4 = Label(Frame0, text='')
    lbLHS4.place(relx=0.1, rely=0.40, relwidth=0.8, relheight=0.1)

    txt0 = Text(Frame0)
    txt0.config(state=DISABLED)
    txt0.place(rely=0.5, relheight=0.5, relwidth=1.0)
    txt0.tag_config("tag_1", foreground="red")

    # moreStatic
    lbMoreStatic1 = Label(AbaqusFrame5, text='Abaqus路径：未选择')
    lbMoreStatic1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lbMoreStatic2 = Label(AbaqusFrame5, text='脚本文件路径：未选择')
    lbMoreStatic2.place(relx=0.2, rely=0.1, relwidth=0.8, relheight=0.1)
    lbMoreStatic3 = Label(AbaqusFrame5, text='')
    lbMoreStatic3.place(relx=0.2, rely=0.18, relwidth=0.8, relheight=0.1)
    lbMoreStatic4 = Label(AbaqusFrame5, text='CAE文件：TurbineStatic.cae',anchor='w',justify='left')
    lbMoreStatic4.place(relx=0.1, rely=0.26, width=320, relheight=0.05)
    lbMoreStatic5 = Label(AbaqusFrame5, text='ODB文件：TurbineThermal.odb',anchor='w',justify='left')
    lbMoreStatic5.place(relx=0.1, rely=0.31, width=320, relheight=0.05)
    lbMoreStatic6 = Label(AbaqusFrame5, text='抽样数据文件：AbaqusInputLatinHypercube.txt',anchor='w',justify='left')
    lbMoreStatic6.place(relx=0.1, rely=0.36, width=320, relheight=0.05)
    lbMoreStatic7 = Label(AbaqusFrame5, text='温度数据文件：fsi-temp.txt')
    lbMoreStatic7.place(relx=0.48, rely=0.26, width=320, relheight=0.05)
    lbMoreStatic8 = Label(AbaqusFrame5, text='压力数据文件：fsi-pres.txt')
    lbMoreStatic8.place(relx=0.48, rely=0.31, width=320, relheight=0.05)
    
    btnMoreStatic1 = Button(AbaqusFrame5, text='选择Abaqus路径', command=ChAbaqusMorePathS)
    btnMoreStatic1.place(relx=0.1, rely=0.04, relwidth=0.2, relheight=0.05)
    btnMoreStatic2 = Button(AbaqusFrame5, text='选择脚本文件', command=ChScriptMorePathS)
    btnMoreStatic2.place(relx=0.1, rely=0.12, relwidth=0.2, relheight=0.05)
    # btn23 = Button(AbaqusFrame2, text='选择脚本文件根目录', command=ChScriptRootPath)
    # btn23.place(relx=0.1, rely=0.20, relwidth=0.2, relheight=0.05)
    btnMoreStatic3 = Button(AbaqusFrame5, text='开始分析', command=StartMoreStatic)
    btnMoreStatic3.place(relx=0.1, rely=0.43, relwidth=0.8, relheight=0.05)
    
    txtAbaqusMore = Text(AbaqusFrame5)
    txtAbaqusMore.config(state=DISABLED)
    txtAbaqusMore.place(rely=0.5, relheight=0.5, relwidth=1.0)
    txtAbaqusMore.tag_config("tag_1", foreground="red")

    # moreFreqStatic
    lbFreqMoreStatic1 = Label(AbaqusFrame6, text='Abaqus路径：未选择')
    lbFreqMoreStatic1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lbFreqMoreStatic2 = Label(AbaqusFrame6, text='脚本文件路径：未选择')
    lbFreqMoreStatic2.place(relx=0.2, rely=0.1, relwidth=0.8, relheight=0.1)
    lbFreqMoreStatic3 = Label(AbaqusFrame6, text='')
    lbFreqMoreStatic3.place(relx=0.2, rely=0.18, relwidth=0.8, relheight=0.1)
    lbFreqMoreStatic4 = Label(AbaqusFrame6, text='CAE文件：TurbineStaticFreq.cae',anchor='w',justify='left')
    lbFreqMoreStatic4.place(relx=0.1, rely=0.26, width=320, relheight=0.05)
    lbFreqMoreStatic5 = Label(AbaqusFrame6, text='ODB文件：TurbineThermal.odb',anchor='w',justify='left')
    lbFreqMoreStatic5.place(relx=0.1, rely=0.31, width=320, relheight=0.05)
    lbFreqMoreStatic6 = Label(AbaqusFrame6, text='抽样数据文件：AbaqusInputLatinHypercube.txt',anchor='w',justify='left')
    lbFreqMoreStatic6.place(relx=0.1, rely=0.36, width=320, relheight=0.05)
    
    btnFreqMoreStatic1 = Button(AbaqusFrame6, text='选择Abaqus路径', command=ChAbaqusFreqMorePathS)
    btnFreqMoreStatic1.place(relx=0.1, rely=0.04, relwidth=0.2, relheight=0.05)
    btnFreqMoreStatic2 = Button(AbaqusFrame6, text='选择脚本文件', command=ChScriptFreqMorePathS)
    btnFreqMoreStatic2.place(relx=0.1, rely=0.12, relwidth=0.2, relheight=0.05)
    # btn23 = Button(AbaqusFrame2, text='选择脚本文件根目录', command=ChScriptRootPath)
    # btn23.place(relx=0.1, rely=0.20, relwidth=0.2, relheight=0.05)
    btnFreqMoreStatic3 = Button(AbaqusFrame6, text='开始分析', command=StartMoreFreq)
    btnFreqMoreStatic3.place(relx=0.1, rely=0.4, relwidth=0.8, relheight=0.05)
    
    txtAbaqusMore2 = Text(AbaqusFrame6)
    txtAbaqusMore2.config(state=DISABLED)
    txtAbaqusMore2.place(rely=0.5, relheight=0.5, relwidth=1.0)
    txtAbaqusMore2.tag_config("tag_1", foreground="red")


    # 静强度窗口
    lbSta1 = Label(Frame1, text='静强度样本文件1：未选择')
    lbSta1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lbSta2 = Label(Frame1, text='静强度样本文件2：未选择')
    lbSta2.place(relx=0.2, rely=0.12, relwidth=0.8, relheight=0.1)
    lbSta3 = Label(Frame1, text='选择概率分布文件：未选择')
    lbSta3.place(relx=0.2, rely=0.22, relwidth=0.8, relheight=0.1)
    lbSta4 = Label(Frame1, text='文件保存路径：未选择')
    lbSta4.place(relx=0.2, rely=0.32, relwidth=0.8, relheight=0.1)
    # lbSta6 = Label(Frame1, text='小样本，用于建立代理模型')
    # lbSta6.place(relx=0.1, rely=0.09, relwidth=0.2, relheight=0.05)
    # lbSta7 = Label(Frame1, text='大样本，用于验证代理模型')
    # lbSta7.place(relx=0.1, rely=0.19, relwidth=0.2, relheight=0.05)
    # lbSta8 = Label(Frame1, text='用于抽样100000样本')
    # lbSta8.place(relx=0.1, rely=0.29, relwidth=0.2, relheight=0.05)


    #DelCheckVar = IntVar()
    #ch = Checkbutton(Frame1, text='分析时显示界面', variable=DelCheckVar, onvalue=1, offvalue=0)
    #ch.place(relx=0.1, rely=0.35)

    btn1 = Button(Frame1, text='选择静强度样本文件1（小样本，用于建立代理模型）', command=chStaticFile1)
    btn1.place(relx=0.1, rely=0.04, relwidth=0.3, relheight=0.05)

    btn2 = Button(Frame1, text='选择静强度样本文件2（大样本，用于验证代理模型）', command=chStaticFile2)
    btn2.place(relx=0.1, rely=0.14, relwidth=0.3, relheight=0.05)

    btn4 = Button(Frame1, text='选择概率分布文件（用于抽样100000样本）', command=chFileStaGauss)
    btn4.place(relx=0.1, rely=0.24, relwidth=0.3, relheight=0.05)

    btn5 = Button(Frame1, text='选择文件保存路径', command=chStaRoot)
    btn5.place(relx=0.1, rely=0.34, relwidth=0.3, relheight=0.05)
    btn21 = Button(Frame1, text='开始分析', command=StartStatic)
    btn21.place(relx=0.1, rely=0.42, relwidth=0.8, relheight=0.05)
    lbSta5 = Label(Frame1, text='')
    lbSta5.place(relx=0.1, rely=0.47, relwidth=0.8, relheight=0.1)

    txt1 = Text(Frame1)
    txt1.config(state=DISABLED)
    txt1.place(rely=0.57, relheight=0.43, relwidth=1.0)
    txt1.tag_config("tag_1", foreground="red")

    # 振动可靠性窗口
    lbFreq1 = Label(Frame2, text='振动样本文件1：未选择')
    lbFreq1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lbFreq2 = Label(Frame2, text='振动样本文件2：未选择')
    lbFreq2.place(relx=0.2, rely=0.09, relwidth=0.8, relheight=0.1)
    lbFreq3 = Label(Frame2, text='选择概率分布文件：未选择')
    lbFreq3.place(relx=0.2, rely=0.26, relwidth=0.8, relheight=0.1)
    lbFreq4 = Label(Frame2, text='文件保存路径：未选择')
    lbFreq4.place(relx=0.2, rely=0.33, relwidth=0.8, relheight=0.1)
    # lbFreq6 = Label(Frame2, text='小样本，用于建立代理模型')
    # lbFreq6.place(relx=0.1, rely=0.09, relwidth=0.2, relheight=0.05)
    # lbFreq7 = Label(Frame2, text='大样本，用于验证代理模型')
    # lbFreq7.place(relx=0.1, rely=0.19, relwidth=0.2, relheight=0.05)
    # lbFreq8 = Label(Frame2, text='用于抽样100000样本')
    # lbFreq8.place(relx=0.1, rely=0.29, relwidth=0.2, relheight=0.05)


    btn1 = Button(Frame2, text='选择振动样本文件1（小样本，用于建立代理模型）', command=chFreqFile1)
    btn1.place(relx=0.1, rely=0.04, relwidth=0.3, relheight=0.05)
    btn2 = Button(Frame2, text='选择振动样本文件2（大样本，用于验证代理模型）', command=chFreqFile2)
    btn2.place(relx=0.1, rely=0.12, relwidth=0.3, relheight=0.05)
    lbFreqEn = Label(Frame2, text='请输入前六节径在文件中所在的列数：')
    lbFreqEn.place(relx=0.04, rely=0.20, relwidth=0.28, relheight=0.05)
    addr = StringVar(value='1,4,7,11,20,29')
    en = Entry(Frame2, textvariable=addr)
    en.place(relx=0.28, rely=0.20, relwidth=0.12, relheight=0.05)
    btn4 = Button(Frame2, text='选择概率分布文件（用于抽样100000样本）', command=chFileFreqGauss)
    btn4.place(relx=0.1, rely=0.28, relwidth=0.3, relheight=0.05)
    btn5 = Button(Frame2, text='选择文件保存路径', command=chFreqRoot)
    btn5.place(relx=0.1, rely=0.36, relwidth=0.3, relheight=0.05)
    btn22 = Button(Frame2, text='开始分析', command=StartFrequency)
    btn22.place(relx=0.1, rely=0.44, relwidth=0.8, relheight=0.05)
    lbFreq5 = Label(Frame2, text='')
    lbFreq5.place(relx=0.1, rely=0.49, relwidth=0.8, relheight=0.1)

    txt2 = Text(Frame2)
    txt2.config(state=DISABLED)
    txt2.place(rely=0.59, relheight=0.41, relwidth=1.0)
    txt2.tag_config("tag_1", foreground="red")

    # 疲劳强度可靠性窗口
    lbFat1 = Label(Frame3, text='静强度样本文件1：未选择')
    lbFat1.place(relx=0.2, rely=0.02, relwidth=0.8, relheight=0.1)
    lbFat2 = Label(Frame3, text='静强度样本文件2：未选择')
    lbFat2.place(relx=0.2, rely=0.10, relwidth=0.8, relheight=0.1)
    lbFat3 = Label(Frame3, text='选择概率分布文件：未选择')
    lbFat3.place(relx=0.2, rely=0.18, relwidth=0.8, relheight=0.1)
    lbFat4 = Label(Frame3, text='材料系数文件：未选择')
    lbFat4.place(relx=0.2, rely=0.26, relwidth=0.8, relheight=0.1)
    lbFat5 = Label(Frame3, text='文件保存路径：未选择')
    lbFat5.place(relx=0.2, rely=0.34, relwidth=0.8, relheight=0.1)
    # lbFat7 = Label(Frame3, text='小样本，用于建立代理模型')
    # lbFat7.place(relx=0.1, rely=0.09, relwidth=0.2, relheight=0.05)
    # lbFat8 = Label(Frame3, text='大样本，用于验证代理模型')
    # lbFat8.place(relx=0.1, rely=0.19, relwidth=0.2, relheight=0.05)
    # lbFat9 = Label(Frame3, text='用于抽样100000样本')
    # lbFat9.place(relx=0.1, rely=0.29, relwidth=0.2, relheight=0.05)
    # lbFat10 = Label(Frame3, text='由《中国航空材料手册》查得')
    # lbFat10.place(relx=0.1, rely=0.39, relwidth=0.2, relheight=0.05)


    btn1 = Button(Frame3, text='选择静强度样本文件1（小样本，用于建立代理模型）', command=chFatFile1)
    btn1.place(relx=0.1, rely=0.04, relwidth=0.3, relheight=0.05)
    btn2 = Button(Frame3, text='选择静强度样本文件2（大样本，用于验证代理模型）', command=chFatFile2)
    btn2.place(relx=0.1, rely=0.12, relwidth=0.3, relheight=0.05)
    btn3 = Button(Frame3, text='选择概率分布文件（用于抽样100000样本）', command=chFileFatGauss)
    btn3.place(relx=0.1, rely=0.20, relwidth=0.3, relheight=0.05)
    btn5 = Button(Frame3, text='选择材料疲劳数据文件（由《中国航空材料手册》查得）', command=chFileMater)
    btn5.place(relx=0.1, rely=0.28, relwidth=0.3, relheight=0.05)
    btn5 = Button(Frame3, text='选择文件保存路径', command=chFatRoot)
    btn5.place(relx=0.1, rely=0.36, relwidth=0.3, relheight=0.05)
    btn23 = Button(Frame3, text='开始分析', command=StartFatigue,)
    btn23.place(relx=0.1, rely=0.44, relwidth=0.8, relheight=0.05)
    lbFat6 = Label(Frame3, text='')
    lbFat6.place(relx=0.1, rely=0.49, relwidth=0.8, relheight=0.1)

    txt3 = Text(Frame3)
    txt3.config(state=DISABLED)
    txt3.place(rely=0.59, relheight=0.41, relwidth=1.0)
    txt3.tag_config("tag_1", foreground="red")

    print("pageloop")
    root.config(menu=mainmenu)
    root.bind('Button-3', popupmenu)  # 根窗体绑定鼠标右击响应事件
    root.mainloop()
   