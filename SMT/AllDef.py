# -*- coding: UTF-8 -*-
import matplotlib.pyplot as plt
import math
import numpy   as np
import pandas  as pd
import seaborn as sns
import scipy.stats as st
import scipy.integrate
import timeit
import pickle
import multiprocessing as mp
import smt.surrogate_models 
import shutil 
import glob
import os 
from fitter import Fitter
from pathlib import Path
from scipy.optimize import fsolve
from smt.surrogate_models import RBF,IDW,RMTB,RMTC,LS,QP, KRG, KPLS, KPLSK
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT                                                 
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

#系统CPU核数
Num_CPU=int(mp.cpu_count()/2)
#分析对象
Object1="Turbine"#涡轮
Object2="涡轮"#涡轮

def ClearFileDir():
  print('\t 清理不必要的文件 \n')  
  filedir=[]
  #以SMT-开头的文件夹：构建IDW, RBF, RMTB, RMTC模型，临时保存目录data_dir
  #以.pkl结尾的文件：除了(IDW, RBF, RMTB, RMTC)之外的模型，临时保存在pkl文件
  #*.png和KrigingOut*.txt是输出
  filedir=glob.glob("SMT-*") +\
      glob.glob("*.pkl") +\
      glob.glob("*.png") +\
      glob.glob("KrigingOut*.txt")
  #print(filedir)
  for path in filedir:
    if os.path.isdir(path):#目录用shutil.rmtree删除
      shutil.rmtree(path)
    elif os.path.isfile(path):#文件用os.remove删除
      os.remove(path)
  #EndFor
#EndDef

def FigConfig():
  config = {
      "font.family": 'serif',
      "font.size": 9,
      "legend.fontsize":7,#图例字体尺寸
      "axes.titlesize":7,#坐标轴字体尺寸
      "font.weight": 'normal',
      "font.serif": ['SimSun'],#中文宋体
      "savefig.dpi":500,#保存图片的分辨率
      "figure.figsize":(6.0, 3.0),#保存图片的尺寸
      "axes.unicode_minus": 'False'#正常显示坐标轴正负号（minus），unicode显示负号异常
      }
  plt.rcParams.update(config)
#EndDef

#根据Abaqus输入参数的随机特征，进行大样本抽样，样本数目为Num1。
#目前针对密度、弹性模量和转速进行抽样
def ReliabSampling(Num1,FileGauss,_Text):
  FigList=[]
  FigCaption=[]

  #用于可靠性分析的抽样
  #读取文本文件，存入矩阵。分隔符为' '，注释符为'#'
  #第1列为平均值，第2列为变异系数
  data = np.loadtxt(FileGauss, comments='#',encoding='utf-8-sig')
  length=data.shape[0]
  data=data[0:length-1,:]
  
  #样本平均值
  Mean1=data[:,0]
  #样本方差
  Vars1=(data[:,0]*data[:,1])**2#(平均值*变异系数)^2
  #print(data[:,0]*data[:,1])
  
  #构建单位对角线矩阵
  Cov1  = np.eye(Mean1.shape[0])
  #提取对角线元素的坐标
  row,col= np.diag_indices_from(Cov1)
  #修改对角线元素为方差
  Cov1[row,col]=Vars1
  
  #采用多维正态分布抽样方法，构建样本
  data = np.random.multivariate_normal(Mean1,Cov1,Num1)

  #画图
  #检测密度、弹性模量、转速的抽样结果，保存为png图片
  data1=data[:,0]
  data2=data[:,1]
  data3=data[:,2]
  StrVar=['密度','弹性模量','转速']
  StrUnit=[' [kg/m^3]',' [GPa]',' [r/min]']
  StrFilename=['rho','E','rpm']
  for index in range(len(data[0,:])):
    DistX=data[:,index]
    
    fig = plt.figure()  # 定义一个图像窗口
    #直方图+概率密度+拟合曲线
    dist=sns.distplot(DistX,hist=True,kde=False, fit=st.norm, fit_kws={'color':'red', 'label':'正态分布拟合','linestyle':'-'})
    plt.xlabel(StrVar[index]+StrUnit[index])
    plt.ylabel('')
    Str1=StrVar[index]+'抽样'
    Str2='(拟合平均值%.2f' %(st.norm.fit(DistX)[0])
    Str3='标准差%.2f' %(st.norm.fit(DistX)[1])+')'
    StrTitle=_Text+Str1+Str2+','+Str3
    plt.title(StrTitle)
    FigName="Fig-Sample-Dist-"+StrFilename[index]+".png"
    ##显示图例
    plt.legend()
    #plt.show()
    plt.savefig(FigName, bbox_inches='tight')
    plt.close()
    FigList.append(FigName)
    FigCaption.append(StrTitle)
  #EndFor
  return data,FigList,FigCaption
##end def

#从文件中读取GH4169材料的4个疲劳系数
#(用于局部应力应变法求疲劳寿命)
def FuncReadFatigueP(FileMater):
  global sigma_f,b,epsilon_f,c
  data = np.loadtxt(FileMater, comments='#',encoding='utf-8-sig')
  #疲劳强度系数[MPa]745.4
  sigma_f=data[0]
  #疲劳强度指数-0.092
  b=data[1]
  #疲劳塑性系数0.161
  epsilon_f=data[2]
  #疲劳塑性指数-0.419
  c=data[3]
#def

#定义局部应力应变法求疲劳寿命的函数
#SWT疲劳寿命函数Func1(x)=0
def Func1(x):
  return sigma_f**2/E*x**(2*b) + epsilon_f*sigma_f*x**(b+c) -epsilon_a*sigma_max
##end def

#python并行计算用的循环体：循环内命令彼此独立，不相干。
#读取data中的数据，作为疲劳寿命计算的参数，
#调用fsolve求解疲劳寿命函数，求得疲劳计算中的工作次数
def do_something(data):
  global E,epsilon_a,sigma_max,sigma_f,b,epsilon_f,c
  E    =data[0]#弹模模量
  epsilon_a=data[1]#应变幅
  sigma_max=data[2]#最大应力
  sigma_f  =data[3]#疲劳强度系数[MPa]
  b    =data[4]#疲劳强度指数
  epsilon_f=data[5]#疲劳塑性系数
  c    =data[6]#疲劳塑性指数
  #print(E,epsilon_a,sigma_max,sigma_f,b,epsilon_f,c)
  #Func1函数中自变量x=2*NFatigue
  #NFatigue为疲劳计算中的工作次数
  NFatigue = fsolve(Func1, 1,xtol=1e-8)/2
  return NFatigue[0]
#end def

def GetDataFreq(StrFile):
  data=np.loadtxt(StrFile, dtype=float, skiprows=1,encoding='utf-8-sig')
  X0   =data[:,len(data[0])-3:len(data[0])]#密度、弹模、转速
  #前6阶节径振型频率存储的位置
  locs=np.array([1,4,7,11,20,29])
  locs=locs-1#python 数组第1个下标是0，
  #print(locs)
  
  #从频率文件中提取前6阶节径振型固有频率
  for index in range(len(locs)):
    ii=locs[index]
    if index==0:
      Freq0=np.c_[data[:,ii]]
    else:
      Freq0=np.c_[Freq0,data[:,ii]]#拼接矩阵
  #EndFor
  print('\t 该文件共有%d组数据 \n' %(len(X0[:,0])))
  
  return X0,Freq0
#EndDef

def Limits(X):
  for index in range(X.shape[1]):
    xmin=min(X[:,index])
    xmax=max(X[:,index])
    temp=np.array([[xmin,xmax]])
    if index==0:
      xlimits=np.r_[temp]
    else:
      xlimits=np.r_[xlimits,temp]
    #EndIf
  #EndFor
  return  xlimits
#EndDef

def mkdir(path):
  # 去除首位空格
  path=path.strip()
  # 去除尾部 \ 符号
  path=path.rstrip("\\")
  # 判断路径是否存在
  isExists=os.path.exists(path)
  # 判断结果
  if not isExists:
    os.makedirs(path) 
    print('\t\t\t %s创建成功\n' %(path))
    return True
  else:
    # 如果目录存在则不创建，并提示目录已存在
    shutil.rmtree(path)
    print('\t\t\t %s虽然存在，但已经删除，并重新创建\n' %(path))
    os.makedirs(path) 
    return False
  #EndIf
#EndDef

#搜索最优的SMT方法
def OptSMT(X,Y,X1,Y1,Strs,ratio):
  FigList=[]
  FigCaption=[]
  TxtList=[]
  TxtCaption=[]
  TxtHead=[]
  
  xlimits=Limits(X)#确定输入量X的极值，n行2列的数据
  StrVar=Strs[0]#SMax,EMax,Mode1
  if   StrVar[0]=='S':
    StrVarOut='最大Miese应力'
    StrType='Static'
    StrUnit='MPa'
  elif StrVar[0]=='E':
    StrVarOut='最大Miese应变'
    StrType='Static'
    StrUnit=''
  elif StrVar[0]=='M':
    StrVarOut='第'+StrVar[-1]+'节径振型固有频率'
    StrType='Frequency'
    StrUnit='Hz'
  #EndIf
  #一些设置参数
  IDW_p=5.0
  RMTB_order=3
  RMTB_num=15
  RBF_d0=5
  RBF_poly=1
  #用来构建代理模型的几个方法名称
  SMTModels=['RBF','IDW','RMTB','RMTC','LS','QP','KRG','KPLS','KPLSK']
  #SMTModels=['RMTC']
  #存储相对均方根误差，求其最小值。
  RMSEMin=[]
  #存储SMT的plk文件名
  StrFileDir=[]
  for index in range(len(SMTModels)):
    t0 = timeit.default_timer()
    StrSMT=SMTModels[index]
    print('\t\t 测试SMT模型：%s \n' %(StrSMT))    
    
    #IDW, RBF, RMTB, RMTC模型用data_dir=FileDir可以临时保存，
    #其它SMT模型需要保存到pkl文件，
    #后续提取最佳SMT模型时直接调用data_dir，加载pkl文件，不用再训练
    if (StrSMT[0]!='I' and StrSMT[0]!='R'):
      FileDir = 'SMT-'+StrType+'-'+StrVar+'-'+StrSMT+'.pkl'
    else:
      FileDir = 'SMT-'+StrType+'-'+StrVar+'-'+StrSMT
    #EndIf
    
    #调用SMT中的方法，getattr函数能够根据字符串调用函数
    str1=getattr(smt.surrogate_models, StrSMT)
    if   StrSMT=='RBF': #RBF, data_dir用于保存模型
      mkdir(FileDir)
      StrFileDir.append(FileDir)
      sm = str1(print_global=False,data_dir=FileDir,d0=RBF_d0,poly_degree=RBF_poly)
      print('\t\t\t %s保存到%s文件夹 \n' %(StrSMT, FileDir))  
    elif StrSMT=='RMTB':  #RMTB, data_dir用于保存模型
      mkdir(FileDir)
      StrFileDir.append(FileDir)
      sm = str1(print_global=False,data_dir=FileDir,xlimits=xlimits,order=RMTB_order,num_ctrl_pts=RMTB_num)  
      print('\t\t\t %s保存到%s文件夹 \n' %(StrSMT, FileDir))  
    elif StrSMT=='RMTC':  #RMTC, data_dir用于保存模型
      mkdir(FileDir)
      StrFileDir.append(FileDir)
      sm = str1(print_global=False,data_dir=FileDir,xlimits=xlimits)    
      print('\t\t\t %s保存到%s文件夹 \n' %(StrSMT, FileDir))    
    elif StrSMT=='IDW':  #IDW, data_dir用于保存模型
      mkdir(FileDir)
      StrFileDir.append(FileDir)
      sm = str1(print_global=False,data_dir=FileDir,p=IDW_p)
      print('\t\t\t %s保存到%s文件夹 \n' %(StrSMT, FileDir))  
    elif StrSMT[0]=='K':  #Kriging 族
      sm = str1(print_global=False,poly= 'quadratic',corr='squar_exp')
      StrFileDir.append(FileDir)
    else:  # LS, QP
      sm = str1(print_global=False)
      StrFileDir.append(FileDir)
    #EndIf
    
    #设置smt输入数据
    sm.set_training_values(X, Y)
    #训练smt
    sm.train()
    
    #保存除(IDW, RBF, RMTB, RMTC)之外的SMT模型到pkl文件，后续可以加载
    if (StrSMT[0]!='I' and StrSMT[0]!='R'):
       with open(FileDir, "wb") as f:
         pickle.dump(sm, f)
       print('\t\t\t %s保存到%s文件 \n' %(StrSMT, FileDir))  
    #EndIf
    print('\t\t\t 校核SMT模型 \n')  
    #根据检测数据的输入X1，预测
    Y1p=sm.predict_values(X1)
    ymin=min(Y1)*ratio
    ymax=max(Y1)*ratio
    
    #判断Kriging预测值与Abaqus计算值之间的差别：相对均方根误差:均方根误差/真实值的平均值
    RMSE=np.sqrt(np.sum((1-Y1p/Y1)**2)/len(Y1))
    RMSEMin.append(RMSE)
    
    StrRMSE=Object2+StrVarOut+'SMT模型校核的相对RMSE='+'%6.4f' %(RMSE*1e3) +'‰('+StrSMT+'方法)'
    StrTitle=StrRMSE
    print('\t\t\t 对校核结果画图 \n')  
    #ratio=[1,1e3]
    StrYlabels=Strs[1]#'最大Mises应力 [MPa]',  '最大Mises应变 [‰]', '第1节径振型固有频率[Hz]'
    FigName='Fig-'+StrType+'-RMSE-'+ StrVar +'-'+ StrSMT +'.png'
    #['Fig-Static-RMSE-SMax.png','Fig-Static-RMSE-EMax.png']
    fig = plt.figure()  # 定义一个图像窗口
    plt.plot(ratio*Y1p,'o',markersize=4,color='red',markerfacecolor='none')
    plt.plot(ratio*Y1, '.',markersize=2,color='blue')
    plt.xlabel('抽样点编号')
    plt.ylabel(StrYlabels)
    plt.legend(["预测值" , "真实值" ] ,loc='best')  
    plt.title(StrTitle)
    plt.ylim(ymin-(ymax-ymin)*0.05, ymax+(ymax-ymin)*0.05)
    plt.savefig(FigName, bbox_inches='tight')
    plt.close()
    FigList.append(FigName)
    FigCaption.append(StrTitle)
    
    #保存检测结果到txt文件
    print('\t\t\t 保存校核结果到txt文件 \n')
    XX=np.arange(1,len(Y1)+1,1)
    data=np.c_[XX,Y1p,Y1]
    TxtName='KrigingOut-'+StrType+'-'+StrVar+'-'+StrSMT+'.txt'
    TabCaption=StrTitle
    TabHead=['样本编号']+ ['%s方法预测值[%s]' %(StrSMT,StrUnit),'Abaqus值[%s]' %(StrUnit)]
    f = open(TxtName, 'w', encoding='utf-8-sig') # 打开文件， 用'w'写文件
    f.write('#%s\n'   %(TabCaption))
    f.write('#    %s    ' %(TabHead[0]))
    f.write('  %s       %s\n' %(TabHead[1],TabHead[2]))
    
    np.savetxt(f, X=data, delimiter=' ', encoding='utf-8-sig')
    f.close()
    TxtList.append(TxtName)
    TxtCaption.append(TabCaption)
    TxtHead.append(TabHead)
    tn = timeit.default_timer()
    print('\t\t\t 完成！耗时:',str(tn-t0),'s.\n')
  #EndFor
    
  #查找误差最小的SMT方法
  index=np.argmin(RMSEMin)
  StrSMT=SMTModels[index]
  str1=getattr(smt.surrogate_models, StrSMT)
  str2='%s%s的最佳SMT模型为:%s方法' %(Object2,StrVarOut,StrSMT)
  OptSMT_str=[str2]
  print('\t %s \n' %(str2))  
  FileDir = StrFileDir[index]
  t0 = timeit.default_timer()

  #加载之前保存的SMT模型，并返回到主函数中。
  if (StrSMT[0]!='I' and StrSMT[0]!='R'):#加载除(IDW, RBF, RMTB, RMTC)之外的SMT模型
    with open(FileDir, "rb") as f:
      sm = pickle.load(f)
  else:
    if StrSMT=='RBF':#加载(IDW, RBF, RMTB, RMTC)中的SMT模型
      sm = str1(print_global=False,data_dir=FileDir,d0=RBF_d0,poly_degree=RBF_poly)
    elif StrSMT=='RMTB':  #RMTB, data_dir用于保存pkl文件
      sm = str1(print_global=False,data_dir=FileDir,xlimits=xlimits,order=RMTB_order,num_ctrl_pts=RMTB_num)  
    elif StrSMT=='RMTC':  #RMTC, data_dir用于保存pkl文件
      sm = str1(print_global=False,data_dir=FileDir,xlimits=xlimits)  
    elif StrSMT=='IDW':#加载(IDW, RBF, RMTB, RMTC)中的SMT模型
      sm = str1(print_global=False,data_dir=FileDir,p=IDW_p)
    sm.set_training_values(X, Y) #输入参数与之前相同
    sm.train()  #实际并未训练，直接加载
  tn = timeit.default_timer()
  print('\t 完成！耗时:',str(tn-t0),'s.\n')
  ####用于再次校核
  ##Y1p=sm.predict_values(X1)
  ##RMSE=np.sqrt(np.sum((1-Y1p/Y1)**2)/len(Y1))
  ##print(RMSE)
  return sm,OptSMT_str,FigList,FigCaption,TxtList,TxtCaption,TxtHead

#EndDef

def StressStrain(Exec,Num1,FileGauss,FileAbaqus1,FileAbaqus2):
  ######################
  # 构建SMT模型
  ######################
  
  #读取Abaqus计算结果1，用于构建kriging模型
  print('\t 读取Abaqus计算得到的最大Mises应力和Mises应变 \n')
  print('\t\t 读取用于构建SMT的数据 \n')
  data =np.loadtxt(FileAbaqus1, dtype=float, skiprows=1,encoding='utf-8-sig')
  X0  =data[:,2:5]#密度、弹模、转速
  Y0  =data[:,0:2]#最大Mises应力和应变  
  print('\t\t 共有%d个样本点 \n' %(len(X0[:,0])))
  
  print('\t\t 读取用于校核SMT的样本数据 \n')
  data =np.loadtxt(FileAbaqus2, dtype=float, skiprows=1,encoding='utf-8-sig')
  X1   =data[:,2:5]#密度、弹模、转速
  Y1   =data[:,0:2]#最大Mises应力和应变
  print('\t\t 共有%d个样本点 \n' %(len(X1[:,0])))
  
  ##处理校核SMT模型的样本，用于生成docx报告
  TxtName='KrigingOut-Samples.txt'
  TabCaption=Object2+'校核SMT模型的样本'
  TabHead=['样本编号','密度 [kg/m^3]','弹性模量 [GPa]','转速 [r/min] ']+\
    ['最大Miese应力 [MPa]','最大Miese应变']
  f = open(TxtName, 'w', encoding='utf-8-sig') # 打开文件， 用'w'写文件
  f.write('#%s\n'   %(TabCaption))
  f.write('#    %s    ' %(TabHead[0]))
  f.write('  %s       %s       %s' %(TabHead[1],TabHead[2],TabHead[3]))
  f.write('  %s       %s\n' %(TabHead[4],TabHead[5]))
  XX=np.arange(1,len(Y1[:,0])+1,1)
  data=np.c_[XX,X1,Y1]
  np.savetxt(f, X=data, delimiter=' ', encoding='utf-8-sig')
  f.close()
  TxtList0=[TxtName]
  TxtCaption0=[TabCaption]
  TxtHead0=[TabHead]
  
  
  #构建最大Mises应力的SMT模型，从几个SMT中搜索最优模型
  print('\n\t 搜索最大Mises应力的最佳SMT模型\n') 
  Strs=['SMax','最大Mises应力 [MPa]']
  ratio=1
  SMax =np.c_[Y0[:,0]]
  SMax1=np.c_[Y1[:,0]]
  sm1,OptSMT_str1,FigList1,FigCaption1,TxtList1,TxtCaption1,TxtHead1=\
  OptSMT(X0,SMax,X1,SMax1,Strs,ratio)
  
  #构建最大Mises应变的SMT模型，从几个SMT中搜索最优模型
  print('\n\t 搜索最大Mises应变的最佳SMT模型\n') 
  Strs=['EMax','最大Mises应变 [‰]']
  ratio=1e3
  EMax =np.c_[Y0[:,1]]
  EMax1=np.c_[Y1[:,1]]
  sm2,OptSMT_str2,FigList2,FigCaption2,TxtList2,TxtCaption2,TxtHead2=\
  OptSMT(X0,EMax,X1,EMax1,Strs,ratio)
  
  ##疲劳可靠性计算抽样
  print('\t 抽样数目%3.2e组 \n' %Num1)
  _Text=Object2+Exec
  XR,FigList3,FigCaption3=ReliabSampling(Num1,FileGauss,_Text)
  
  #Kriging预测抽样样本的最大Mises应力
  print('\t Kriging预测抽样样本的最大Mises应力 \n')
  #YR=np.zeros([XR.shape[0],2])
  temp1=sm1.predict_values(XR)
  
  #Kriging预测抽样样本的最大Mises应变
  print('\t Kriging预测抽样样本的最大Mises应变 \n')
  temp2=sm2.predict_values(XR)
  YR=np.c_[temp1,temp2]
  OptOut_str=OptSMT_str1+OptSMT_str2
  FigList   =FigList1   +FigList2   +FigList3
  FigCaption=FigCaption1+FigCaption2+FigCaption3
  TxtList   =TxtList0   +TxtList1   +TxtList2
  TxtCaption=TxtCaption0+TxtCaption1+TxtCaption2
  TxtHead   =TxtHead0   +TxtHead1   +TxtHead2
  return XR,YR,OptOut_str,FigList,FigCaption,TxtList,TxtCaption,TxtHead
#EndDef

def StressReliability(YR,FileGauss):
  FigList=[]
  FigCaption=[]
  #读取GH4169的抗拉强度分布特性
  data = np.loadtxt(FileGauss,skiprows=8, comments='#',encoding='utf-8-sig')
  mu1=data[0] #data[0]是平均值
  sigma1=data[0]*data[1] #data[1]是标准差

  #确定抗拉强度分布
  x1_min=mu1-5.0*sigma1
  x1_max=mu1+5.0*sigma1
  x1 = np.linspace(x1_min,x1_max,1000)
  y1 = st.norm.pdf(x1,mu1,sigma1)

  print('\t 绘制最大Mises应力概率分布图 \n') 
  fig = plt.figure()
  #最大Mises应力概率分布
  dist=sns.distplot(YR,hist=True,kde=False, fit=st.norm,hist_kws={'facecolor':'gray','alpha':0.3}, 
  fit_kws={'lw':0.5,'color':'purple', 'label':'正态分布拟合','linestyle':'-'})
  
  plt.plot(x1,y1,color='red', linewidth=0.5)
  plt.fill_between(x1,y1,facecolor='gray',alpha=0.3)
  
  StrTitle=Object2+"最大Mises应力分布"
  plt.xlabel('应力 [MPa]')
  plt.ylabel('概率密度')
  plt.title(StrTitle)  
  plt.legend(['最大Mises应力', '抗拉强度'] ,loc='best')
  FigName='Fig-Static-Dist.png'
  plt.savefig(FigName, bbox_inches='tight')
  plt.close()
  FigList.append(FigName)
  FigCaption.append(StrTitle)


  #用norm分布拟合最大Mises应力数据，返回norm分布参数
  (mu2, sigma2) = st.norm.fit(YR)
  
  beta=abs(mu1-mu2)/np.sqrt(sigma1**2+sigma2**2)
  R_S=st.norm.cdf(beta)#静强度可靠度

  print('\n保存强度可靠度到txt文本\n')
  TxtName='KrigingOut-Static-Reliability.txt'
  f = open(TxtName, 'w', encoding='utf-8-sig') # 打开文件， 用'w'写文件
  f.write('#静强度可靠度\n')
  f.write('%f\n' %(R_S))
  f.close() 

  return R_S,FigList,FigCaption
#EndDef

def FrequencyReliability(IFreq,X,Y,X1,Y1,delta,Num1,FileGauss):

  #第1列为平均值，第2列为变异系数
  data = np.loadtxt(FileGauss, comments='#',encoding='utf-8-sig')
  rpm_avg=data[2,0]
  #转速变异系数
  rpm_cv=data[2,1]
  #转速标准差
  rpm_std=rpm_avg*rpm_cv
  

  print('处理第%d节径振型固有频率 \n' %(IFreq))

  Strs=['Mode%d' %(IFreq),'第%d节径振型固有频率[Hz]' %(IFreq)]
  ratio=1
  sm,OptSMT_str,FigList1,FigCaption1,TxtList,TxtCaption,TxtHead=OptSMT(X,Y,X1,Y1,Strs,ratio)
  ##振动可靠性计算抽样
  print('\t 振动可靠性计算抽样 \n')
  _Text=Object2+'第%d节径振动可靠性计算' %(IFreq)
  XR,FigList2,FigCaption2=ReliabSampling(Num1,FileGauss,_Text)

  FigList=FigList1+FigList2
  FigCaption=FigCaption1+FigCaption2
  
  #Kriging预测抽样样本的固有频率
  print('\t Kriging预测抽样样本的固有频率 \n')
  YR=sm.predict_values(XR)#样本固有频率 ω_n
  print('\t 三重点共振可靠度 \n')
  Freq_1=np.c_[XR[:,2]]/60.0#1倍转动频率 ω 
  Error=YR/Freq_1-IFreq#n节径共振时的三重点共振条件:ω_n/ω=k=n。IFreq=n。若error=0，表示共振
  Num2  = np.sum(Error > 0.01)#统计非零元素数目
  R_tri = Num2/Num1#非零元素比例：未共振的可靠度。
  #print(R_tri)
  
  #用norm分布拟合数据，返回norm分布参数
  (mu, sigma) = st.norm.fit(YR)
  
  #共振避开率可靠度，采用应力干涉法
  #delta=0.05#避开率Δ
  #高于(1+Δ)倍k次谐波的正态分布
  mu2   =(1+delta)*IFreq*rpm_avg/60.0
  sigma2=(1+delta)*IFreq*rpm_std/60.0
  #低于(1-Δ)倍k次谐波的正态分布
  mu1   =(1-delta)*IFreq*rpm_avg/60.0
  sigma1=(1-delta)*IFreq*rpm_std/60.0
  
  beta1=abs(mu-mu1)/np.sqrt(sigma**2+sigma1**2)
  beta2=abs(mu-mu2)/np.sqrt(sigma**2+sigma2**2)
  
  print('\t 第%d节径振型固有频率概率分布图 \n' %(IFreq)) 
  fig = plt.figure()
  #根据实际数据画疲劳寿命概率分布
  dist=sns.distplot(YR,hist=True,kde=False, fit=st.norm,hist_kws={'facecolor':'gray','alpha':0.3}, 
  fit_kws={'lw':0.5,'color':'purple', 'label':'正态分布拟合','linestyle':'-'})

  #根据norm拟合各节径振型固有频率概率分布
  x0 = np.linspace(min(YR),max(YR),1000)
  loc=mu#平均值
  scale=sigma#标准差
  y0 = st.norm.pdf(x0,loc,scale)
  
  #根据转速的正态分布，确定高于(1+Δ)倍k次谐波的分布
  x1_min=mu1-5.0*sigma1
  x1_max=mu1+5.0*sigma1
  x1 = np.linspace(x1_min,x1_max,1000)
  y1 = st.norm.pdf(x1,mu1,sigma1)
  
  x2_min=mu2-5.0*sigma2
  x2_max=mu2+5.0*sigma2
  x2 = np.linspace(x2_min,x2_max,1000)
  y2 = st.norm.pdf(x2,mu2,sigma2)
  plt.plot(x1,y1,color='red', linewidth=0.5)
  plt.plot(x2,y2,color='blue',linewidth=0.5)
  plt.fill_between(x1,y1,facecolor='gray',alpha=0.3)
  plt.fill_between(x2,y2,facecolor='gray',alpha=0.3)
  StrTitle='%s第%d节径振型固有频率概率分布' %(Object2,IFreq)
  plt.xlabel('频率 [Hz]')
  plt.ylabel('概率密度')
  plt.title(StrTitle)  
  plt.legend(['%d节径振型固有频率' %(IFreq), '%d次谐波频率×%3.2f' %(IFreq,1-delta),
  '%d次谐波频率×%3.2f' %(IFreq,1+delta) ] ,loc='best')
  FigName='Fig-Frequency-Dist-Mode'+'%d' %(IFreq)+'.png'
  plt.savefig(FigName, bbox_inches='tight')
  plt.close()
  FigList.append(FigName)
  FigCaption.append(StrTitle)
  
  #根据标准正态分布的累积密度，计算避开k次谐波频率可靠度
  R_delta1=st.norm.cdf(beta1)#低于(1-Δ)倍k次谐波的可靠度
  R_delta2=st.norm.cdf(beta2)#高于(1+Δ)倍k次谐波的可靠度

  #节径振动避开率可靠度
  R_delta=R_delta1*R_delta2
  return R_tri,R_delta,OptSMT_str,FigList,FigCaption,TxtList,TxtCaption,TxtHead
  
#EndDef

def FatigueReliability(XR,YR,FileMater):
  ##读取材料疲劳参数
  print('\t 读取材料疲劳参数 \n')  
  FuncReadFatigueP(FileMater)
  
  Num1=XR.shape[0]#样本数量
  
  ##构建疲劳寿命计算的参数矩阵
  print('\t 构建计算疲劳寿命的参数矩阵 \n') 
  E0=XR[:,1]*1e3#弹性模量[MPa]
  sigma0  =YR[:,0]#最大Mises应力[MPa]
  epsilon0  =YR[:,1]/2#Mises应变幅
  sigma_f0  =np.ones(Num1)*sigma_f  #疲劳强度系数[MPa]
  b0    =np.ones(Num1)*b    #疲劳强度指数
  epsilon_f0=np.ones(Num1)*epsilon_f#疲劳塑性系数
  c0    =np.ones(Num1)*c    #疲劳塑性指数
  
  ##疲劳寿命计算(多核并行)
  print('\t 疲劳寿命计算(%d核并行)，样本数目=%3.2e （耗时略多）\n' %( Num_CPU-1,Num1)) 
  items=zip(E0,epsilon0,sigma0,sigma_f0,b0,epsilon_f0,c0)#封装数据
  p = mp.Pool(Num_CPU-1)
  t0 = timeit.default_timer()
  NFatig = p.map(do_something, items)
  p.close()
  p.join()
  tn   = timeit.default_timer()
  print('\t\t 完成！耗时:',str(tn-t0),'s.\n')
  
  np.savetxt('Fatigue.txt',NFatig,fmt="%d",delimiter=" ")
  
  FigList,FigCaption,TxtList,TxtCaption,TxtHead=OptFit(NFatig)
  
  return FigList,FigCaption,TxtList,TxtCaption,TxtHead
#EndDef

#对疲劳寿命分布进行拟合
def OptFit(data):
  FigList=[]
  FigCaption=[]
  TxtList=[]
  TxtCaption=[]
  TxtHead=[]
  print('\n采用Fitter对疲劳寿命分布进行拟合（耗时略多）\n')
  #偏态分布，x最大值太大了，需要截取:根据累积概率密度曲线，取接近最大累积密度的某点(99.9%或99.8%)为x坐标轴最大值
  fig1 = plt.figure()
  x1,y1=sns.kdeplot(data,cumulative=True).get_lines()[0].get_data()
  #plt.show()
  plt.close()
  
  y1max=max(y1)
  index=np.where(y1>y1max*0.995)
  xmax=x1[index[0][0]]
  xmin=int(min(data))
  
  #绘制直方图及核密度估计曲线
  fig = plt.figure()
  #画直方图，核密度
  sns.distplot(data,norm_hist=True,kde=False,bins=500,color="y")
  sns.kdeplot(data,clip=[0,xmax],cumulative=False)
  StrTitle='疲劳寿命分布直方图及核密度估计曲线'
  FigName='Fig-Fatigue-HistKde.png'
  plt.xlabel('疲劳循环次数')
  plt.ylabel('概率密度')
  plt.title(StrTitle)
  plt.xlim(0, xmax)
  plt.savefig(FigName, bbox_inches='tight')
  #plt.show()
  plt.close()
  FigList.append(FigName)
  FigCaption.append(StrTitle)
  
  #使用Fitter拟合概率分布核密度曲线
  #https://docs.scipy.org/doc/scipy/reference/stats.html
  #https://www.zhihu.com/question/284452109
  distributions=['genhalflogistic','logistic','genlogistic','gamma','invweibull','erlang','exponweib','fatiguelife','weibull_min','weibull_max']
  f = Fitter(data, xmin=0, xmax=xmax, distributions=distributions,timeout =30,bins=300,density=True)  # 创建Fitter类
  f.fit()  # 调用fit函数拟合分布
  df=f.summary(Nbest=5,plot=False)  # 前5个最佳拟合结果，存储到df中(dataframe格式)，不直接绘图
  #f.hist()
  #f.plot_pdf(Nbest=5, lw=2, method='sumsquare_error')
  
  #根据拟合参数，绘制前5个最佳核密度拟合曲线
  #用f.summary(plot=True)也能绘制，这里想试试getattr函数
  x1 = np.linspace(0,xmax,num=int(xmax))
  #https://blog.csdn.net/u014281392/article/details/75331570
  distributions_opt=df.index#提取前5个最佳拟合结果的分布名称
  
  fig = plt.figure()
  #只画直方图
  sns.distplot(data,norm_hist=True,kde=False,bins=500,color="y")
  #提取拟合分布的建模数据,并画图
  for index in range(len(distributions_opt)):
    str0=distributions_opt[index]
    param=f.fitted_param[str0]
    str1=getattr(st, str0)
    y1 = str1.pdf(x1,*param)
    plt.plot(x1, y1,linewidth=1.0)
  #EndFor
  StrTitle=Object2+'疲劳寿命分布直方图及核密度估计曲线'
  FigName='Fig-Fatigue-HistKdeFit.png'
  plt.xlabel('疲劳循环次数')
  plt.ylabel('概率密度')
  plt.title(StrTitle)
  plt.xlim(0, xmax)
  plt.legend(distributions_opt ,loc='best')
  plt.savefig(FigName, bbox_inches='tight')
  plt.close()
  FigList.append(FigName)
  FigCaption.append(StrTitle)
  
  ##累积概率密度曲线画图
  ##先画核密度概率分布曲线
  fig = plt.figure()
  #概率核密度曲线从0开始的第1段加密，将数据点保存为x1,y1
  x1,y1=sns.kdeplot(data,bw_adjust=0.1,clip=[0,xmin],gridsize=xmin+1,cumulative=False).get_lines()[0].get_data()
  
  #概率核密度曲线的第2段放粗一些，将数据点保存为x2,y2
  x2,y2=sns.kdeplot(data,bw_adjust=0.1,clip=[xmin+1,xmax],gridsize=300,cumulative=False).get_lines()[1].get_data()
  #plt.show()
  plt.close()
  
  #将两端数据拼接在一起
  x1=np.r_[x1,x2]#拼接数组，增加行数用np.r_[a,b]，增加列数用np.c_[a,b]
  y1=np.r_[y1,y2]
  
  
  ##再对核密度曲线求积分得到累积概率密度
  area=[]
  #https://www.cnpython.com/qa/1477247
  #https://vimsky.com/examples/usage/python-scipy.integrate.trapz.html
  #scipy.integrate.trapz(y, x=None, dx=1.0, axis=-1)
  for i in range(len(x1)):
    x=x1[0:i+1]#取x1的第0,1,2...,i列，注意没有取第i+1列
    y=y1[0:i+1]
    area_tmp=scipy.integrate.trapz(y,x)
    area.append(area_tmp)
  #EndFor
  ##根据累积概率密度曲线，确定可靠性曲线
  y2=1.0-np.c_[area]
  
  print('\n画疲劳寿命累积概率密度曲线\n')
  fig = plt.figure()
  plt.plot(x1,area,color='red',linewidth=1.0)
  StrTitle=Object2+'疲劳寿命累积概率密度曲线'
  FigName='Fig-Fatigue-Cumulative.png'
  plt.xlabel('疲劳循环次数')
  plt.ylabel('概率')
  plt.title(StrTitle)
  plt.xlim(0, xmax)
  plt.savefig(FigName, bbox_inches='tight')
  #plt.show()
  plt.close()
  FigList.append(FigName)
  FigCaption.append(StrTitle)
  
  print('\n画疲劳寿命可靠度曲线\n')
  fig = plt.figure()
  plt.plot(x1,y2,color='red',linewidth=1.0)
  StrTitle=Object2+'疲劳寿命可靠度曲线'
  FigName='Fig-Fatigue-Reliability.png'
  plt.xlabel('疲劳循环次数')
  plt.ylabel('可靠度')
  plt.title(StrTitle)
  plt.xlim(0, xmax)
  plt.savefig(FigName, bbox_inches='tight')
  plt.close()
  FigList.append(FigName)
  FigCaption.append(StrTitle)
  
  print('\n保存疲劳寿命概率密度曲线到txt文本\n')
  data=np.c_[x1,y1]
  TxtName='KrigingOut-Fatigue-KDE.txt'
  TabCaption=Object2+'疲劳寿命概率密度'
  TabHead=['循环次数','概率密度']
  comments='#%s \n#%s     %s' %(TabCaption,TabHead[0],TabHead[1])
  formats='%d      %12.10f'
  SaveTxt(data,TxtName,comments,formats)
  TxtList.append(TxtName)
  TxtCaption.append(TabCaption)
  TxtHead.append(TabHead)
  
  print('\n保存疲劳寿命累积概率密度曲线到txt文本\n')
  data=np.c_[x1,area]
  TxtName='KrigingOut-Fatigue-Cumulative.txt'
  TabCaption=Object2+'疲劳寿命累积概率密度'
  TabHead=['循环次数','累积概率密度']
  comments='#%s \n#%s     %s' %(TabCaption,TabHead[0],TabHead[1])
  formats='%d      %12.10f'
  SaveTxt(data,TxtName,comments,formats)
  TxtList.append(TxtName)
  TxtCaption.append(TabCaption)
  TxtHead.append(TabHead)
  
  
  print('\n保存疲劳寿命可靠度曲线到txt文本\n')
  data=np.c_[x1,y2]
  TxtName='KrigingOut-Fatigue-Reliability.txt'
  TabCaption=Object2+'疲劳寿命可靠度'
  TabHead=['循环次数','可靠度']
  comments='#%s \n#%s     %s' %(TabCaption,TabHead[0],TabHead[1])
  formats='%d      %12.10f'
  SaveTxt(data,TxtName,comments,formats)
  TxtList.append(TxtName)
  TxtCaption.append(TabCaption)
  TxtHead.append(TabHead)
  
  
  return FigList,FigCaption,TxtList,TxtCaption,TxtHead
#EndDef

def SaveTxt(data,TxtName,comments,formats):
  f = open(TxtName, 'w', encoding='utf-8-sig') # 打开文件， 用'w'写文件
  f.write('%s\n' %(comments))
  np.savetxt(f, X=data, fmt=formats, delimiter=' ', encoding='utf-8-sig')
  f.close()
#EndDEF


def InsertFig(Doc,Bookmark_Name,Bookmark_Text, Fig, Caption, index):
  # 插入图片
  paragraph = Doc.add_paragraph()
  run = paragraph.add_run()
  run.add_picture(Fig, width=Inches(5.4))
  paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
  paragraph.paragraph_format.space_after=Pt(0)  #段后0磅间距
  
  # 插入题注与标签
  #str0='图'
  #InsertCaption(doc, caption, index,str0)
  InsertCaption2(Doc,Bookmark_Name,Bookmark_Text,index,Caption)#题注含章标题编号
#EndDef

# 插入表格
#https://www.freesion.com/article/7261971185/
#https://blog.csdn.net/liaosen/article/details/121242200
def InsertTab(Doc,Bookmark_Name,Bookmark_Text,TxtName, TxtCaption, TxtHead, index):
  # 插入题注与标签
  #InsertCaption(Doc, TxtCaption, index,str0)
  InsertCaption1(Doc,Bookmark_Name,Bookmark_Text,index,TxtCaption)#题注不含章标题编号
  
  data = np.loadtxt(TxtName, comments='#',encoding='utf-8-sig')
  rows=data.shape[0]#原始数据的行数
  cols=data.shape[1]#原始数据的列数
  NumAll=rows*cols#原始数据总数
  
  if   cols==2:##如果只有2列，扩展成6列
    rmax=rows#原始数据最大行数
    rows=math.ceil(rows/3)#调整行数
    cols=cols*3#调整列数
    TxtHead=TxtHead+TxtHead+TxtHead
    data0=data#重组data数组
    data=np.zeros([rows,cols])
    #print(rmax,rows,cols)
    data[:,0:2]=data0[0:rows,:]
    data[:,2:4]=data0[rows:rows*2,:]
    data[0:rmax-rows*2,4:6]=data0[rows*2:rmax,:]
  elif cols==3:##如果只有3列，扩展成6列
    rmax=rows#原始数据最大行数
    rows=math.ceil(rows/2)#调整行数
    cols=cols*2#调整列数
    TxtHead=TxtHead+TxtHead
    data0=data#重组data数组
    data=np.zeros([rows,cols])
    #print(rmax,rows,cols)
    data[:,0:3]=data0[0:rows,:]
    data[0:rmax-rows,3:6]=data0[rows:rmax,:]
  elif cols==4:##如果只有4列，扩展成8列
    rmax=rows#原始数据最大行数
    rows=math.ceil(rows/2)#调整行数
    cols=cols*2#调整列数
    TxtHead=TxtHead+TxtHead
    data0=data#重组data数组
    data=np.zeros([rows,cols])
    #print(rmax,rows,cols)
    data[:,0:4]=data0[0:rows,:]
    data[0:rmax-rows,4:8]=data0[rows:rmax,:]
    
  #创建表格
  table=Doc.add_table(rows=rows+1,cols=cols,style='Table Grid')
  #将所有的单元格抽取出来，快速填表，修改单元字体大小，对齐方式
  #https://theprogrammingexpert.com/write-table-fast-python-docx/
  #https://cloud.tencent.com/developer/ask/sof/253016
  table_cells = table._cells
  
  for i_col in range(len(TxtHead)):
    cells=table.cell(0,i_col)
    cells.text="%s" %(TxtHead[i_col])
    cells.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cells.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cells.paragraphs[0].runs[0].font.size=Pt(7)
  #EndFor
  for i_row in range(len(data[:,0])):
    for i_col in range(len(data[0,:])):
      ii=(i_col+1)+i_row*data.shape[1]#计数器，处理到第n个数据
      i_cell=i_col + (i_row+1) * data.shape[1]#表格位置，注意第1行的编号包含进去了
      cells=table_cells[i_cell]
      cells.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
      if ii<=NumAll: #最后一行，可能有空列，需要根据数据总数判断
        cells.text="%8.6e" %(data[i_row,i_col])
      else:
        cells.text="-"
      cells.paragraphs[0].runs[0].font.size=Pt(7)
      cells.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
      #
    #EndFor
  #EndFor
  table.alignment=WD_TABLE_ALIGNMENT.CENTER
  
  #表后插入空白行
  paragraph = Doc.add_paragraph()
  paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
  paragraph.paragraph_format.space_after=Pt(0)  #段后0磅间距

#EndDef


#书签锚点起始
def Bookmark_Start(Tag,index, Bookmark_Name):
  bmrk = OxmlElement('w:bookmarkStart')
  bmrk.set(qn('w:id'), str(index))
  bmrk.set(qn('w:name'), Bookmark_Name)
  #return bmrk
  Tag.append(bmrk)
#EndDef

#书签锚点结束
def Bookmark_End(Tag,index):
  bmrk = OxmlElement('w:bookmarkEnd')
  bmrk.set(qn('w:id'), str(index))
  #return bmrk
  Tag.append(bmrk)
#EndDef

#创建题注域代码: "图{ SEQ 图 \* ARABIC }"   或者 
#创建题注域代码: "表{ SEQ 表 \* ARABIC }"
#不包含章标题编号
def CaptionField1(Tag,str0,index):
  
  #创建第1部分："图"  或者 "表"
  text = OxmlElement('w:r')
  text.text = str0
  text.set(qn('xml:space'), 'preserve')
  Tag.append(text)
  
  #创建第2部分："{"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r2.append(fldChar)
  Tag.append(r2)
  
  #创建第3部分：" SEQ 图 \* ARABIC "  或者 " SEQ 表 \* ARABIC "
  r2 = OxmlElement('w:r')
  instrText = OxmlElement('w:instrText')
  instrText.text = ' SEQ '+str0+' \* ARABIC '
  instrText.set(qn('xml:space'), 'preserve')
  r2.append(instrText)
  Tag.append(r2)

  #创建第4部分：
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'separate')
  r2.append(fldChar)
  Tag.append(r2)

  #创建第5部分：
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = str(index)
  r2.append(t)
  Tag.append(r2)
  
  #创建第6部分："}"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r2.append(fldChar)
  Tag.append(r2)
  ####
#EndDef


#创建题注域代码: "图{ STYLEREF 1 }-{ SEQ 图 \* ARABIC \s 1 }"   或者 
#创建题注域代码: "表{ STYLEREF 1 }-{ SEQ 表 \* ARABIC \s 1 }"
#  包含章标题编号
def CaptionField2(Tag,str0,index):
  
  #创建第1部分："图"  或者 "表"
  text = OxmlElement('w:r')
  text.text = str0
  text.set(qn('xml:space'), 'preserve')
  Tag.append(text)
  
  #创建第2部分：章标题编号
  #创建第2.1部分："{"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r2.append(fldChar)
  Tag.append(r2)
  
  #创建第2.2部分：" STYLEREF 1 "
  r2 = OxmlElement('w:r')
  instrText = OxmlElement('w:instrText')
  instrText.text = ' STYLEREF 1 \s'
  instrText.set(qn('xml:space'), 'preserve')
  r2.append(instrText)
  Tag.append(r2)

  #创建第2.3部分：
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'separate')
  r2.append(fldChar)
  Tag.append(r2)

  #创建第2.4部分：
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = str(index)
  r2.append(t)
  Tag.append(r2)
  
  #创建第2.5部分："}"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r2.append(fldChar)
  Tag.append(r2)
  ####
  
  #创建第3部分："-"
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = '-'
  r2.append(t)
  Tag.append(r2)
  
  #创建第4部分：题注编号
  #创建第4.1部分："{"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r2.append(fldChar)
  Tag.append(r2)
  
  #创建第4.2部分：" SEQ 图 \* ARABIC \s 1 "  或者 " SEQ 表 \* ARABIC \s 1 "
  r2 = OxmlElement('w:r')
  instrText = OxmlElement('w:instrText')
  instrText.text = ' SEQ '+str0+' \* ARABIC \s 1 '
  instrText.set(qn('xml:space'), 'preserve')
  r2.append(instrText)
  Tag.append(r2)

  #创建第4.3部分：
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'separate')
  r2.append(fldChar)
  Tag.append(r2)

  #创建第4.4部分：
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = str(index)
  r2.append(t)
  Tag.append(r2)
  
  #创建第4.5部分："}"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r2.append(fldChar)
  Tag.append(r2)
  ####
#EndDef


#在CaptionField2的题注中包含章标题编号
#下面插入章标题
def InsertHead(Doc):
  #创建段落
  paragraph = Doc.add_paragraph()
  paragraph.add_run().font.color.rgb = RGBColor(0,0,0)
  Tag=paragraph._p
  #段落居中
  #加个空格，防止在题注前回车，出现2行书签
  r1 = OxmlElement('w:pPr')
  fldChar = OxmlElement('w:pStyle')
  fldChar.set(qn('w:val'), 'Heading1')
  r1.append(fldChar)#创建1级标题
  
  r2 = OxmlElement('w:numPr')
  fldChar = OxmlElement('w:ilvl')
  fldChar.set(qn('w:val'), '0')
  r2.append(fldChar)
  fldChar = OxmlElement('w:numId')
  fldChar.set(qn('w:val'), '7')
  r2.append(fldChar)
  r1.append(r2)#创建自动编号
  
  #1级标题的格式
  fldChar = OxmlElement('w:ind')
  fldChar.set(qn('w:left'), '425')
  fldChar.set(qn('w:leftChars'), '0')
  fldChar.set(qn('w:hanging'), '425')
  fldChar.set(qn('w:firstLineChars'), '0')
  r1.append(fldChar)
  ##居中
  #fldChar = OxmlElement('w:jc')
  #fldChar.set(qn('w:val'), 'center')
  #r1.append(fldChar)
  
  #自动编号的颜色为黑色
  r2 = OxmlElement('w:rPr')
  fldChar = OxmlElement('w:color')
  fldChar.set(qn('w:val'), '000000')
  r2.append(fldChar)
  r1.append(r2)
  
  Tag.append(r1)
  
  
  r1 = OxmlElement('w:r')
  #1级标题颜色为黑色
  r2 = OxmlElement('w:rPr')
  fldChar = OxmlElement('w:color')
  fldChar.set(qn('w:val'), '000000')
  r2.append(fldChar)
  r1.append(r2)
  #1级标题的题目
  t = OxmlElement('w:t')
  t.text = "章"
  r1.append(t)
  Tag.append(r1)
  
  #另起一段
  Doc.add_paragraph()
  
#EndDef

# 插入题注与标签
#插入域，再插入书签，最后引用书签
#https://blog.csdn.net/weixin_39742727/article/details/109914739
#https://blog.csdn.net/igoizzz/article/details/117905625
def InsertCaption1(Doc,Bookmark_Name,Bookmark_Text,index,Caption):
  #创建段落
  paragraph = Doc.add_paragraph()
  #段落居中
  paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
  #加个空格，防止在题注前回车，出现2行书签
  text = OxmlElement('w:r')
  text.text = ' '
  text.set(qn('xml:space'), 'preserve')
  paragraph._p.append(text)
  
  #书签的位置锚点
  Tag=paragraph._p
  #书签开始
  Bookmark_Start(Tag,index, Bookmark_Name)
  #创建题注域代码
  CaptionField1(Tag,Bookmark_Text,index)#不包含章标题编号
  #书简结束
  Bookmark_End(Tag,index)
  
  #题注文字部分
  paragraph.add_run(" " + Caption)
  paragraph.paragraph_format.space_after=Pt(0)  #段后0磅间距
#EndDef

# 插入题注与标签
#插入域，再插入书签，最后引用书签
#https://blog.csdn.net/weixin_39742727/article/details/109914739
#https://blog.csdn.net/igoizzz/article/details/117905625
def InsertCaption2(Doc,Bookmark_Name,Bookmark_Text,index,Caption):
  #创建段落
  paragraph = Doc.add_paragraph()
  #段落居中
  paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
  #加个空格，防止在题注前回车，出现2行书签
  text = OxmlElement('w:r')
  text.text = ' '
  text.set(qn('xml:space'), 'preserve')
  paragraph._p.append(text)
  
  #书签的位置锚点
  Tag=paragraph._p
  #书签开始
  Bookmark_Start(Tag,index, Bookmark_Name)
  #创建题注域代码
  CaptionField2(Tag,Bookmark_Text,index)#包含章标题编号
  #书简结束
  Bookmark_End(Tag,index)
  
  #题注文字部分
  paragraph.add_run(" " + Caption)
  paragraph.paragraph_format.space_after=Pt(0)  #段后0磅间距
#EndDef

#创建交叉引用的域代码："{ REF Fig1 \h }"
def ReferenceField(Tag,Bookmark_Name,Bookmark_Text,index):
  #创建第1部分："{"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r2.append(fldChar)
  Tag.append(r2)
  
  #创建第2部分：" REF Fig1 \h "，其中Bookmark_Name是书签名称
  r2 = OxmlElement('w:r')
  instrText = OxmlElement('w:instrText')
  instrText.text = ' REF ' + Bookmark_Name + ' \h '
  r2.append(instrText)
  Tag.append(r2)
  
  #创建第3部分：
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'separate')
  r2.append(fldChar)
  Tag.append(r2)

  #创建第4部分：" REF Fig1 \h "，其中Bookmark_Text是 引用的书签内容
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = Bookmark_Text + str(index)
  r2.append(t)
  Tag.append(r2)

  #创建第5部分："}"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r2.append(fldChar)
  Tag.append(r2)
#EndDef

def CrossReference(Doc,_Text,Bookmark_Name,Bookmark_Text,index):
  #创建段落
  paragraph = Doc.add_paragraph()
  
  run = paragraph.add_run(_Text + "如")
  Tag = run._r
  ReferenceField(Tag,Bookmark_Name,Bookmark_Text,index)
  paragraph.add_run("所示。")
  paragraph.paragraph_format.space_after=Pt(0)  #段后0磅间距
#EndDef

#在页脚中添加页码
#https://stackoverflow.com/questions/50776715/setting-pgnumtype-property-in-python-docx-is-without-effect
def AddFooterNumber(run):
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'Page'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "Seq"
    fldChar2.append(t)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

#在文档中设置页脚
#https://baijiahao.baidu.com/s?id=1665454009794833226&wfr=spider&for=pc
def InsertPageNumber(Doc):
  footer = Doc.sections[0].footer # 获取第一个节的页脚
  footer.is_linked_to_previous = True  #编号续前一节
  paragraph = footer.paragraphs[0] # 获取页脚的第一个段落
  paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER#页脚居中对齐
  run_footer=paragraph.add_run() # 添加页脚内容
  AddFooterNumber(run_footer)
  font = run_footer.font
  font.name = 'Times New Roman'#新罗马字体
  font.size = Pt(10)#10号字体
  font.bold = False#加粗
#EndDef

#在docx中新建：图的交叉引用，表的交叉引用，图及其题注，表及其题注，
def AddFigTab(Exec,OptOut_str,FigList,FigCaption,TxtList,TxtCaption,TxtHead,DocxName):
  #path = os.getcwd()
  #print(path)
  
  ##输出图片
  Doc = Document()
  Doc.styles['Normal'].font.name = u'宋体'
  Doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
  Doc.styles['Normal'].font.size = Pt(10)# 小四
  Doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
  
  #设置文档页码
  InsertPageNumber(Doc)
  
  #插入章标题，自动编号，后续题注中使用
  #打开docx文件时，需要"更新域"2次
  InsertHead(Doc)
  
  #文字说明
  for i in range(0, len(OptOut_str)):
    paragraph = Doc.add_paragraph()
    run = paragraph.add_run(OptOut_str[i])
    paragraph.paragraph_format.space_after=Pt(0)  #段后0磅间距
  
  #建立图的交叉引用
  for i in range(0, len(FigList)):
    index=i+1
    Bookmark_Name=Object1+Exec+'Fig'+str(index)
    Bookmark_Text='图'
    CrossReference(Doc,FigCaption[i],Bookmark_Name,Bookmark_Text,index)
  
  #新建图及其题注
  for i in range(0, len(FigList)):
    index=i+1
    Bookmark_Name=Object1+Exec+'Fig'+str(index)
    Bookmark_Text='图'
    InsertFig(Doc,Bookmark_Name,Bookmark_Text,FigList[i],FigCaption[i],index)
    
  paths=DocxName+"-图.docx"
  Doc.save(paths)
  
  
  ##输出表格
  Doc = Document()
  Doc.styles['Normal'].font.name = u'宋体'
  Doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
  Doc.styles['Normal'].font.size = Pt(10)# 小四
  Doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
  
  #设置文档页码
  InsertPageNumber(Doc)
  
  #建立表的交叉引用
  for i in range(0, len(TxtList)):
    index=i+1
    Bookmark_Name=Object1+Exec+'Tab'+str(index)
    Bookmark_Text='表'
    CrossReference(Doc,TxtCaption[i],Bookmark_Name,Bookmark_Text,index)

  paragraph = Doc.add_paragraph()
  #段落居中
  #paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
  #run = paragraph.add_run("附表")
  #run.bold=True
  #run.font.name=u'宋体'
  #run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
  #run.font.color.rgb = RGBColor(0,0,0)
  
  #新建表及其题注
  for i in range(0, len(TxtList)):
    #TxtCaption是2维列表
    index=i+1
    Bookmark_Name=Object1+Exec+'Tab'+str(index)
    Bookmark_Text='表'
    InsertTab(Doc,Bookmark_Name,Bookmark_Text,TxtList[i],TxtCaption[i],TxtHead[i],index)
  
  paths=DocxName+"-表.docx"
  Doc.save(paths)
  return paths
#EndDef