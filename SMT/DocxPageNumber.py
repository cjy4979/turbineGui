import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Inches, Pt

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
def _add_field(run, field):
    """ add a field to a run
    """
    fldChar1 = OxmlElement('w:fldChar')  # creates a new element
    fldChar1.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = field

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = "Seq"
    fldChar2.append(t)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')


    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)



def _add_number_range(run):
    """ add a number range field to a run
    """
    _add_field(run, r'Page')

if __name__=='__main__':
  Doc = docx.Document()
  
  #new_section = Doc.add_section()  # Added new section for assigning different footer on each page.
  new_section=Doc.sections[0]
  sectPr = new_section._sectPr
  
  pgNumType = OxmlElement('w:pgNumType')
  pgNumType.set(qn('w:fmt'), 'decimal')
  pgNumType.set(qn('w:start'), '1')
  
  sectPr.append(pgNumType)
  
  new_footer = new_section.footer  # Get footer-area of the recent section in document
  new_footer.is_linked_to_previous = False  
  footer_para = new_footer.add_paragraph()  
  run_footer = footer_para.add_run("Your footer here")
  _add_number_range(run_footer)
  font = run_footer.font
  font.name = 'Arial'
  font.size = Pt(8)
  footer_para.paragraph_format.page_break_before = True
  
  Doc.save("Doc-PageNumber.docx")