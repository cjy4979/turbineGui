import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

#书签锚点起始
def Bookmark_Start(Tag,index, Bookmark_Name):
    bmrk = OxmlElement('w:bookmarkStart')
    bmrk.set(qn('w:id'), str(index))
    bmrk.set(qn('w:name'), Bookmark_Name)
    #return bmrk
    Tag.append(bmrk)
#EndDef

#书签锚点结束
def Bookmark_End(Tag,index):
    bmrk = OxmlElement('w:bookmarkEnd')
    bmrk.set(qn('w:id'), str(index))
    #return bmrk
    Tag.append(bmrk)
#EndDef

#创建题注域代码: "图{ SEQ 图 \* ARABIC }"   或者 
#创建题注域代码: "表{ SEQ 表 \* ARABIC }"
#不包含章标题编号
def CaptionField1(Tag,str0,index):
  
  #创建第1部分："图"  或者 "表"
  text = OxmlElement('w:r')
  text.text = str0
  text.set(qn('xml:space'), 'preserve')
  Tag.append(text)
  
  #创建第2部分："{"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r2.append(fldChar)
  Tag.append(r2)
  
  #创建第3部分：" SEQ 图 \* ARABIC "  或者 " SEQ 表 \* ARABIC "
  r2 = OxmlElement('w:r')
  instrText = OxmlElement('w:instrText')
  instrText.text = ' SEQ '+str0+' \* ARABIC '
  instrText.set(qn('xml:space'), 'preserve')
  r2.append(instrText)
  Tag.append(r2)

  #创建第4部分：
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'separate')
  r2.append(fldChar)
  Tag.append(r2)

  #创建第5部分：
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = str(index)
  r2.append(t)
  Tag.append(r2)
  
  #创建第6部分："}"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r2.append(fldChar)
  Tag.append(r2)
  ####
#EndDef


#创建题注域代码: "图{ STYLEREF 1 }-{ SEQ 图 \* ARABIC }"   或者 
#创建题注域代码: "表{ STYLEREF 1 }-{ SEQ 表 \* ARABIC }"
#  包含章标题编号
def CaptionField2(Tag,str0,index):
  
  #创建第1部分："图"  或者 "表"
  text = OxmlElement('w:r')
  text.text = str0
  text.set(qn('xml:space'), 'preserve')
  Tag.append(text)
  
  #创建第2部分：章标题编号
  #创建第2.1部分："{"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r2.append(fldChar)
  Tag.append(r2)
  
  #创建第2.2部分：" STYLEREF 1 "
  r2 = OxmlElement('w:r')
  instrText = OxmlElement('w:instrText')
  instrText.text = ' STYLEREF 1 \s'
  instrText.set(qn('xml:space'), 'preserve')
  r2.append(instrText)
  Tag.append(r2)

  #创建第2.3部分：
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'separate')
  r2.append(fldChar)
  Tag.append(r2)

  #创建第2.4部分：
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = str(index)
  r2.append(t)
  Tag.append(r2)
  
  #创建第2.5部分："}"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r2.append(fldChar)
  Tag.append(r2)
  ####
  
  #创建第3部分："-"
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = '-'
  r2.append(t)
  Tag.append(r2)
  
  #创建第4部分：题注编号
  #创建第4.1部分："{"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r2.append(fldChar)
  Tag.append(r2)
  
  #创建第4.2部分：" SEQ 图 \* ARABIC "  或者 " SEQ 表 \* ARABIC "
  r2 = OxmlElement('w:r')
  instrText = OxmlElement('w:instrText')
  instrText.text = ' SEQ '+str0+' \* ARABIC \s 1'
  instrText.set(qn('xml:space'), 'preserve')
  r2.append(instrText)
  Tag.append(r2)

  #创建第4.3部分：
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'separate')
  r2.append(fldChar)
  Tag.append(r2)

  #创建第4.4部分：
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = str(index)
  r2.append(t)
  Tag.append(r2)
  
  #创建第4.5部分："}"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r2.append(fldChar)
  Tag.append(r2)
  ####
#EndDef


#创建图或表的题注
def InsertCaption(Doc,Bookmark_Name,Bookmark_Text,index,Caption):
  #创建段落
  paragraph = Doc.add_paragraph()
  #段落居中
  paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
  #加个空格，防止在题注前回车，出现2行书签
  text = OxmlElement('w:r')
  #text.text = ' '
  #text.set(qn('xml:space'), 'preserve')
  paragraph._p.append(text)
  
  #书签的位置锚点
  Tag=paragraph._p
  #书签开始
  Bookmark_Start(Tag,index, Bookmark_Name)
  #创建题注域代码
  CaptionField2(Tag,Bookmark_Text,index)
  #书简结束
  Bookmark_End(Tag,index)
  
  paragraph.add_run(" " + Caption)
  
  ##另起一段
  #Doc.add_paragraph()
#EndDef

#创建交叉引用的域代码："{ REF Fig1 \h }"
def ReferenceField(Tag,Bookmark_Name,Bookmark_Text,index):
  #创建第1部分："{"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r2.append(fldChar)
  Tag.append(r2)
  
  #创建第2部分：" REF Fig1 \h "，其中Bookmark_Name是书签名称
  r2 = OxmlElement('w:r')
  instrText = OxmlElement('w:instrText')
  instrText.text = ' REF ' + Bookmark_Name + ' \h '
  r2.append(instrText)
  Tag.append(r2)
  
  #创建第3部分：
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'separate')
  r2.append(fldChar)
  Tag.append(r2)

  #创建第4部分：" REF Fig1 \h "，其中Bookmark_Text是 引用的书签内容
  r2 = OxmlElement('w:r')
  t = OxmlElement('w:t')
  t.text = Bookmark_Text + str(index)
  r2.append(t)
  Tag.append(r2)

  #创建第5部分："}"
  r2 = OxmlElement('w:r')
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r2.append(fldChar)
  Tag.append(r2)
#EndDef

def CrossReference(Doc,_Text,Bookmark_Name,Bookmark_Text,index):
  #创建段落
  paragraph = Doc.add_paragraph()
  
  run = paragraph.add_run(_Text + "如")
  Tag = run._r
  ReferenceField(Tag,Bookmark_Name,Bookmark_Text,index)
  paragraph.add_run("所示。")
#EndDef

#在CaptionField2的题注中包含章标题编号
#下面插入章标题
def InsertHead(Doc,Head_i):
  #创建段落
  paragraph = Doc.add_paragraph()
  paragraph.add_run().font.color.rgb = RGBColor(0,0,0)
  Tag=paragraph._p
  #段落居中
  #加个空格，防止在题注前回车，出现2行书签
  r1 = OxmlElement('w:pPr')
  fldChar = OxmlElement('w:pStyle')
  fldChar.set(qn('w:val'), 'Heading1')
  r1.append(fldChar)#创建1级标题
  
  r2 = OxmlElement('w:numPr')
  fldChar = OxmlElement('w:ilvl')
  fldChar.set(qn('w:val'), '0')
  r2.append(fldChar)
  fldChar = OxmlElement('w:numId')
  fldChar.set(qn('w:val'), '7')
  r2.append(fldChar)
  r1.append(r2)#创建自动编号
  
  #1级标题的格式
  fldChar = OxmlElement('w:ind')
  fldChar.set(qn('w:left'), '425')
  fldChar.set(qn('w:leftChars'), '0')
  fldChar.set(qn('w:hanging'), '425')
  fldChar.set(qn('w:firstLineChars'), '0')
  r1.append(fldChar)
  ##居中
  #fldChar = OxmlElement('w:jc')
  #fldChar.set(qn('w:val'), 'center')
  #r1.append(fldChar)
  
  #自动编号的颜色为黑色
  r2 = OxmlElement('w:rPr')
  fldChar = OxmlElement('w:color')
  fldChar.set(qn('w:val'), '000000')
  r2.append(fldChar)
  r1.append(r2)
  
  Tag.append(r1)
  
  
  r1 = OxmlElement('w:r')
  #1级标题颜色为黑色
  r2 = OxmlElement('w:rPr')
  fldChar = OxmlElement('w:color')
  fldChar.set(qn('w:val'), '000000')
  r2.append(fldChar)
  r1.append(r2)
  #1级标题的题目
  t = OxmlElement('w:t')
  t.text = "章"
  r1.append(t)
  Tag.append(r1)
  
  #另起一段
  Doc.add_paragraph()
  
  return Head_i+1
#EndDef

if __name__=='__main__':
  Doc = docx.Document()
  
  #Head = Doc.add_heading("",level=1)# 这里不填标题内容
  #Head.alignment = WD_ALIGN_PARAGRAPH.LEFT
  #run = Head.add_run("章")
  #run.font.name=u'宋体'
  #run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
  #run.font.color.rgb = RGBColor(0,0,0)
  #
  #Head._p.pPr.pStyle.set(qn('w:val'), u'Heading1')
  
  ##paragraph = Doc.add_paragraph("章")
  ##paragraph.style = Doc.styles['Heading 1']
  
  
  Head_i=0
  index=0
  Head_i=InsertHead(Doc,Head_i)

  #对交叉引用
  index=index+1
  Bookmark_Name='Fig'+str(Head_i)+'-'+str(index)
  Bookmark_Text='图'
  _Text="第1个图"
  CrossReference(Doc,_Text,Bookmark_Name,Bookmark_Text,index)

  #创建题注
  #index=index+1
  Bookmark_Name='Fig'+str(Head_i)+'-'+str(index)
  Bookmark_Text='图'
  Caption='随机参数分布'
  InsertCaption(Doc,Bookmark_Name,Bookmark_Text,index,Caption)

  #对交叉引用
  index=index+1
  Bookmark_Name='Fig'+str(Head_i)+'-'+str(index)
  Bookmark_Text='图'
  _Text="第1个图"
  CrossReference(Doc,_Text,Bookmark_Name,Bookmark_Text,index)

  #创建题注
  #index=index+1
  Bookmark_Name='Fig'+str(Head_i)+'-'+str(index)
  Bookmark_Text='图'
  Caption='随机参数分布'
  InsertCaption(Doc,Bookmark_Name,Bookmark_Text,index,Caption)

  ##另起一章
  Head_i=InsertHead(Doc,Head_i)
  
  #创建题注交叉引用
  index=index+1
  Bookmark_Name='Fig'+str(Head_i)+'-'+str(index)
  Bookmark_Text='图'
  _Text="第1个图"
  CrossReference(Doc,_Text,Bookmark_Name,Bookmark_Text,index)
  
  #创建题注
  #index=index+1
  Bookmark_Name='Fig'+str(Head_i)+'-'+str(index)
  Bookmark_Text='图'
  Caption='随机参数分布'
  InsertCaption(Doc,Bookmark_Name,Bookmark_Text,index,Caption)

  #创建题注交叉引用
  index=index+1
  Bookmark_Name='Fig'+str(Head_i)+'-'+str(index)
  Bookmark_Text='图'
  _Text="第1个图"
  CrossReference(Doc,_Text,Bookmark_Name,Bookmark_Text,index)
  
  #创建题注
  #index=index+1
  Bookmark_Name='Fig'+str(Head_i)+'-'+str(index)
  Bookmark_Text='图'
  Caption='随机参数分布'
  InsertCaption(Doc,Bookmark_Name,Bookmark_Text,index,Caption)

  Doc.save('Doc-BookMark.Docx')

#参考如下的
#https://www.yeahshecodes.com/python/using-python-Docx-create-word-Documents
#https://github.com/python-openxml/python-Docx/issues/403
